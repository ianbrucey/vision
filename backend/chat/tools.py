"""
Vision — SDK-Native Database Tools.

Each tool is an Agent SDK @tool handler registered on a per-session "vision"
MCP server. The server is created by create_vision_server(case_id), which
captures the case_id in each handler's closure — no context propagation needed.

Security:
  - case_id is hardcoded in the handler closure — the agent never sees or provides it
  - Every resource-ID tool verifies the resource belongs to the current case
  - Errors return is_error=True so the agent can react without crashing the loop
"""

from __future__ import annotations

import json
import re
from datetime import date, datetime
from typing import Any

from claude_agent_sdk import create_sdk_mcp_server, tool, ToolAnnotations

from core.db import connect, ensure_schema


# ---------------------------------------------------------------------------
# DB helpers
# ---------------------------------------------------------------------------

def _conn():
    """Return a connection in autocommit mode.

    Autocommit ensures write operations (INSERT, UPDATE, DELETE) are
    immediately persisted — no explicit COMMIT needed. The connection
    commits its implicit transaction (started by SET search_path) before
    enabling autocommit, since psycopg2 starts with autocommit=False.
    """
    ensure_schema()
    conn = connect()
    conn.commit()  # commit the implicit txn from SET search_path
    conn.autocommit = True
    return conn


def _query(sql: str, params: tuple | None = None) -> list[dict]:
    import psycopg2.extras

    conn = _conn()
    try:
        with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
            cur.execute(sql, params)
            return [dict(row) for row in cur.fetchall()]
    finally:
        conn.close()


def _query_one(sql: str, params: tuple | None = None) -> dict | None:
    rows = _query(sql, params)
    return rows[0] if rows else None


def _serialize(obj: Any) -> Any:
    if isinstance(obj, (date, datetime)):
        return obj.isoformat()
    if isinstance(obj, bytes):
        return obj.decode("utf-8", errors="replace")
    if isinstance(obj, dict):
        return {k: _serialize(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [_serialize(v) for v in obj]
    return obj


def _json(obj: Any) -> str:
    return json.dumps(_serialize(obj), indent=2, default=str)


def _result(data: dict) -> dict:
    return {"content": [{"type": "text", "text": _json(data)}]}


def _error(message: str) -> dict:
    return {"content": [{"type": "text", "text": message}], "is_error": True}


# ---------------------------------------------------------------------------
# Embedding helper
# ---------------------------------------------------------------------------

def _embed_query(query_text: str) -> str:
    """Embed query text via Mistral. Returns pgvector literal string."""
    from search.embed import _get_client, _truncate, EMBED_MODEL

    client = _get_client()
    resp = client.embeddings.create(
        model=EMBED_MODEL,
        inputs=[_truncate(query_text)],
    )
    vec = resp.data[0].embedding
    return "[" + ",".join(repr(float(v)) for v in vec) + "]"


# ---------------------------------------------------------------------------
# FAR constants — used by far_lookup and far_status tools
# ---------------------------------------------------------------------------

FAR_CASE_NAME = "FAR — Federal Acquisition Regulation"
FAR_ZIP_URL = (
    "https://www.acquisition.gov/sites/default/files/"
    "current/far/zip/html/FARHTML.zip"
)
KB_CASE_NAME = "Knowledge Base — Cross-Case Reference"


# ---------------------------------------------------------------------------
# Server factory — creates a fully scoped MCP server per session
# ---------------------------------------------------------------------------


def create_vision_server(case_id: int | None = None):
    """Create a vision MCP server with tools scoped to case_id.

    Each tool handler captures case_id from this function's closure.
    The agent never sees or provides a case_id — it's hardcoded per session.

    When case_id is None or 0 (system agent), tools operate across ALL cases —
    no case-level scoping.
    """

    _is_system = case_id is None or case_id == 0

    # -- verification helpers (capture case_id) -------------------------------

    def _doc_in_case(document_id: int) -> bool:
        if _is_system:
            return _query_one("SELECT 1 FROM documents WHERE id = %s", (document_id,)) is not None
        return _query_one(
            "SELECT 1 FROM documents WHERE id = %s AND case_id = %s",
            (document_id, case_id),
        ) is not None

    def _block_in_case(block_id: int) -> bool:
        if _is_system:
            return _query_one(
                """SELECT 1 FROM blocks b
                   JOIN documents d ON b.document_id = d.id
                   WHERE b.id = %s""",
                (block_id,),
            ) is not None
        return _query_one(
            """SELECT 1 FROM blocks b
               JOIN documents d ON b.document_id = d.id
               WHERE b.id = %s AND d.case_id = %s""",
            (block_id, case_id),
        ) is not None

    def _section_in_case(section_id: int) -> bool:
        if _is_system:
            return _query_one(
                """SELECT 1 FROM sections s
                   JOIN documents d ON s.document_id = d.id
                   WHERE s.id = %s""",
                (section_id,),
            ) is not None
        return _query_one(
            """SELECT 1 FROM sections s
               JOIN documents d ON s.document_id = d.id
               WHERE s.id = %s AND d.case_id = %s""",
            (section_id, case_id),
        ) is not None

    def _strategy_in_case(strategy_id: int) -> bool:
        return _query_one(
            "SELECT 1 FROM strategies WHERE id = %s AND case_id = %s",
            (strategy_id, case_id),
        ) is not None

    # -- Layer 1: Orientation ------------------------------------------------

    @tool(
        "list_workspaces",
        "List all workspaces for the current case. Workspaces scope "
        "drafts, views, and documents to a specific sub-matter within "
        "a case (e.g. 'RFP Response', 'Motion to Dismiss'). "
        "Use this to see what workspaces exist and which is active.",
        {},
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_workspaces(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_workspaces as _list_ws
                rows = _list_ws(conn, case_id)
            finally:
                conn.close()
            return _result({"count": len(rows), "workspaces": rows})
        except Exception as exc:
            return _error(f"list_workspaces failed: {exc}")

    @tool(
        "create_workspace",
        "Create a new workspace for the current case. Workspaces scope "
        "drafts, views, and documents to a specific sub-matter (e.g. "
        "'RFP Response', 'Credit Dispute', 'Motion to Dismiss'). "
        "Use when the user starts a new sub-matter or asks to organize "
        "work separately.",
        {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Workspace name. Be descriptive: 'RFP Response — VA T4NG', "
                    "'Motion to Dismiss — Smith v. Jones'.",
                },
                "description": {
                    "type": "string",
                    "description": "Optional: what this workspace is for.",
                },
            },
            "required": ["name"],
        },
    )
    async def create_workspace(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                import psycopg2.extras
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    cur.execute(
                        """INSERT INTO workspaces (case_id, name, phase, description, status)
                           VALUES (%s, %s, 'other', %s, 'active')
                           RETURNING id, name""",
                        (case_id, args["name"], args.get("description")),
                    )
                    row = dict(cur.fetchone())
            finally:
                conn.close()
            return _result({"workspace": row})
        except Exception as exc:
            return _error(f"create_workspace failed: {exc}")

    @tool(
        "get_case",
        "Complete overview of the current case: metadata, parties, "
        "allegations, documents list, events timeline, strategies, "
        "and workspaces. "
        "Use this FIRST to understand what you're working with before "
        "searching.",
        {},
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_case(args: dict[str, Any]) -> dict[str, Any]:
        try:
            case = _query_one(
                "SELECT * FROM cases WHERE id = %s", (case_id,)
            )
            if not case:
                return _error(f"Case {case_id} not found.")

            case["parties"] = _query(
                "SELECT * FROM parties WHERE case_id = %s ORDER BY name",
                (case_id,),
            )
            case["allegations"] = _query(
                """SELECT * FROM allegations
                   WHERE case_id = %s ORDER BY sort_order, allegation_id""",
                (case_id,),
            )
            case["documents"] = _query(
                """SELECT id, name, page_count, document_type,
                          ocr_status, source, created_at
                   FROM documents WHERE case_id = %s
                   ORDER BY created_at DESC""",
                (case_id,),
            )
            case["events"] = _query(
                """SELECT * FROM events
                   WHERE case_id = %s
                   ORDER BY event_date, sequence_hint""",
                (case_id,),
            )
            case["strategies"] = _query(
                """SELECT id, name, strategy_type, posture, jurisdiction,
                          status, objective, filing_deadline, created_at
                   FROM strategies WHERE case_id = %s
                   ORDER BY created_at DESC""",
                (case_id,),
            )
            case["workspaces"] = _query(
                """SELECT id, name, phase, description, status,
                          created_at, updated_at
                   FROM workspaces WHERE case_id = %s
                   ORDER BY created_at""",
                (case_id,),
            )
            case.pop("narrative", None)
            return _result({"case": case})
        except Exception as exc:
            return _error(f"get_case failed: {exc}")

    @tool(
        "list_documents",
        "List all documents in the current case. Optionally filter by "
        "document type. Use this to see what evidence is available before "
        "searching within it.",
        {
            "type": "object",
            "properties": {
                "document_type": {
                    "type": "string",
                    "description": "Filter by type: medical_record, contract, "
                    "transcript, correspondence, pleading, tax_return, etc.",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_documents(args: dict[str, Any]) -> dict[str, Any]:
        doc_type = args.get("document_type")
        try:
            params: list[Any] = [case_id]
            filt = ""
            if doc_type:
                filt = " AND document_type = %s"
                params.append(doc_type)
            rows = _query(
                f"""SELECT id, name, page_count, document_type,
                           ocr_status, source, created_at
                    FROM documents WHERE case_id = %s{filt}
                    ORDER BY created_at DESC""",
                tuple(params),
            )
            return _result({
                "count": len(rows),
                "filter": doc_type or "all",
                "documents": rows,
            })
        except Exception as exc:
            return _error(f"list_documents failed: {exc}")

    # -- Layer 2: Search -----------------------------------------------------

    @tool(
        "search_blocks",
        "Full-text keyword/phrase search across all documents in the case. "
        "Best for: specific names, dates, medical terms, legal phrases, "
        "or any exact wording. Returns ranked results with page numbers, "
        "document names, section titles, and highlighted text snippets. "
        "Use this when you know the words you're looking for. "
        "For thematic/conceptual queries, use semantic_search instead.",
        {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Keywords or phrase. "
                    "Examples: 'cardiac arrest', 'breach of contract', "
                    "'Dr. Chen', 'January 2024'.",
                },
                "document_id": {
                    "type": "integer",
                    "description": "Restrict to a specific document.",
                },
                "limit": {
                    "type": "integer",
                    "description": "Max results (default 20, max 50).",
                },
            },
            "required": ["query"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def search_blocks(args: dict[str, Any]) -> dict[str, Any]:
        query_text = args["query"]
        document_id = args.get("document_id")
        limit = min(args.get("limit", 20), 50)

        try:
            tsq = " & ".join(
                w + ":*" for w in query_text.strip().split() if w.isalnum()
            )
            if not tsq:
                return _error("Query contains no searchable terms.")

            params: list[Any] = [case_id, tsq]
            doc_filt = ""
            if document_id is not None:
                if not _doc_in_case(document_id):
                    return _error(
                        f"Document {document_id} not in case {case_id}."
                    )
                doc_filt = " AND b.document_id = %s"
                params.append(document_id)

            sql = f"""SELECT b.id, b.document_id, b.page, b.block_type,
                             b.section_id,
                             ts_rank(b.text_tsv,
                                     to_tsquery('english', %s)) AS rank,
                             ts_headline('english', b.text_content,
                                         to_tsquery('english', %s),
                                         'MaxWords=40, MinWords=10') AS snippet,
                             d.name AS document_name,
                             s.title AS section_title
                      FROM blocks b
                      JOIN documents d ON b.document_id = d.id
                      LEFT JOIN sections s ON b.section_id = s.id
                      WHERE d.case_id = %s
                        AND b.text_tsv @@ to_tsquery('english', %s)
                        {doc_filt}
                      ORDER BY rank DESC
                      LIMIT %s"""
            params_full = [case_id, tsq, tsq] + params[2:] + [limit]
            rows = _query(sql, tuple(params_full))

            return _result({
                "query": query_text,
                "count": len(rows),
                "results": rows,
            })
        except Exception as exc:
            return _error(f"search_blocks failed: {exc}")

    @tool(
        "semantic_search",
        "Concept/meaning search across all documents using vector "
        "embeddings. Best for: thematic queries where you don't know "
        "the exact wording, finding related concepts, or discovering "
        "relevant sections that use different terminology. "
        "Returns sections ranked by semantic similarity. "
        "For exact terms, use search_blocks. For both, use search_hybrid.",
        {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Concept or question. "
                    "Examples: 'standard of care for post-operative infection', "
                    "'evidence of prior knowledge of the defect'.",
                },
                "document_id": {
                    "type": "integer",
                    "description": "Restrict to a specific document.",
                },
                "limit": {
                    "type": "integer",
                    "description": "Max results (default 15, max 30).",
                },
            },
            "required": ["query"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def semantic_search(args: dict[str, Any]) -> dict[str, Any]:
        query_text = args["query"]
        document_id = args.get("document_id")
        limit = min(args.get("limit", 15), 30)

        try:
            vec_literal = _embed_query(query_text)
        except Exception as exc:
            return _error(
                f"Embedding failed — embeddings may not be generated yet. "
                f"Use search_blocks for keyword search instead. Error: {exc}"
            )

        try:
            params: list[Any] = [case_id, vec_literal]
            doc_filt = ""
            if document_id is not None:
                if not _doc_in_case(document_id):
                    return _error(
                        f"Document {document_id} not in case {case_id}."
                    )
                doc_filt = " AND s.document_id = %s"
                params.append(document_id)

            sql = f"""SELECT s.id AS section_id, s.document_id, s.title,
                             s.heading_level, s.page_start, s.page_end,
                             s.block_count, s.heading_chain,
                             ROUND((1 - (s.embedding <=> %s::vector))::numeric,
                                   4) AS similarity,
                             d.name AS document_name
                      FROM sections s
                      JOIN documents d ON s.document_id = d.id
                      WHERE d.case_id = %s
                        AND s.embedding IS NOT NULL
                        {doc_filt}
                      ORDER BY s.embedding <=> %s::vector
                      LIMIT %s"""
            params_full = (
                [case_id, vec_literal] + params[2:] + [vec_literal, limit]
            )
            rows = _query(sql, tuple(params_full))

            if not rows:
                return _result({
                    "query": query_text,
                    "count": 0,
                    "note": (
                        "No sections with embeddings found. Embeddings "
                        "may not have been generated yet for this case. "
                        "Try search_blocks for keyword search."
                    ),
                })

            return _result({
                "query": query_text,
                "count": len(rows),
                "results": rows,
            })
        except Exception as exc:
            return _error(f"semantic_search failed: {exc}")

    @tool(
        "search_hybrid",
        "Combined keyword + semantic search. Runs both modalities and "
        "merges results into a single ranked list with similarity scores "
        "and keyword snippets. Best for: important queries where missing "
        "a relevant result matters, or when unsure which modality fits.",
        {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Search query used for both modalities.",
                },
                "document_id": {
                    "type": "integer",
                    "description": "Restrict to a specific document.",
                },
                "limit": {
                    "type": "integer",
                    "description": "Max results (default 20, max 40).",
                },
            },
            "required": ["query"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def search_hybrid(args: dict[str, Any]) -> dict[str, Any]:
        query_text = args["query"]
        document_id = args.get("document_id")
        limit = min(args.get("limit", 20), 40)

        if document_id is not None and not _doc_in_case(document_id):
            return _error(f"Document {document_id} not in case {case_id}.")

        # Keyword pass
        kw_results: list[dict] = []
        try:
            tsq = " & ".join(
                w + ":*" for w in query_text.strip().split() if w.isalnum()
            )
            if tsq:
                params: list[Any] = [case_id, tsq]
                doc_filt = ""
                if document_id is not None:
                    doc_filt = " AND b.document_id = %s"
                    params.append(document_id)
                kw_results = _query(
                    f"""SELECT b.id AS block_id, b.document_id, b.page,
                               b.section_id,
                               ts_headline('english', b.text_content,
                                           to_tsquery('english', %s),
                                           'MaxWords=30, MinWords=5') AS snippet,
                               d.name AS document_name
                        FROM blocks b
                        JOIN documents d ON b.document_id = d.id
                        WHERE d.case_id = %s
                          AND b.text_tsv @@ to_tsquery('english', %s)
                          {doc_filt}
                        ORDER BY ts_rank(b.text_tsv,
                                         to_tsquery('english', %s)) DESC
                        LIMIT %s""",
                    tuple(
                        [case_id, tsq] + params[2:] + [tsq, limit * 2]
                    ),
                )
        except Exception:
            pass

        # Semantic pass
        sem_results: list[dict] = []
        try:
            vec_literal = _embed_query(query_text)
            params_s: list[Any] = [case_id, vec_literal]
            doc_filt_s = ""
            if document_id is not None:
                doc_filt_s = " AND s.document_id = %s"
                params_s.append(document_id)
            sem_results = _query(
                f"""SELECT s.id AS section_id, s.document_id, s.title,
                           s.heading_level, s.page_start, s.page_end,
                           s.block_count, s.heading_chain,
                           ROUND((1 - (s.embedding <=> %s::vector))::numeric,
                                 4) AS similarity,
                           d.name AS document_name
                    FROM sections s
                    JOIN documents d ON s.document_id = d.id
                    WHERE d.case_id = %s
                      AND s.embedding IS NOT NULL
                      {doc_filt_s}
                    ORDER BY s.embedding <=> %s::vector
                    LIMIT %s""",
                tuple(
                    [case_id, vec_literal] + params_s[2:] + [vec_literal, limit]
                ),
            )
        except Exception:
            pass

        # Merge
        kw_by_section: dict[int, list[dict]] = {}
        for r in kw_results:
            sid = r.get("section_id")
            if sid:
                kw_by_section.setdefault(sid, []).append(r)

        merged = []
        seen = set()
        for sec in sem_results:
            sid = sec["section_id"]
            if sid in seen:
                continue
            seen.add(sid)
            merged.append({
                "section_id": sid,
                "document_id": sec["document_id"],
                "document_name": sec["document_name"],
                "title": sec.get("title"),
                "page_start": sec.get("page_start"),
                "page_end": sec.get("page_end"),
                "block_count": sec.get("block_count"),
                "heading_chain": sec.get("heading_chain"),
                "similarity": sec.get("similarity"),
                "keyword_snippets": [
                    {
                        "block_id": b["block_id"],
                        "page": b["page"],
                        "snippet": b.get("snippet", ""),
                    }
                    for b in kw_by_section.get(sid, [])[:3]
                ],
            })

        for r in kw_results:
            sid = r.get("section_id")
            if sid and sid not in seen:
                seen.add(sid)
                merged.append({
                    "section_id": sid,
                    "document_id": r["document_id"],
                    "document_name": r["document_name"],
                    "keyword_only": True,
                    "keyword_snippets": [
                        {
                            "block_id": r["block_id"],
                            "page": r["page"],
                            "snippet": r.get("snippet", ""),
                        }
                    ],
                })

        return _result({
            "query": query_text,
            "count": len(merged[:limit]),
            "keyword_hits": len(kw_results),
            "semantic_hits": len(sem_results),
            "results": merged[:limit],
        })

    # -- Layer 3: Structure --------------------------------------------------

    @tool(
        "get_document_structure",
        "Section outline (table of contents) for a document. Shows "
        "heading hierarchy, page ranges, and block counts per section. "
        "Use this to understand a document's organization before "
        "reading specific sections.",
        {
            "type": "object",
            "properties": {
                "document_id": {
                    "type": "integer",
                    "description": "Document ID from get_case or list_documents.",
                },
            },
            "required": ["document_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_document_structure(args: dict[str, Any]) -> dict[str, Any]:
        document_id = args["document_id"]
        try:
            if not _doc_in_case(document_id):
                return _error(
                    f"Document {document_id} not in case {case_id}."
                )

            doc = _query_one(
                """SELECT id, name, page_count, document_type
                   FROM documents WHERE id = %s""",
                (document_id,),
            )
            sections = _query(
                """SELECT id, title, heading_level, page_start, page_end,
                          block_count, heading_chain
                   FROM sections WHERE document_id = %s
                   ORDER BY page_start, id""",
                (document_id,),
            )
            return _result({
                "document": doc,
                "section_count": len(sections),
                "sections": sections,
            })
        except Exception as exc:
            return _error(f"get_document_structure failed: {exc}")

    @tool(
        "search_sections",
        "Find sections by title text. Supports fuzzy matching — finds "
        "sections even when the title isn't exact. Use this to jump to "
        "a specific part of a document (e.g., 'find the operative "
        "report section').",
        {
            "type": "object",
            "properties": {
                "document_id": {
                    "type": "integer",
                    "description": "Document to search within.",
                },
                "title": {
                    "type": "string",
                    "description": "Title text to search for. Partial matches work.",
                },
            },
            "required": ["document_id", "title"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def search_sections(args: dict[str, Any]) -> dict[str, Any]:
        document_id = args["document_id"]
        title = args["title"]
        try:
            if not _doc_in_case(document_id):
                return _error(
                    f"Document {document_id} not in case {case_id}."
                )

            rows = _query(
                """SELECT id, title, heading_level, page_start, page_end,
                          block_count, similarity(title, %s) AS sim
                   FROM sections
                   WHERE document_id = %s AND title %% %s
                   ORDER BY sim DESC
                   LIMIT 20""",
                (title, document_id, title),
            )
            return _result({
                "document_id": document_id,
                "query": title,
                "count": len(rows),
                "sections": rows,
            })
        except Exception as exc:
            return _error(f"search_sections failed: {exc}")

    # -- Layer 4: Read -------------------------------------------------------

    @tool(
        "get_block_context",
        "Read a specific block with surrounding text on adjacent pages. "
        "Returns the target block plus neighbors for full context. "
        "ALWAYS use this after search to verify a match before citing it. "
        "Never cite from a search snippet alone — read in context first.",
        {
            "type": "object",
            "properties": {
                "block_id": {
                    "type": "integer",
                    "description": "Block ID from search_blocks results.",
                },
            },
            "required": ["block_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_block_context(args: dict[str, Any]) -> dict[str, Any]:
        block_id = args["block_id"]
        try:
            if not _block_in_case(block_id):
                return _error(f"Block {block_id} not in case {case_id}.")

            target = _query_one(
                "SELECT * FROM blocks WHERE id = %s", (block_id,)
            )
            if not target:
                return _error(f"Block {block_id} not found.")

            doc = _query_one(
                "SELECT id, name FROM documents WHERE id = %s",
                (target["document_id"],),
            )
            section = None
            if target["section_id"]:
                section = _query_one(
                    "SELECT id, title, heading_level, heading_chain "
                    "FROM sections WHERE id = %s",
                    (target["section_id"],),
                )

            neighbors = _query(
                """SELECT b.id, b.block_type, b.page, b.text_content,
                          b.datalab_id, s.title AS section_title
                   FROM blocks b
                   LEFT JOIN sections s ON b.section_id = s.id
                   WHERE b.document_id = %s
                     AND b.page BETWEEN %s AND %s
                   ORDER BY b.page, b.id""",
                (
                    target["document_id"],
                    max(0, (target["page"] or 0) - 1),
                    (target["page"] or 0) + 1,
                ),
            )
            return _result({
                "target": target,
                "document": doc,
                "section": section,
                "context_block_count": len(neighbors),
                "context": neighbors,
            })
        except Exception as exc:
            return _error(f"get_block_context failed: {exc}")

    @tool(
        "get_blocks_in_section",
        "Read all blocks within a specific section. Use after "
        "get_document_structure or search_sections to read an entire "
        "section's content. Returns blocks in page order.",
        {
            "type": "object",
            "properties": {
                "section_id": {
                    "type": "integer",
                    "description": "Section ID from get_document_structure, "
                    "search_sections, or semantic_search results.",
                },
                "block_types": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Optional: filter to specific block types "
                    "(e.g., ['Text'] to skip headers).",
                },
                "limit": {
                    "type": "integer",
                    "description": "Max blocks (default 100, max 200).",
                },
            },
            "required": ["section_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_blocks_in_section(args: dict[str, Any]) -> dict[str, Any]:
        section_id = args["section_id"]
        block_types = args.get("block_types")
        limit = min(args.get("limit", 100), 200)

        try:
            if not _section_in_case(section_id):
                return _error(
                    f"Section {section_id} not in case {case_id}."
                )

            section = _query_one(
                "SELECT id, title, heading_level, page_start, page_end, "
                "block_count FROM sections WHERE id = %s",
                (section_id,),
            )

            params: list[Any] = [section_id]
            type_filt = ""
            if block_types:
                type_filt = " AND block_type = ANY(%s)"
                params.append(block_types)
            params.append(limit)

            blocks = _query(
                f"""SELECT id, block_type, page, text_content, datalab_id
                    FROM blocks
                    WHERE section_id = %s{type_filt}
                    ORDER BY page, id
                    LIMIT %s""",
                tuple(params),
            )
            return _result({
                "section": section,
                "block_count_returned": len(blocks),
                "blocks": blocks,
            })
        except Exception as exc:
            return _error(f"get_blocks_in_section failed: {exc}")

    # -- Layer 5: Strategy ---------------------------------------------------

    @tool(
        "get_strategies",
        "List all legal strategy trees built for the current case. "
        "Each strategy models a legal claim. Use this to see what "
        "analysis exists before diving into a specific strategy.",
        {},
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_strategies(args: dict[str, Any]) -> dict[str, Any]:
        try:
            rows = _query(
                """SELECT id, name, strategy_type, posture, jurisdiction,
                          status, objective, filing_deadline, created_at
                   FROM strategies WHERE case_id = %s
                   ORDER BY created_at DESC""",
                (case_id,),
            )
            return _result({"count": len(rows), "strategies": rows})
        except Exception as exc:
            return _error(f"get_strategies failed: {exc}")

    @tool(
        "get_strategy_tree",
        "Complete recursive proposition tree for a strategy. Returns "
        "claims, elements, AND/OR gates, fact mappings, and current "
        "status per node. Use this to analyze a specific legal claim "
        "in detail.",
        {
            "type": "object",
            "properties": {
                "strategy_id": {
                    "type": "integer",
                    "description": "Strategy ID from get_strategies.",
                },
            },
            "required": ["strategy_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_strategy_tree(args: dict[str, Any]) -> dict[str, Any]:
        strategy_id = args["strategy_id"]
        try:
            if not _strategy_in_case(strategy_id):
                return _error(
                    f"Strategy {strategy_id} not in case {case_id}."
                )

            strategy = _query_one(
                "SELECT * FROM strategies WHERE id = %s", (strategy_id,)
            )
            if not strategy:
                return _error(f"Strategy {strategy_id} not found.")

            propositions = _query(
                """WITH RECURSIVE tree AS (
                       SELECT id, parent_proposition_id, proposition_type,
                              gate_type, party_id, label, proposition_text,
                              current_status, sort_order,
                              0 AS depth,
                              ARRAY[sort_order, id] AS path
                       FROM strategy_propositions
                       WHERE strategy_id = %s
                         AND parent_proposition_id IS NULL
                       UNION ALL
                       SELECT sp.id, sp.parent_proposition_id,
                              sp.proposition_type, sp.gate_type,
                              sp.party_id, sp.label, sp.proposition_text,
                              sp.current_status, sp.sort_order,
                              t.depth + 1,
                              t.path || sp.sort_order || sp.id
                       FROM strategy_propositions sp
                       JOIN tree t ON sp.parent_proposition_id = t.id
                       WHERE sp.strategy_id = %s
                   )
                   SELECT * FROM tree ORDER BY path""",
                (strategy_id, strategy_id),
            )
            strategy["propositions"] = propositions
            strategy["proposition_count"] = len(propositions)
            return _result({"strategy": strategy})
        except Exception as exc:
            return _error(f"get_strategy_tree failed: {exc}")

    # -- Layer 6: Drafting ---------------------------------------------------

    @tool(
        "list_drafts",
        "List all drafts for the current case. Returns id, name, document_type, "
        "status, block count, and timestamps. Does not return full content — "
        "use get_draft for that. Use this to see what drafts exist before "
        "creating or editing one.",
        {},
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_drafts(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_drafts as _list_drafts
                rows = _list_drafts(conn, case_id)
            finally:
                conn.close()
            return _result({"count": len(rows), "drafts": rows})
        except Exception as exc:
            return _error(f"list_drafts failed: {exc}")

    @tool(
        "get_draft",
        "Read a draft's full content including all blocks. Use this before "
        "editing a draft so you can see the current state and target specific "
        "block IDs.",
        {
            "type": "object",
            "properties": {
                "draft_id": {
                    "type": "integer",
                    "description": "Draft ID from list_drafts.",
                },
            },
            "required": ["draft_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_draft(args: dict[str, Any]) -> dict[str, Any]:
        draft_id = args["draft_id"]
        try:
            conn = _conn()
            try:
                from core.db import get_draft as _get_draft
                draft = _get_draft(conn, draft_id)
            finally:
                conn.close()
            if not draft:
                return _error(f"Draft {draft_id} not found.")
            if draft["case_id"] != case_id:
                return _error(f"Draft {draft_id} not in case {case_id}.")
            return _result({"draft": draft})
        except Exception as exc:
            return _error(f"get_draft failed: {exc}")

    @tool(
        "create_draft",
        "Create a new draft. The content is an array of blocks, each with "
        "a unique id (short string), a type, and text content. "
        "Use this to produce a structured document the user can review "
        "and edit in the Drafts tab.\n\n"
        "Block types:\n"
        "  section_heading    — Centered, bold, underlined section title\n"
        "  numbered_paragraph — Auto-numbered paragraph (1., 2., 3.)\n"
        "  list_item          — Letter-labeled item (a), (b), (c)\n"
        "  signature          — Signature block with top-border line\n\n"
        "Example content array:\n"
        '  [{"id":"h1","type":"section_heading","content":"BACKGROUND"},\n'
        '   {"id":"p1","type":"numbered_paragraph","content":"Vision is a..."}]',
        {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Display title for the draft.",
                },
                "document_type": {
                    "type": "string",
                    "enum": ["letter", "pleading", "contract", "memo", "other"],
                    "description": "Type of document.",
                },
                "content": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "type": {
                                "type": "string",
                                "enum": [
                                    "section_heading",
                                    "numbered_paragraph",
                                    "list_item",
                                    "signature",
                                ],
                            },
                            "content": {"type": "string"},
                        },
                        "required": ["id", "type", "content"],
                    },
                    "description": "Array of blocks in document order.",
                },
            },
            "required": ["name", "document_type", "content"],
        },
    )
    async def create_draft(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import insert_draft as _insert_draft
                draft_id = _insert_draft(
                    conn,
                    case_id=case_id,
                    name=args["name"],
                    document_type=args["document_type"],
                    content=args["content"],
                    created_by="agent",
                )
            finally:
                conn.close()
            return _result({
                "draft_id": draft_id,
                "name": args["name"],
                "document_type": args["document_type"],
                "block_count": len(args["content"]),
            })
        except Exception as exc:
            return _error(f"create_draft failed: {exc}")

    @tool(
        "update_draft",
        "Modify a draft — update its name, status, or replace specific "
        "blocks. For targeted edits, provide only the blocks that changed "
        "(they'll be matched by id). To replace the entire content, provide "
        "a full content array. Use get_draft first to see current block IDs.",
        {
            "type": "object",
            "properties": {
                "draft_id": {
                    "type": "integer",
                    "description": "Draft ID to update.",
                },
                "name": {
                    "type": "string",
                    "description": "New display title (optional).",
                },
                "status": {
                    "type": "string",
                    "enum": ["draft", "review", "final"],
                    "description": "New status (optional).",
                },
                "content": {
                    "type": "array",
                    "items": {
                        "type": "object",
                        "properties": {
                            "id": {"type": "string"},
                            "type": {
                                "type": "string",
                                "enum": [
                                    "section_heading",
                                    "numbered_paragraph",
                                    "list_item",
                                    "signature",
                                ],
                            },
                            "content": {"type": "string"},
                        },
                        "required": ["id", "type", "content"],
                    },
                    "description": "Full replacement content array (optional). "
                    "Replaces all blocks. Use for significant rewrites.",
                },
            },
            "required": ["draft_id"],
        },
    )
    async def update_draft(args: dict[str, Any]) -> dict[str, Any]:
        draft_id = args["draft_id"]
        try:
            conn = _conn()
            try:
                from core.db import get_draft as _get_draft
                draft = _get_draft(conn, draft_id)
            finally:
                conn.close()
            if not draft:
                return _error(f"Draft {draft_id} not found.")
            if draft["case_id"] != case_id:
                return _error(f"Draft {draft_id} not in case {case_id}.")

            kwargs = {}
            if "name" in args:
                kwargs["name"] = args["name"]
            if "status" in args:
                kwargs["status"] = args["status"]
            if "content" in args:
                kwargs["content"] = args["content"]

            if not kwargs:
                return _error("No fields to update.")

            conn = _conn()
            try:
                from core.db import update_draft as _update_draft
                updated = _update_draft(conn, draft_id, **kwargs)
            finally:
                conn.close()

            return _result({
                "draft_id": draft_id,
                "name": updated["name"],
                "block_count": len(updated.get("content", [])),
                "updated_at": str(updated.get("updated_at", "")),
            })
        except Exception as exc:
            return _error(f"update_draft failed: {exc}")

    # -- Layer 6.5: Workspace ------------------------------------------------

    @tool(
        "list_workspace_items",
        "List workspace items for the current case. Optionally filter by "
        "folder_id or file_type. Returns id, "
        "name, file_type, folder_id, document_type, status, block count, and "
        "timestamps. Does not return full content — use get_workspace_item "
        "for that.",
        {
            "type": "object",
            "properties": {
                "folder_id": {
                    "type": "integer",
                    "description": "Filter by folder ID.",
                },
                "file_type": {
                    "type": "string",
                    "enum": ["markdown", "structured_draft", "html", "json_view", "pdf"],
                    "description": "Filter by file type.",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_workspace_items(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_drafts as _list_drafts
                folder_id = args.get("folder_id")
                rows = _list_drafts(conn, case_id, folder_id=folder_id)
            finally:
                conn.close()
            file_type = args.get("file_type")
            if file_type is not None:
                rows = [r for r in rows if r.get("file_type") == file_type]
            return _result({"count": len(rows), "items": rows})
        except Exception as exc:
            return _error(f"list_workspace_items failed: {exc}")

    @tool(
        "get_workspace_item",
        "Read a workspace item's full content. The content structure depends "
        "on file_type:\n"
        "  markdown         — [{\"markdown\": \"# Title...\"}]\n"
        "  structured_draft — [{id, type, content}, ...]\n"
        "  html             — [{\"html\": \"<html>...\"}]\n"
        "  json_view        — {documentMetadata: {title, sourceId?, lastUpdated?}, "
        "views: [{viewType, title, description?, data}]}\n"
        "    viewType: 'table' | 'list' | 'cards' | 'chart'\n"
        "    table data: {headers: [string], rows: [{id, ...column}]}\n"
        "    list data:  {listStyle: 'checkbox'|'ordered'|'bullet', "
        "items: [{id, text, completed?, notes?}]}\n"
        "    cards data: {pairs: [{key, value, emphasis?}]}\n"
        "    chart data: {chartType: 'bar'|'line'|'pie', "
        "headers: [string], rows: [{id, ...column}]}  (same structure as table)\n"
        "    emphasis: 'default' | 'warning' | 'danger' | 'success' | 'info'\n"
        "Use this before editing so you can see the current state.",
        {
            "type": "object",
            "properties": {
                "item_id": {
                    "type": "integer",
                    "description": "Workspace item ID from list_workspace_items.",
                },
            },
            "required": ["item_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_workspace_item(args: dict[str, Any]) -> dict[str, Any]:
        item_id = args["item_id"]
        try:
            conn = _conn()
            try:
                from core.db import get_draft as _get_draft
                item = _get_draft(conn, item_id)
            finally:
                conn.close()
            if not item:
                return _error(f"Workspace item {item_id} not found.")
            if item["case_id"] != case_id:
                return _error(f"Workspace item {item_id} not in case {case_id}.")
            return _result({"item": item})
        except Exception as exc:
            return _error(f"get_workspace_item failed: {exc}")

    @tool(
        "create_workspace_item",
        "Create a new workspace item. Content structure depends on file_type:\n"
        "  markdown         — [{\"markdown\": \"# Title\\n\\nContent...\"}]\n"
        "  structured_draft — [{\"id\":\"b1\",\"type\":\"section_heading\","
        "\"content\":\"TITLE\"}, ...]\n"
        "  html             — [{\"html\": \"<html>...</html>\"}]\n"
        "  json_view        — {documentMetadata: {title}, "
        "views: [{viewType, title, data}]}  (direct object — NOT array-wrapped)\n\n"
        "For json_view, see the dynamic-views skill for the full envelope schema. "
        "For all other types, content must be wrapped in a single-element array.",
        {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Display name for the item.",
                },
                "file_type": {
                    "type": "string",
                    "enum": ["markdown", "structured_draft", "html", "json_view", "pdf"],
                    "description": "Type of content this item holds.",
                },
                "folder_id": {
                    "type": "integer",
                    "description": "Which folder ID to place the item in (optional).",
                },
                "document_type": {
                    "type": "string",
                    "enum": ["letter", "pleading", "contract", "memo",
                             "capability_statement", "other"],
                    "description": "Legal document type (optional, defaults to other).",
                },
                "content": {
                    "type": "object",
                    "description": "Content envelope matching file_type. "
                    "For json_view: a direct object {documentMetadata, views[]}. "
                    "For all other types: array-wrapped.",
                },
                "workspace_id": {
                    "type": "integer",
                    "description": "Workspace ID to scope this item to (optional). "
                    "If omitted, uses the default workspace. Use list_workspaces "
                    "to see available workspaces.",
                },
            },
            "required": ["name", "file_type", "content"],
        },
    )
    async def create_workspace_item(args: dict[str, Any]) -> dict[str, Any]:
        try:
            content_raw = args.get("content", {})
            file_type = args.get("file_type", "markdown")

            # json_view stores content as a direct object — never array-wrap it.
            if file_type == "json_view" and isinstance(content_raw, dict):
                content_to_store = content_raw
            elif isinstance(content_raw, list):
                content_to_store = content_raw
            elif isinstance(content_raw, dict):
                content_to_store = [content_raw]
            else:
                content_to_store = [{"raw": str(content_raw)}]

            conn = _conn()
            try:
                from core.db import insert_draft as _insert_draft
                item_id = _insert_draft(
                    conn,
                    case_id=case_id,
                    name=args["name"],
                    document_type=args.get("document_type", "other"),
                    content=content_to_store,
                    created_by="agent",
                    file_type=args["file_type"],
                    folder="artifacts",
                    workspace_id=args.get("workspace_id"),
                    folder_id=args.get("folder_id"),
                )
            finally:
                conn.close()
            return _result({
                "item_id": item_id,
                "name": args["name"],
                "file_type": args["file_type"],
                "folder_id": args.get("folder_id"),
                "block_count": len(content_to_store),
            })
        except Exception as exc:
            return _error(f"create_workspace_item failed: {exc}")

    @tool(
        "update_workspace_item",
        "Modify a workspace item — update its name, content, folder_id, or "
        "status. For targeted edits, provide only what changed. To replace "
        "the entire content, provide a full content envelope matching the "
        "item's file_type. Use get_workspace_item first to see current state.\n\n"
        "For json_view: send content as a direct object {documentMetadata, views[]} "
        "(NOT array-wrapped). For all other types: content is array-wrapped.",
        {
            "type": "object",
            "properties": {
                "item_id": {
                    "type": "integer",
                    "description": "Workspace item ID to update.",
                },
                "name": {
                    "type": "string",
                    "description": "New display name (optional).",
                },
                "content": {
                    "type": "object",
                    "description": "Full replacement content envelope (optional). "
                    "For json_view: direct object. For other types: array-wrapped.",
                },
                "folder_id": {
                    "type": "integer",
                    "description": "Move to a different folder ID (optional). Null means root.",
                },
                "status": {
                    "type": "string",
                    "enum": ["draft", "review", "final"],
                    "description": "New status (optional).",
                },
            },
            "required": ["item_id"],
        },
    )
    async def update_workspace_item(args: dict[str, Any]) -> dict[str, Any]:
        item_id = args["item_id"]
        try:
            # Verify case scope
            conn = _conn()
            try:
                from core.db import get_draft as _get_draft
                item = _get_draft(conn, item_id)
            finally:
                conn.close()
            if not item:
                return _error(f"Workspace item {item_id} not found.")
            if item["case_id"] != case_id:
                return _error(f"Workspace item {item_id} not in case {case_id}.")

            kwargs = {}
            if "name" in args:
                kwargs["name"] = args["name"]
            if "status" in args:
                kwargs["status"] = args["status"]
            if "folder_id" in args:
                kwargs["folder_id"] = args["folder_id"]
            if "content" in args:
                content_raw = args["content"]
                item_file_type = item.get("file_type", "")
                if item_file_type == "json_view" and isinstance(content_raw, dict):
                    kwargs["content"] = content_raw  # direct object — never wrap
                elif isinstance(content_raw, list):
                    kwargs["content"] = content_raw
                elif isinstance(content_raw, dict):
                    kwargs["content"] = [content_raw]
                else:
                    kwargs["content"] = [{"raw": str(content_raw)}]

            if not kwargs:
                return _error("No fields to update.")

            conn = _conn()
            try:
                from core.db import update_draft as _update_draft
                updated = _update_draft(conn, item_id, **kwargs)
            finally:
                conn.close()

            return _result({
                "item_id": item_id,
                "name": updated["name"],
                "file_type": updated.get("file_type", ""),
                "folder_id": updated.get("folder_id"),
                "block_count": len(updated.get("content", [])),
                "updated_at": str(updated.get("updated_at", "")),
            })
        except Exception as exc:
            return _error(f"update_workspace_item failed: {exc}")

    # -- Layer 7: Tasks ------------------------------------------------------

    @tool(
        "list_tasks",
        "List all tasks for the current case, ordered by urgency. "
        "Overdue tasks first, then nearest deadline, then priority. "
        "Use this to check what needs to be done, what's pending, "
        "or what follow-ups were created from previous analysis.",
        {
            "type": "object",
            "properties": {
                "status": {
                    "type": "string",
                    "enum": ["open", "in_progress", "blocked", "complete"],
                    "description": "Filter by status (optional).",
                },
                "assignee_id": {
                    "type": "string",
                    "description": "Filter by assignee user ID (optional).",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_tasks(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_tasks as _list_tasks
                rows = _list_tasks(
                    conn, case_id,
                    status=args.get("status"),
                    assignee_id=args.get("assignee_id"),
                )
            finally:
                conn.close()
            return _result({"count": len(rows), "tasks": rows})
        except Exception as exc:
            return _error(f"list_tasks failed: {exc}")

    @tool(
        "create_task",
        "Create a new task in the case tracker. Use this to create "
        "follow-up items after analysis, reminders for missing evidence, "
        "or action items the user needs to handle.",
        {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Task title. Be specific about what needs to happen.",
                },
                "notes": {
                    "type": "string",
                    "description": "Detailed notes or context (optional).",
                },
                "deadline": {
                    "type": "string",
                    "description": "Deadline date in YYYY-MM-DD format (optional).",
                },
                "priority": {
                    "type": "string",
                    "enum": ["low", "medium", "high", "urgent"],
                    "description": "Priority level. Default medium.",
                },
                "document_ids": {
                    "type": "array",
                    "items": {"type": "integer"},
                    "description": "Document IDs to attach (optional). "
                    "Use document IDs from get_case or list_documents.",
                },
            },
            "required": ["title"],
        },
    )
    async def create_task(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import insert_task as _insert_task
                task_id = _insert_task(
                    conn,
                    case_id=case_id,
                    title=args["title"],
                    notes=args.get("notes"),
                    deadline=args.get("deadline"),
                    priority=args.get("priority", "medium"),
                    created_by="agent",
                )
                doc_ids = args.get("document_ids", [])
                if doc_ids:
                    from core.db import attach_task_documents as _attach
                    _attach(conn, task_id, doc_ids)
            finally:
                conn.close()
            return _result({
                "task_id": task_id,
                "title": args["title"],
                "priority": args.get("priority", "medium"),
                "deadline": args.get("deadline"),
                "documents_attached": len(doc_ids) if doc_ids else 0,
            })
        except Exception as exc:
            return _error(f"create_task failed: {exc}")

    @tool(
        "update_task",
        "Update a task's status, notes, or assignment. Use this to "
        "mark tasks complete when follow-up is done, or update notes "
        "with findings.",
        {
            "type": "object",
            "properties": {
                "task_id": {
                    "type": "integer",
                    "description": "Task ID from list_tasks.",
                },
                "status": {
                    "type": "string",
                    "enum": ["open", "in_progress", "blocked", "complete"],
                    "description": "New status (optional).",
                },
                "notes": {
                    "type": "string",
                    "description": "Updated notes (optional).",
                },
                "priority": {
                    "type": "string",
                    "enum": ["low", "medium", "high", "urgent"],
                    "description": "Updated priority (optional).",
                },
                "deadline": {
                    "type": "string",
                    "description": "Updated deadline YYYY-MM-DD (optional).",
                },
            },
            "required": ["task_id"],
        },
    )
    async def update_task(args: dict[str, Any]) -> dict[str, Any]:
        task_id = args["task_id"]
        try:
            kwargs = {}
            for f in ("status", "notes", "priority", "deadline"):
                if f in args:
                    kwargs[f] = args[f]
            if not kwargs:
                return _error("No fields to update.")
            conn = _conn()
            try:
                from core.db import update_task as _update_task
                updated = _update_task(conn, task_id, **kwargs)
            finally:
                conn.close()
            if not updated:
                return _error(f"Task {task_id} not found.")
            return _result({
                "task_id": task_id,
                "status": updated.get("status"),
                "updated_at": str(updated.get("updated_at", "")),
            })
        except Exception as exc:
            return _error(f"update_task failed: {exc}")

    @tool(
        "delete_task",
        "Delete a task. Use when a task is no longer relevant or was created "
        "in error.",
        {
            "type": "object",
            "properties": {
                "task_id": {
                    "type": "integer",
                    "description": "Task ID from list_tasks.",
                },
            },
            "required": ["task_id"],
        },
    )
    async def delete_task(args: dict[str, Any]) -> dict[str, Any]:
        task_id = args["task_id"]
        try:
            conn = _conn()
            try:
                from core.db import delete_task as _delete_task
                ok = _delete_task(conn, task_id)
            finally:
                conn.close()
            if not ok:
                return _error(f"Task {task_id} not found.")
            return _result({"deleted": True, "task_id": task_id})
        except Exception as exc:
            return _error(f"delete_task failed: {exc}")

    # -- Layer 8.5: Calendar Events & Reminders ------------------------------

    @tool(
        "create_calendar_event",
        "Create a calendar event in the case. Use this to schedule hearings, "
        "depositions, deadlines, meetings, or other events. The event appears "
        "on the case calendar and can have reminders attached.\n\n"
        "IMPORTANT: Always use the user's local timezone (typically Eastern: "
        "-04:00 in summer, -05:00 in winter). Include the timezone offset in "
        "start_time and end_time. For all-day events like filing deadlines, "
        "set all_day=true and use midnight for start_time.",
        {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Event title. Be specific (e.g., "
                    "'Deposition — Jane Smith' not just 'Meeting').",
                },
                "start_time": {
                    "type": "string",
                    "description": "Event start time in ISO 8601 format with "
                    "timezone offset (e.g., '2026-07-15T09:00:00-04:00'). "
                    "Required.",
                },
                "end_time": {
                    "type": "string",
                    "description": "Event end time (optional). ISO 8601 with "
                    "timezone. Omit for point-in-time events.",
                },
                "all_day": {
                    "type": "boolean",
                    "description": "Set true for all-day events like filing "
                    "deadlines. When true, start_time should be midnight of "
                    "the event date. Default false.",
                },
                "category": {
                    "type": "string",
                    "enum": ["hearing", "deposition", "deadline", "meeting", "other"],
                    "description": "Event category. Choose based on the type "
                    "of event. Default 'other'.",
                },
                "description": {
                    "type": "string",
                    "description": "Detailed notes — purpose, preparation "
                    "needed, attendees, Zoom links, etc. (optional).",
                },
                "location": {
                    "type": "string",
                    "description": "Physical address, courtroom number, or "
                    "virtual meeting link (optional).",
                },
            },
            "required": ["title", "start_time"],
        },
    )
    async def create_calendar_event(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import insert_calendar_event as _insert
                event_id = _insert(
                    conn,
                    case_id=case_id,
                    title=args["title"],
                    start_time=args["start_time"],
                    end_time=args.get("end_time"),
                    all_day=args.get("all_day", False),
                    category=args.get("category", "other"),
                    description=args.get("description"),
                    location=args.get("location"),
                    created_by="agent",
                )
            finally:
                conn.close()
            return _result({
                "event_id": event_id,
                "title": args["title"],
                "start_time": args["start_time"],
                "category": args.get("category", "other"),
                "all_day": args.get("all_day", False),
            })
        except Exception as exc:
            return _error(f"create_calendar_event failed: {exc}")

    @tool(
        "list_calendar_events",
        "List calendar events for the current case. Use to answer questions "
        "like 'what's on the calendar this week?' or 'when is the next hearing?' "
        "Supports filtering by date range and category.",
        {
            "type": "object",
            "properties": {
                "start_date": {
                    "type": "string",
                    "description": "Filter: events from this date (YYYY-MM-DD, "
                    "inclusive). Use to find upcoming events.",
                },
                "end_date": {
                    "type": "string",
                    "description": "Filter: events until this date (YYYY-MM-DD, "
                    "inclusive).",
                },
                "category": {
                    "type": "string",
                    "enum": ["hearing", "deposition", "deadline", "meeting", "other"],
                    "description": "Filter by event category (optional).",
                },
            },
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_calendar_events(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_calendar_events as _list
                rows = _list(
                    conn, case_id,
                    start_date=args.get("start_date"),
                    end_date=args.get("end_date"),
                    category=args.get("category"),
                )
            finally:
                conn.close()
            return _result({"count": len(rows), "events": rows})
        except Exception as exc:
            return _error(f"list_calendar_events failed: {exc}")

    @tool(
        "get_calendar_event",
        "Get a single calendar event by ID with its attached reminders. "
        "Use to see full event details and any reminders linked to it.",
        {
            "type": "object",
            "properties": {
                "event_id": {
                    "type": "integer",
                    "description": "Calendar event ID from list_calendar_events.",
                },
            },
            "required": ["event_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_calendar_event(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import get_calendar_event as _get
                event = _get(conn, args["event_id"])
            finally:
                conn.close()
            if not event:
                return _error(f"Calendar event {args['event_id']} not found.")
            return _result({"event": event})
        except Exception as exc:
            return _error(f"get_calendar_event failed: {exc}")

    @tool(
        "create_reminder",
        "Create a reminder for the case. Reminders can be standalone or "
        "attached to a calendar event. Use this when the user asks to be "
        "reminded about something.\n\n"
        "IMPORTANT: Compute the absolute remind_at time yourself. If the user "
        "says 'remind me 48 hours before the hearing', look up the hearing "
        "event, subtract 48 hours from its start_time, and use that as "
        "remind_at. The schema only stores absolute times — never intervals.\n\n"
        "Use the same timezone as the user (typically Eastern).",
        {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Reminder title. Be specific about what "
                    "the user needs to do.",
                },
                "remind_at": {
                    "type": "string",
                    "description": "Absolute time to fire the reminder in ISO "
                    "8601 format with timezone offset (e.g., "
                    "'2026-07-13T09:00:00-04:00'). Compute this from the "
                    "event's start_time if the user gives an interval like "
                    "'48 hours before'.",
                },
                "event_id": {
                    "type": "integer",
                    "description": "Calendar event ID to attach this reminder "
                    "to (optional). Omit for standalone reminders.",
                },
                "category": {
                    "type": "string",
                    "enum": ["hearing", "deposition", "deadline", "meeting", "other"],
                    "description": "Same categories as events. Default 'other'.",
                },
                "description": {
                    "type": "string",
                    "description": "Additional context about what the user "
                    "needs to do (optional).",
                },
            },
            "required": ["title", "remind_at"],
        },
    )
    async def create_reminder(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import insert_reminder as _insert
                reminder_id = _insert(
                    conn,
                    case_id=case_id,
                    title=args["title"],
                    remind_at=args["remind_at"],
                    event_id=args.get("event_id"),
                    category=args.get("category", "other"),
                    description=args.get("description"),
                    created_by="agent",
                )
            finally:
                conn.close()
            return _result({
                "reminder_id": reminder_id,
                "title": args["title"],
                "remind_at": args["remind_at"],
                "event_id": args.get("event_id"),
            })
        except Exception as exc:
            return _error(f"create_reminder failed: {exc}")

    @tool(
        "list_reminders",
        "List reminders for the current case. Use to see what reminders "
        "are pending, which have fired, or which were dismissed. Supports "
        "filtering by status and category.",
        {
            "type": "object",
            "properties": {
                "status": {
                    "type": "string",
                    "enum": ["pending", "fired", "dismissed"],
                    "description": "Filter by status (optional). Default: all.",
                },
                "category": {
                    "type": "string",
                    "enum": ["hearing", "deposition", "deadline", "meeting", "other"],
                    "description": "Filter by category (optional).",
                },
                "event_id": {
                    "type": "integer",
                    "description": "Filter: reminders attached to a specific "
                    "calendar event (optional).",
                },
            },
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_reminders(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_reminders as _list
                rows = _list(
                    conn, case_id,
                    status=args.get("status"),
                    category=args.get("category"),
                    event_id=args.get("event_id"),
                )
            finally:
                conn.close()
            return _result({"count": len(rows), "reminders": rows})
        except Exception as exc:
            return _error(f"list_reminders failed: {exc}")

    @tool(
        "get_reminder",
        "Get a single reminder by ID. Use to see full reminder details.",
        {
            "type": "object",
            "properties": {
                "reminder_id": {
                    "type": "integer",
                    "description": "Reminder ID from list_reminders.",
                },
            },
            "required": ["reminder_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_reminder(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import get_reminder as _get
                reminder = _get(conn, args["reminder_id"])
            finally:
                conn.close()
            if not reminder:
                return _error(f"Reminder {args['reminder_id']} not found.")
            return _result({"reminder": reminder})
        except Exception as exc:
            return _error(f"get_reminder failed: {exc}")

    # -- Layer 8: Correspondence ---------------------------------------------

    @tool(
        "list_correspondence_threads",
        "List all correspondence threads for the current case. Returns id, "
        "title, status, item count, and last activity. Use to see existing "
        "threads before adding items.",
        {
            "type": "object",
            "properties": {
                "status": {
                    "type": "string",
                    "enum": ["active", "archived"],
                    "description": "Filter by status (optional). Defaults to all.",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_correspondence_threads(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                clauses = ["ct.case_id = %s"]
                params: list[Any] = [case_id]
                if args.get("status"):
                    clauses.append("ct.status = %s")
                    params.append(args["status"])
                where = " AND ".join(clauses)
                sql = f"""SELECT ct.*,
                                 (SELECT count(*) FROM correspondence_items
                                  WHERE thread_id = ct.id) AS item_count,
                                 (SELECT max(updated_at) FROM correspondence_items
                                  WHERE thread_id = ct.id) AS last_activity
                          FROM correspondence_threads ct
                          WHERE {where}
                          ORDER BY ct.updated_at DESC"""
                import psycopg2.extras
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    cur.execute(sql, tuple(params))
                    rows = [dict(r) for r in cur.fetchall()]
            finally:
                conn.close()
            return _result({"count": len(rows), "threads": rows})
        except Exception as exc:
            return _error(f"list_correspondence_threads failed: {exc}")

    @tool(
        "create_correspondence_thread",
        "Create a new correspondence thread. A thread groups related "
        "correspondence items (e.g., 'Discovery letters to opposing counsel'). "
        "Use this before logging individual items.",
        {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Thread title. Be descriptive.",
                },
            },
            "required": ["title"],
        },
    )
    async def create_correspondence_thread(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                import psycopg2.extras
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    cur.execute(
                        """INSERT INTO correspondence_threads (case_id, title)
                           VALUES (%s, %s) RETURNING *""",
                        (case_id, args["title"]),
                    )
                    thread = dict(cur.fetchone())
                conn.commit()
            finally:
                conn.close()
            return _result({"thread": thread})
        except Exception as exc:
            return _error(f"create_correspondence_thread failed: {exc}")

    @tool(
        "update_correspondence_thread",
        "Update a correspondence thread's title or status. Archive threads "
        "when they're no longer active.",
        {
            "type": "object",
            "properties": {
                "thread_id": {
                    "type": "integer",
                    "description": "Thread ID from list_correspondence_threads.",
                },
                "title": {
                    "type": "string",
                    "description": "New title (optional).",
                },
                "status": {
                    "type": "string",
                    "enum": ["active", "archived"],
                    "description": "New status (optional).",
                },
            },
            "required": ["thread_id"],
        },
    )
    async def update_correspondence_thread(args: dict[str, Any]) -> dict[str, Any]:
        try:
            sets = []
            params: list[Any] = []
            if "title" in args:
                sets.append("title = %s"); params.append(args["title"])
            if "status" in args:
                sets.append("status = %s"); params.append(args["status"])
            if not sets:
                return _error("No fields to update.")
            sets.append("updated_at = now()")
            params.append(args["thread_id"])
            conn = _conn()
            try:
                import psycopg2.extras
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    cur.execute(
                        f"UPDATE correspondence_threads SET {', '.join(sets)} "
                        f"WHERE id = %s AND case_id = %s RETURNING *",
                        tuple(params + [case_id]),
                    )
                    row = cur.fetchone()
                conn.commit()
            finally:
                conn.close()
            if not row:
                return _error("Thread not found in this case.")
            return _result({"thread": dict(row)})
        except Exception as exc:
            return _error(f"update_correspondence_thread failed: {exc}")

    @tool(
        "list_correspondence_items",
        "List all items in a correspondence thread. Returns sender, receiver, "
        "direction, notes, dates, and attached documents.",
        {
            "type": "object",
            "properties": {
                "thread_id": {
                    "type": "integer",
                    "description": "Thread ID from list_correspondence_threads.",
                },
            },
            "required": ["thread_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_correspondence_items(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                import psycopg2.extras
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    # Verify thread belongs to case
                    cur.execute(
                        "SELECT id FROM correspondence_threads WHERE id = %s AND case_id = %s",
                        (args["thread_id"], case_id),
                    )
                    if not cur.fetchone():
                        return _error("Thread not found in this case.")
                    cur.execute(
                        """SELECT ci.*,
                                  sp.name AS sender_name,
                                  rp.name AS receiver_name,
                                  (SELECT jsonb_agg(
                                      jsonb_build_object(
                                          'id', ca.id,
                                          'document_id', ca.document_id,
                                          'document_name', d.name
                                      )
                                   )
                                   FROM correspondence_attachments ca
                                   JOIN documents d ON ca.document_id = d.id
                                   WHERE ca.item_id = ci.id
                                  ) AS attachments
                           FROM correspondence_items ci
                           LEFT JOIN parties sp ON ci.sender_party_id = sp.id
                           LEFT JOIN parties rp ON ci.receiver_party_id = rp.id
                           WHERE ci.thread_id = %s
                           ORDER BY ci.date_sent DESC, ci.date_received DESC,
                                    ci.created_at DESC""",
                        (args["thread_id"],),
                    )
                    rows = [dict(r) for r in cur.fetchall()]
            finally:
                conn.close()
            return _result({"count": len(rows), "items": rows})
        except Exception as exc:
            return _error(f"list_correspondence_items failed: {exc}")

    @tool(
        "create_correspondence_item",
        "Log a new correspondence item in a thread. Records who sent/received "
        "it, the direction, dates, notes, and optionally attaches documents. "
        "Use party IDs from get_case (parties array).",
        {
            "type": "object",
            "properties": {
                "thread_id": {
                    "type": "integer",
                    "description": "Thread ID from list_correspondence_threads.",
                },
                "direction": {
                    "type": "string",
                    "enum": ["sent", "received"],
                    "description": "Whether this was sent by us or received.",
                },
                "sender_party_id": {
                    "type": "integer",
                    "description": "Party ID of the sender (optional).",
                },
                "receiver_party_id": {
                    "type": "integer",
                    "description": "Party ID of the receiver (optional).",
                },
                "notes": {
                    "type": "string",
                    "description": "Notes or summary of the correspondence (optional).",
                },
                "date_sent": {
                    "type": "string",
                    "description": "Date sent in YYYY-MM-DD format (optional).",
                },
                "date_received": {
                    "type": "string",
                    "description": "Date received in YYYY-MM-DD format (optional).",
                },
                "document_ids": {
                    "type": "array",
                    "items": {"type": "integer"},
                    "description": "Document IDs to attach (optional).",
                },
            },
            "required": ["thread_id", "direction"],
        },
    )
    async def create_correspondence_item(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                import psycopg2.extras
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    # Verify thread belongs to case
                    cur.execute(
                        "SELECT id FROM correspondence_threads WHERE id = %s AND case_id = %s",
                        (args["thread_id"], case_id),
                    )
                    if not cur.fetchone():
                        return _error("Thread not found in this case.")
                    cur.execute(
                        """INSERT INTO correspondence_items
                           (thread_id, sender_party_id, receiver_party_id,
                            direction, notes, date_sent, date_received)
                           VALUES (%s, %s, %s, %s, %s, %s, %s)
                           RETURNING *""",
                        (args["thread_id"], args.get("sender_party_id"),
                         args.get("receiver_party_id"), args["direction"],
                         args.get("notes"), args.get("date_sent"),
                         args.get("date_received")),
                    )
                    item = dict(cur.fetchone())
                    doc_ids = args.get("document_ids", [])
                    for did in doc_ids:
                        cur.execute(
                            """INSERT INTO correspondence_attachments
                               (item_id, document_id) VALUES (%s, %s)
                               ON CONFLICT DO NOTHING""",
                            (item["id"], did),
                        )
                conn.commit()
            finally:
                conn.close()
            return _result({
                "item": item,
                "documents_attached": len(doc_ids) if doc_ids else 0,
            })
        except Exception as exc:
            return _error(f"create_correspondence_item failed: {exc}")

    @tool(
        "update_correspondence_item",
        "Update a correspondence item's fields. Use to correct dates, add "
        "notes, or change the direction.",
        {
            "type": "object",
            "properties": {
                "item_id": {
                    "type": "integer",
                    "description": "Item ID from list_correspondence_items.",
                },
                "notes": {"type": "string", "description": "Updated notes."},
                "direction": {
                    "type": "string",
                    "enum": ["sent", "received"],
                    "description": "Updated direction.",
                },
                "date_sent": {
                    "type": "string",
                    "description": "Updated date sent YYYY-MM-DD.",
                },
                "date_received": {
                    "type": "string",
                    "description": "Updated date received YYYY-MM-DD.",
                },
                "sender_party_id": {
                    "type": "integer",
                    "description": "Updated sender party ID.",
                },
                "receiver_party_id": {
                    "type": "integer",
                    "description": "Updated receiver party ID.",
                },
            },
            "required": ["item_id"],
        },
    )
    async def update_correspondence_item(args: dict[str, Any]) -> dict[str, Any]:
        try:
            sets = []
            params: list[Any] = []
            for f in ("notes", "direction", "date_sent", "date_received",
                       "sender_party_id", "receiver_party_id"):
                if f in args and args[f] is not None:
                    sets.append(f"{f} = %s"); params.append(args[f])
            if not sets:
                return _error("No fields to update.")
            sets.append("updated_at = now()")
            params.append(args["item_id"])
            params.append(case_id)
            conn = _conn()
            try:
                import psycopg2.extras
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    cur.execute(
                        f"""UPDATE correspondence_items ci SET {', '.join(sets)}
                            FROM correspondence_threads ct
                            WHERE ci.thread_id = ct.id
                              AND ci.id = %s AND ct.case_id = %s
                            RETURNING ci.*""",
                        tuple(params),
                    )
                    row = cur.fetchone()
                conn.commit()
            finally:
                conn.close()
            if not row:
                return _error("Item not found in this case.")
            return _result({"item": dict(row)})
        except Exception as exc:
            return _error(f"update_correspondence_item failed: {exc}")

    @tool(
        "delete_correspondence_item",
        "Delete a correspondence item. Use if an entry was logged in error.",
        {
            "type": "object",
            "properties": {
                "item_id": {
                    "type": "integer",
                    "description": "Item ID from list_correspondence_items.",
                },
            },
            "required": ["item_id"],
        },
    )
    async def delete_correspondence_item(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                import psycopg2.extras
                with conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor) as cur:
                    cur.execute(
                        """DELETE FROM correspondence_items ci
                           USING correspondence_threads ct
                           WHERE ci.thread_id = ct.id
                             AND ci.id = %s AND ct.case_id = %s
                           RETURNING ci.id""",
                        (args["item_id"], case_id),
                    )
                    row = cur.fetchone()
                conn.commit()
            finally:
                conn.close()
            if not row:
                return _error("Item not found in this case.")
            return _result({"deleted": True, "item_id": args["item_id"]})
        except Exception as exc:
            return _error(f"delete_correspondence_item failed: {exc}")

    # -- Layer 9: Company Profile --------------------------------------------

    @tool(
        "list_company_profiles",
        "List all company profiles. Returns id, name, description, status, "
        "and timestamps for each profile. Does not return full content — use "
        "get_company_profile for the detailed profile data including CAGE/UEI, "
        "NAICS codes, certifications, past performance, and key personnel.",
        {
            "type": "object",
            "properties": {},
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_company_profiles(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_company_profiles as _list_cp
                profiles = _list_cp(conn)
            finally:
                conn.close()
            # Strip full content — return metadata only
            summaries = [
                {k: v for k, v in p.items() if k != "content"}
                for p in profiles
            ]
            return _result({"count": len(summaries), "profiles": summaries})
        except Exception as exc:
            return _error(f"list_company_profiles failed: {exc}")

    @tool(
        "get_company_profile",
        "Get a company profile's full data including the structured content "
        "JSONB. The content field contains: company_name, legal_name, dba, "
        "tax_id, cage_code, uei, psc_codes, naics_codes (array), "
        "certifications (array with expirations), past_performance (array of "
        "{client, contract_value, description, period_of_performance}), "
        "key_personnel (array of {name, title, years_experience, clearance}), "
        "contact ({address_line1, city, state, zip, phone, email}), and "
        "field_status indicating confidence for each filled field. "
        "Use this to get company details for capability statements, past "
        "performance references, and proposal responses.",
        {
            "type": "object",
            "properties": {
                "profile_id": {
                    "type": "integer",
                    "description": "Company profile ID from list_company_profiles.",
                },
            },
            "required": ["profile_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_company_profile(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import get_company_profile as _get_cp
                profile = _get_cp(conn, args["profile_id"])
            finally:
                conn.close()
            if not profile:
                return _error(f"Company profile {args['profile_id']} not found.")
            return _result({"profile": profile})
        except Exception as exc:
            return _error(f"get_company_profile failed: {exc}")

    @tool(
        "get_case_profile",
        "Get the company profile attached to the current case. Cases have a "
        "profile_id column that links to a company_profiles row. This is the "
        "profile to use for any solicitation response, capability statement, "
        "or proposal. Returns the full profile content including CAGE/UEI, "
        "NAICS codes, certifications, past performance, and key personnel. "
        "Use this FIRST before drafting any response — never invent company "
        "information.",
        {
            "type": "object",
            "properties": {},
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_case_profile(args: dict[str, Any]) -> dict[str, Any]:
        try:
            case_row = _query_one(
                "SELECT profile_id FROM cases WHERE id = %s", (case_id,)
            )
            if not case_row or not case_row.get("profile_id"):
                return _error(
                    "No company profile attached to this case. "
                    "Create a company profile first at /api/profiles and "
                    "attach it to the case."
                )
            conn = _conn()
            try:
                from core.db import get_company_profile as _get_cp
                profile = _get_cp(conn, case_row["profile_id"])
            finally:
                conn.close()
            if not profile:
                return _error(
                    f"Attached profile {case_row['profile_id']} not found — "
                    "it may have been deleted."
                )
            return _result({"profile": profile})
        except Exception as exc:
            return _error(f"get_case_profile failed: {exc}")

    # -- Knowledge Base helper (capture case_id for cross-case knowledge) -----

    def _ensure_kb_case() -> int:
        """Get or create the cross-case Knowledge Base. Returns case_id."""
        row = _query_one(
            "SELECT id FROM cases WHERE name = %s AND case_type = 'other'",
            (KB_CASE_NAME,),
        )
        if row:
            return row["id"]
        from core.db import insert_case as _insert_case
        conn = _conn()
        try:
            kb_id = _insert_case(conn, name=KB_CASE_NAME, case_type="other")
        finally:
            conn.close()
        return kb_id

    # -- Layer 9.5: Knowledge Base -------------------------------------------

    @tool(
        "create_knowledge_entry",
        "Persist a piece of knowledge as a searchable, tagged document in "
        "the cross-case Knowledge Base. Use this when you learn something "
        "worth remembering: an industry strategy, a competitor insight, a "
        "procurement lesson, an agency contact pattern, a pricing approach, "
        "or any reusable GovCon intelligence. Each entry gets tags for "
        "filtered retrieval later. The content is markdown — structure it "
        "with headers, bullet lists, and links as appropriate.",
        {
            "type": "object",
            "properties": {
                "title": {
                    "type": "string",
                    "description": "Short descriptive title. Think of it as a filename.",
                },
                "content": {
                    "type": "string",
                    "description": (
                        "Full markdown content. Use headers, lists, code blocks, "
                        "and links. This is the body of the knowledge entry."
                    ),
                },
                "tags": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": (
                        "Tags for categorization and retrieval. Use lowercase "
                        "kebab-case or short labels. Examples: 'cybersecurity', "
                        "'pricing-strategy', 'agency-patterns', 'incumbent-analysis', "
                        "'set-aside-tactics', 'naics-541511'. At least one tag required."
                    ),
                },
                "source_url": {
                    "type": "string",
                    "description": "Optional URL where this knowledge was sourced from.",
                },
                "document_type": {
                    "type": "string",
                    "enum": ["strategy", "insight", "lesson_learned",
                             "reference", "template", "other"],
                    "description": "Classification. Defaults to 'insight'.",
                },
            },
            "required": ["title", "content", "tags"],
        },
    )
    async def create_knowledge_entry(args: dict[str, Any]) -> dict[str, Any]:
        try:
            kb_case_id = _ensure_kb_case()
            tags = args["tags"]
            knowledge_type = args.get("document_type", "insight")

            content_to_store = [{"markdown": args["content"]}]

            conn = _conn()
            try:
                from core.db import insert_draft as _insert_draft
                item_id = _insert_draft(
                    conn,
                    case_id=kb_case_id,
                    name=args["title"],
                    document_type="other",
                    content=content_to_store,
                    created_by="agent",
                    file_type="markdown",
                    folder="artifacts",
                    status="final",
                )
            finally:
                conn.close()

            # Write tags + knowledge type to metadata
            import json as _json
            conn = _conn()
            try:
                with conn.cursor() as cur:
                    cur.execute(
                        "UPDATE drafts SET metadata = %s::jsonb WHERE id = %s",
                        (_json.dumps({
                            "tags": tags,
                            "knowledge_type": knowledge_type,
                            "source_url": args.get("source_url", ""),
                            "created_by": "agent",
                        }), item_id),
                    )
            finally:
                conn.close()

            return _result({
                "entry_id": item_id,
                "title": args["title"],
                "tags": tags,
                "document_type": doc_type,
                "case_id": kb_case_id,
            })
        except Exception as exc:
            return _error(f"create_knowledge_entry failed: {exc}")

    @tool(
        "search_knowledge",
        "Search the Knowledge Base by tags, text, or both. Returns "
        "matching entries ranked by relevance. Use this before creating "
        "new entries to avoid duplicates, or when you need to recall "
        "previously stored strategies, insights, or reference material. "
        "Pass tags to filter to a specific topic; pass query for full-text "
        "search across all entry content; pass both for tagged search.",
        {
            "type": "object",
            "properties": {
                "query": {
                    "type": "string",
                    "description": "Full-text search query across title and content.",
                },
                "tags": {
                    "type": "array",
                    "items": {"type": "string"},
                    "description": "Filter to entries matching ANY of these tags.",
                },
                "document_type": {
                    "type": "string",
                    "enum": ["strategy", "insight", "lesson_learned",
                             "reference", "template", "other"],
                    "description": "Filter by document type.",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def search_knowledge(args: dict[str, Any]) -> dict[str, Any]:
        try:
            kb_case_id = _ensure_kb_case()
            query = args.get("query", "").strip()
            tags = args.get("tags") or []
            doc_type = args.get("document_type")

            # Build SQL
            conditions = ["d.case_id = %s"]
            params: list[Any] = [kb_case_id]
            order = "d.updated_at DESC"

            if query:
                conditions.append(
                    """(to_tsvector('english', d.name || ' ' ||
                       COALESCE(d.content::text, ''))
                       @@ plainto_tsquery('english', %s))"""
                )
                params.append(query)
                order = (
                    f"ts_rank(to_tsvector('english', d.name || ' ' || "
                    f"COALESCE(d.content::text, '')), "
                    f"plainto_tsquery('english', %s)) DESC"
                )

            if tags:
                tag_clauses = []
                for tag in tags:
                    tag_clauses.append("d.metadata->>'tags' ILIKE %s")
                    params.append(f"%{tag}%")
                conditions.append(f"({' OR '.join(tag_clauses)})")

            if doc_type:
                conditions.append("d.metadata->>'knowledge_type' = %s")
                params.append(doc_type)

            where = " AND ".join(conditions)
            sql = (
                f"SELECT d.id, d.name, d.document_type, d.metadata, "
                f"d.file_type, d.folder, d.status, d.created_at, d.updated_at, "
                f"COALESCE((d.content->0->>'markdown'), '') AS body "
                f"FROM drafts d WHERE {where} ORDER BY {order} LIMIT 20"
            )

            rows = _query(sql, tuple(params))

            results = []
            for r in rows:
                body = r.get("body") or ""
                meta = r.get("metadata") or {}
                if isinstance(meta, str):
                    import json as _json
                    meta = _json.loads(meta)
                results.append({
                    "entry_id": r["id"],
                    "title": r["name"],
                    "knowledge_type": meta.get("knowledge_type", "insight"),
                    "tags": meta.get("tags", []),
                    "source_url": meta.get("source_url", ""),
                    "excerpt": body[:300] + ("..." if len(body) > 300 else ""),
                    "created_at": str(r.get("created_at", "")),
                    "updated_at": str(r.get("updated_at", "")),
                })

            return _result({
                "query": query or None,
                "tags_filter": tags or None,
                "count": len(results),
                "results": results,
            })
        except Exception as exc:
            return _error(f"search_knowledge failed: {exc}")

    @tool(
        "list_knowledge_tags",
        "List all unique tags currently used across the Knowledge Base, "
        "with counts. Use this to understand what topics have been "
        "captured before searching or creating entries.",
        {
            "type": "object",
            "properties": {},
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_knowledge_tags(args: dict[str, Any]) -> dict[str, Any]:
        try:
            kb_case_id = _ensure_kb_case()
            rows = _query(
                """SELECT d.metadata->>'tags' AS tags_json
                   FROM drafts d WHERE d.case_id = %s
                     AND d.metadata->>'tags' IS NOT NULL""",
                (kb_case_id,),
            )
            import json as _json
            tag_counts: dict[str, int] = {}
            for r in rows:
                raw = r.get("tags_json") or "[]"
                try:
                    tag_list = _json.loads(raw) if isinstance(raw, str) else raw
                except _json.JSONDecodeError:
                    continue
                for tag in tag_list:
                    tag_counts[tag] = tag_counts.get(tag, 0) + 1

            sorted_tags = sorted(tag_counts.items(), key=lambda x: -x[1])
            return _result({
                "total_entries_with_tags": sum(
                    1 for r in rows if r.get("tags_json")
                ),
                "unique_tags": len(sorted_tags),
                "tags": [
                    {"tag": t, "count": c} for t, c in sorted_tags
                ],
            })
        except Exception as exc:
            return _error(f"list_knowledge_tags failed: {exc}")

    # -- Layer 10: FAR (Federal Acquisition Regulation) -----------------------

    @tool(
        "far_lookup",
        "Look up the authoritative text of a FAR citation. Pass a FAR "
        "reference like '15.101', '52.212-1', 'Subpart 15.1', or 'Part 15'. "
        "Returns the exact regulatory text with proper citation format. "
        "Use this when a solicitation cites a FAR clause and you need the "
        "verbatim text, or when you need to verify a FAR reference is real. "
        "NEVER paraphrase a FAR clause from memory — always look it up.",
        {
            "type": "object",
            "properties": {
                "citation": {
                    "type": "string",
                    "description": (
                        "FAR citation to look up. Examples: '15.101', "
                        "'52.212-1', 'Subpart 15.1', 'Part 15', "
                        "'52.212-1(b)', '1.106'. Supports Part numbers, "
                        "section numbers, subpart references, and clause "
                        "numbers with dashes."
                    ),
                },
            },
            "required": ["citation"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def far_lookup(args: dict[str, Any]) -> dict[str, Any]:
        citation = args["citation"].strip()
        # Normalize: strip "FAR " prefix, "(a)" suffixes
        citation = re.sub(r"^FAR\s+", "", citation, flags=re.IGNORECASE)
        parenthetical = re.search(r"\(([a-z]+)\)$", citation)
        citation_clean = citation.split("(")[0].strip() if parenthetical else citation

        try:
            # Strategy 1: exact match on sections.metadata->>'far_number'
            rows = _query(
                """SELECT s.id AS section_id, s.title, s.search_text,
                          s.metadata->>'far_number' AS far_number,
                          s.heading_level, s.document_id,
                          d.name AS document_name,
                          d.metadata->>'part' AS part_num
                   FROM sections s
                   JOIN documents d ON d.id = s.document_id
                   JOIN cases c ON c.id = d.case_id
                   WHERE c.name = %s
                     AND s.metadata->>'far_number' = %s
                   LIMIT 5""",
                (FAR_CASE_NAME, citation_clean),
            )

            # Quality filter: exclude Part 52 parse artifacts where
            # far_number is a parenthetical like "(c)" or "(d)"
            QUALITY_FILTER = (
                "AND (s.metadata->>'far_number' ~ '^[0-9]' "
                "     OR s.metadata->>'far_number' IS NULL)"
            )

            if not rows:
                # Strategy 2: partial match — try without subsection
                base = citation_clean.rsplit("-", 1)[0] if "-" in citation_clean else citation_clean
                rows = _query(
                    f"""SELECT s.id AS section_id, s.title, s.search_text,
                              s.metadata->>'far_number' AS far_number,
                              s.heading_level, s.document_id,
                              d.name AS document_name,
                              d.metadata->>'part' AS part_num
                       FROM sections s
                       JOIN documents d ON d.id = s.document_id
                       JOIN cases c ON c.id = d.case_id
                       WHERE c.name = %s
                         AND s.metadata->>'far_number' LIKE %s
                         {QUALITY_FILTER}
                       ORDER BY s.heading_level
                       LIMIT 10""",
                    (FAR_CASE_NAME, base + "%"),
                )

            if not rows:
                # Strategy 3: title search (FAR numbers embedded in h1 text)
                rows = _query(
                    f"""SELECT s.id AS section_id, s.title, s.search_text,
                              s.metadata->>'far_number' AS far_number,
                              s.heading_level, s.document_id,
                              d.name AS document_name,
                              d.metadata->>'part' AS part_num
                       FROM sections s
                       JOIN documents d ON d.id = s.document_id
                       JOIN cases c ON c.id = d.case_id
                       WHERE c.name = %s
                         AND s.title LIKE %s
                         {QUALITY_FILTER}
                       ORDER BY s.heading_level
                       LIMIT 10""",
                    (FAR_CASE_NAME, "%" + citation_clean + "%"),
                )

            if not rows:
                # Strategy 4: full-text search fallback (last resort)
                rows = _query(
                    f"""SELECT s.id AS section_id, s.title, s.search_text,
                              s.metadata->>'far_number' AS far_number,
                              s.heading_level, s.document_id,
                              d.name AS document_name,
                              d.metadata->>'part' AS part_num,
                              ts_rank(
                                  to_tsvector('english', s.search_text),
                                  plainto_tsquery('english', %s)
                              ) AS rank
                       FROM sections s
                       JOIN documents d ON d.id = s.document_id
                       JOIN cases c ON c.id = d.case_id
                       WHERE c.name = %s
                         AND to_tsvector('english', s.search_text)
                             @@ plainto_tsquery('english', %s)
                         {QUALITY_FILTER}
                       ORDER BY rank DESC
                       LIMIT 10""",
                    (citation, FAR_CASE_NAME, citation),
                )

            if not rows:
                return _error(
                    f"Citation '{citation}' not found in the FAR corpus. "
                    "Check the citation format. If the FAR hasn't been "
                    "ingested yet, run: python -m scripts.far_ingest"
                )

            # Build the response
            results = []
            for r in rows:
                far_num = r.get("far_number") or ""
                title = r.get("title") or ""
                search_text = r.get("search_text") or ""

                # Get blocks for full text
                blocks = _query(
                    """SELECT text_content FROM blocks
                       WHERE section_id = %s
                       ORDER BY id""",
                    (r["section_id"],),
                )
                block_texts = [b["text_content"] for b in blocks if b.get("text_content")]

                results.append({
                    "citation": f"FAR {far_num}" if far_num else title,
                    "title": title,
                    "part": r.get("part_num"),
                    "document": r.get("document_name"),
                    "heading_level": r.get("heading_level"),
                    "full_text": (
                        title + "\n\n" +
                        "\n".join(block_texts)
                    ) if block_texts else search_text,
                    "block_count": len(block_texts),
                })

            return _result({
                "query": citation,
                "match_type": "exact" if rows and rows[0].get("far_number") == citation_clean else "broad",
                "count": len(results),
                "results": results,
            })

        except Exception as exc:
            return _error(f"far_lookup failed: {exc}")

    @tool(
        "far_status",
        "Check whether the FAR corpus has been ingested into the database. "
        "Returns the FAR case ID, part count, section count, block count, "
        "and last update time. If the FAR hasn't been ingested, returns "
        "instructions for running the ingest script. Use this before any "
        "FAR-dependent work to verify the corpus is available.",
        {
            "type": "object",
            "properties": {},
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def far_status(args: dict[str, Any]) -> dict[str, Any]:
        try:
            case_row = _query_one(
                "SELECT id, created_at FROM cases WHERE name = %s AND case_type = 'other'",
                (FAR_CASE_NAME,),
            )

            if not case_row:
                return _result({
                    "ingested": False,
                    "message": (
                        "The FAR corpus has not been ingested. "
                        "To ingest it, run from the backend directory:\n\n"
                        "    python -m scripts.far_ingest\n\n"
                        "This will download the FAR HTML ZIP from "
                        "acquisition.gov, parse all 53 Parts into "
                        "sections and blocks, and generate Mistral "
                        "embeddings for semantic search. "
                        "Estimated time: ~2 min ingest + ~7 min embed."
                    ),
                })

            case_id = case_row["id"]

            # Count documents, sections, blocks
            stats = _query_one(
                """SELECT
                       (SELECT count(*) FROM documents WHERE case_id = %s) AS doc_count,
                       (SELECT count(*) FROM sections s
                        JOIN documents d ON d.id = s.document_id
                        WHERE d.case_id = %s) AS section_count,
                       (SELECT count(*) FROM blocks b
                        JOIN documents d ON d.id = b.document_id
                        WHERE d.case_id = %s) AS block_count,
                       (SELECT count(*) FROM sections s
                        JOIN documents d ON d.id = s.document_id
                        WHERE d.case_id = %s AND s.embedding IS NOT NULL) AS embedded_count
                """,
                (case_id, case_id, case_id, case_id),
            )

            return _result({
                "ingested": True,
                "case_id": case_id,
                "case_name": FAR_CASE_NAME,
                "created_at": str(case_row.get("created_at", "")),
                "documents": stats["doc_count"] if stats else 0,
                "sections": stats["section_count"] if stats else 0,
                "blocks": stats["block_count"] if stats else 0,
                "embedded": stats["embedded_count"] if stats else 0,
                "current_fac": "FAC 2026-01 (2026-03-13)",
                "source_url": FAR_ZIP_URL,
                "re_ingest_command": "python -m scripts.far_ingest",
            })

        except Exception as exc:
            return _error(f"far_status failed: {exc}")

    # -- Layer 10.5: Federal Statutes (FCRA, FDCPA) -----------------------

    STATUTE_CASE_NAME = "FCRA & FDCPA — Consumer Protection Statutes"

    def _get_statute_case_id() -> int | None:
        """Get the case ID for the statute reference case."""
        row = _query_one(
            "SELECT id FROM cases WHERE name = %s AND case_type = 'other'",
            (STATUTE_CASE_NAME,),
        )
        return row["id"] if row else None

    @tool(
        "statute_lookup",
        "Look up the authoritative text of a federal consumer protection "
        "statute — FCRA (Fair Credit Reporting Act, 15 USC §§ 1681-1681x) "
        "or FDCPA (Fair Debt Collection Practices Act, 15 USC §§ 1692-1692p). "
        "Pass a section number like '1681a', '1681b', '1692e', '1692g'. "
        "Returns the exact statutory text with proper citation format. "
        "Use this when analyzing credit reports, debt collection issues, "
        "or any consumer financial protection matter. "
        "NEVER paraphrase a statute from memory — always look it up.",
        {
            "type": "object",
            "properties": {
                "section": {
                    "type": "string",
                    "description": "Statute section number. Examples: '1681a' (FCRA definitions), "
                    "'1681b' (permissible purposes), '1681e' (compliance procedures), "
                    "'1692e' (FDCPA false representations), '1692g' (validation of debts). "
                    "Just the number — no '§' or '15 USC' prefix.",
                },
            },
            "required": ["section"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def statute_lookup(args: dict[str, Any]) -> dict[str, Any]:
        section = args["section"].strip()
        try:
            statute_case_id = _get_statute_case_id()
            if not statute_case_id:
                return _error(
                    "Statute reference case not found. Run: "
                    "python -m scripts.statute_ingest --statute all"
                )

            # Search for the section by datalab_id pattern: /section/1681a
            rows = _query(
                """SELECT s.id AS section_id, s.title, s.search_text,
                          s.document_id, d.name AS document_name
                   FROM sections s
                   JOIN documents d ON d.id = s.document_id
                   WHERE d.case_id = %s
                     AND s.datalab_id = %s
                   LIMIT 3""",
                (statute_case_id, f"/section/{section}"),
            )

            if not rows:
                # Try partial match
                rows = _query(
                    """SELECT s.id AS section_id, s.title, s.search_text,
                              s.document_id, d.name AS document_name
                       FROM sections s
                       JOIN documents d ON d.id = s.document_id
                       WHERE d.case_id = %s
                         AND s.datalab_id LIKE %s
                       LIMIT 10""",
                    (statute_case_id, f"/section/{section}%"),
                )

            if not rows:
                return _error(
                    f"Section '{section}' not found in the statute corpus. "
                    f"Available: FCRA (§§ 1681-1681x) and FDCPA (§§ 1692-1692p). "
                    f"Check the section number and try again."
                )

            # Get blocks for full text
            results = []
            for r in rows:
                blocks = _query(
                    """SELECT text_content FROM blocks
                       WHERE section_id = %s ORDER BY id""",
                    (r["section_id"],),
                )
                block_texts = [b["text_content"] for b in blocks if b.get("text_content")]

                results.append({
                    "section": f"§ {section}",
                    "title": r["title"],
                    "statute": r["document_name"],
                    "full_text": (
                        (r["title"] or "") + "\n\n" +
                        "\n".join(block_texts)
                    ) if block_texts else (r["search_text"] or ""),
                    "block_count": len(block_texts),
                })

            return _result({
                "query": section,
                "count": len(results),
                "results": results,
            })

        except Exception as exc:
            return _error(f"statute_lookup failed: {exc}")

    # -- Layer 6.6: Folders --------------------------------------------------

    @tool(
        "list_folders",
        "List folders for the current case, optionally scoped to a workspace "
        "and/or parent folder. Folders are hierarchical — use parent_id to "
        "navigate the tree. Pass parent_id=null for root folders.",
        {
            "type": "object",
            "properties": {
                "parent_id": {
                    "type": "integer",
                    "description": "Parent folder ID. Omit for root folders.",
                },
                "workspace_id": {
                    "type": "integer",
                    "description": "Workspace ID to scope to (optional).",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_folders(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_folders as _list_folders
                rows = _list_folders(
                    conn, case_id,
                    workspace_id=args.get("workspace_id"),
                    parent_id=args.get("parent_id"),
                )
            finally:
                conn.close()
            return _result({"count": len(rows), "folders": rows})
        except Exception as exc:
            return _error(f"list_folders failed: {exc}")

    @tool(
        "create_folder",
        "Create a new folder in the case workspace. Folders can be nested "
        "by setting parent_id. Use this to organize workspace items.",
        {
            "type": "object",
            "properties": {
                "name": {
                    "type": "string",
                    "description": "Folder name.",
                },
                "parent_id": {
                    "type": "integer",
                    "description": "Parent folder ID. Omit for root folder.",
                },
                "workspace_id": {
                    "type": "integer",
                    "description": "Workspace ID (optional).",
                },
            },
            "required": ["name"],
        },
    )
    async def create_folder(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import insert_folder as _insert_folder
                folder_id = _insert_folder(
                    conn,
                    case_id=case_id,
                    name=args["name"],
                    parent_id=args.get("parent_id"),
                    workspace_id=args.get("workspace_id"),
                )
            finally:
                conn.close()
            return _result({"folder_id": folder_id, "name": args["name"]})
        except Exception as exc:
            return _error(f"create_folder failed: {exc}")

    # -- Layer 11: Business Vault -------------------------------------------

    @tool(
        "list_vault_items",
        "List items in the business vault for the current case. "
        "Optionally filter by kind (e.g. 'bank_account', 'net_30', "
        "'insurance_policy', 'vendor', 'lease', 'subscription'). "
        "Returns id, kind, name, status, notes, data, document_count, "
        "and timestamps. Does not return attached documents — use "
        "get_vault_item for that.",
        {
            "type": "object",
            "properties": {
                "kind": {
                    "type": "string",
                    "description": "Filter by kind. Free-form text — common values: "
                    "bank_account, net_30_account, insurance_policy, vendor, "
                    "operating_agreement, lease, subscription, contract, other.",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_vault_items(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import list_vault_items as _list_vault_items
                rows = _list_vault_items(
                    conn,
                    case_id=case_id,
                    kind=args.get("kind"),
                )
            finally:
                conn.close()
            return _result({"count": len(rows), "items": rows})
        except Exception as exc:
            return _error(f"list_vault_items failed: {exc}")

    @tool(
        "get_vault_item",
        "Read a vault item's full details including attached documents. "
        "Use before editing to see the current state.",
        {
            "type": "object",
            "properties": {
                "vault_id": {
                    "type": "integer",
                    "description": "Vault item ID from list_vault_items.",
                },
            },
            "required": ["vault_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_vault_item(args: dict[str, Any]) -> dict[str, Any]:
        vault_id = args["vault_id"]
        try:
            conn = _conn()
            try:
                from core.db import get_vault_item as _get_vault_item
                item = _get_vault_item(conn, vault_id)
            finally:
                conn.close()
            if not item:
                return _error(f"Vault item {vault_id} not found.")
            # Case-scope check: item must belong to this case (or be case-null)
            if item.get("case_id") is not None and item["case_id"] != case_id:
                return _error(f"Vault item {vault_id} not in case {case_id}.")
            return _result({"item": item})
        except Exception as exc:
            return _error(f"get_vault_item failed: {exc}")

    @tool(
        "create_vault_item",
        "Create a new item in the business vault. Use this to record "
        "bank accounts, Net 30 accounts, insurance policies, vendors, "
        "leases, subscriptions, operating agreements, or any other "
        "business document or relationship. "
        "The 'kind' field is free-form. Common values: bank_account, "
        "net_30_account, insurance_policy, vendor, operating_agreement, "
        "lease, subscription, contract, other. "
        "The 'data' field is a JSON object for kind-specific structured "
        "data — populate whatever fields are relevant for that kind. "
        "Example for bank_account: "
        '{"bank":"Chase","type":"checking","account_last_4":"1234","routing":"..."}. '
        "Example for net_30_account: "
        '{"vendor":"Uline","credit_limit":5000,"opened":"2025-03"}.',
        {
            "type": "object",
            "properties": {
                "kind": {
                    "type": "string",
                    "description": "What kind of item this is. Free-form text.",
                },
                "name": {
                    "type": "string",
                    "description": "Display name, e.g. 'Chase Business Checking'.",
                },
                "status": {
                    "type": "string",
                    "description": "Status. Default: 'active'. Common: active, "
                    "inactive, expired, closed, pending.",
                },
                "notes": {
                    "type": "string",
                    "description": "Free-text notes or summary.",
                },
                "data": {
                    "type": "object",
                    "description": "Kind-specific structured data as a JSON object.",
                },
            },
            "required": ["kind", "name"],
        },
    )
    async def create_vault_item(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import insert_vault_item as _insert_vault_item
                item_id = _insert_vault_item(
                    conn,
                    case_id=case_id,
                    kind=args["kind"],
                    name=args["name"],
                    status=args.get("status", "active"),
                    notes=args.get("notes"),
                    data=args.get("data"),
                    created_by="agent",
                )
            finally:
                conn.close()
            return _result({
                "vault_id": item_id,
                "kind": args["kind"],
                "name": args["name"],
                "status": args.get("status", "active"),
            })
        except Exception as exc:
            return _error(f"create_vault_item failed: {exc}")

    @tool(
        "update_vault_item",
        "Modify a vault item — update its kind, name, status, notes, or "
        "structured data. Use get_vault_item first to see current state. "
        "All fields are optional — only provide what changed.",
        {
            "type": "object",
            "properties": {
                "vault_id": {
                    "type": "integer",
                    "description": "Vault item ID to update.",
                },
                "kind": {
                    "type": "string",
                    "description": "New kind value (optional).",
                },
                "name": {
                    "type": "string",
                    "description": "New display name (optional).",
                },
                "status": {
                    "type": "string",
                    "description": "New status (optional).",
                },
                "notes": {
                    "type": "string",
                    "description": "New notes (optional).",
                },
                "data": {
                    "type": "object",
                    "description": "New structured data — replaces entirely (optional).",
                },
            },
            "required": ["vault_id"],
        },
    )
    async def update_vault_item(args: dict[str, Any]) -> dict[str, Any]:
        vault_id = args["vault_id"]
        try:
            # Case-scope check
            conn = _conn()
            try:
                from core.db import get_vault_item as _get_vault_item
                item = _get_vault_item(conn, vault_id)
            finally:
                conn.close()
            if not item:
                return _error(f"Vault item {vault_id} not found.")
            if item.get("case_id") is not None and item["case_id"] != case_id:
                return _error(f"Vault item {vault_id} not in case {case_id}.")

            kwargs = {}
            for field in ("kind", "name", "status", "notes", "data"):
                if field in args:
                    kwargs[field] = args[field]

            if not kwargs:
                return _error("No fields to update.")

            conn = _conn()
            try:
                from core.db import update_vault_item as _update_vault_item
                updated = _update_vault_item(conn, vault_id, **kwargs)
            finally:
                conn.close()

            return _result({
                "vault_id": vault_id,
                "name": updated["name"],
                "kind": updated["kind"],
                "status": updated["status"],
                "updated_at": str(updated.get("updated_at", "")),
            })
        except Exception as exc:
            return _error(f"update_vault_item failed: {exc}")

    @tool(
        "attach_vault_documents",
        "Link existing case documents to a vault item. Use after creating "
        "a vault item to connect supporting documents (e.g. a bank statement "
        "PDF or an insurance policy declaration page). Idempotent — attaching "
        "the same document twice is a no-op.",
        {
            "type": "object",
            "properties": {
                "vault_id": {
                    "type": "integer",
                    "description": "Vault item ID to attach documents to.",
                },
                "document_ids": {
                    "type": "array",
                    "items": {"type": "integer"},
                    "description": "List of document IDs to attach.",
                },
            },
            "required": ["vault_id", "document_ids"],
        },
    )
    async def attach_vault_documents(args: dict[str, Any]) -> dict[str, Any]:
        vault_id = args["vault_id"]
        try:
            # Case-scope check
            conn = _conn()
            try:
                from core.db import get_vault_item as _get_vault_item
                item = _get_vault_item(conn, vault_id)
            finally:
                conn.close()
            if not item:
                return _error(f"Vault item {vault_id} not found.")
            if item.get("case_id") is not None and item["case_id"] != case_id:
                return _error(f"Vault item {vault_id} not in case {case_id}.")

            conn = _conn()
            try:
                from core.db import attach_vault_documents as _attach_vault_documents
                count = _attach_vault_documents(conn, vault_id, args["document_ids"])
            finally:
                conn.close()
            return _result({"vault_id": vault_id, "attached": count})
        except Exception as exc:
            return _error(f"attach_vault_documents failed: {exc}")

    # -- Layer 12: Journal ----------------------------------------------------

    @tool(
        "list_journal_entries",
        "List journal entries for the current case, newest first. "
        "The journal tracks session starts, milestones, decisions, and "
        "findings across sessions. Use this at the START of every session "
        "to understand what was previously worked on. "
        "Optionally filter by entry_type.",
        {
            "type": "object",
            "properties": {
                "limit": {
                    "type": "integer",
                    "description": "Max entries to return (default 20, max 50).",
                },
                "entry_type": {
                    "type": "string",
                    "enum": [
                        "session_start", "session_end",
                        "milestone", "decision", "phase_change",
                        "finding", "note",
                    ],
                    "description": "Filter to a specific entry type (optional).",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_journal_entries(args: dict[str, Any]) -> dict[str, Any]:
        try:
            limit = min(args.get("limit", 20), 50)
            entry_type = args.get("entry_type")
            conn = _conn()
            try:
                from core.db import list_journal_entries as _list_journal
                rows = _list_journal(conn, case_id, limit=limit, entry_type=entry_type)
            finally:
                conn.close()
            return _result({"count": len(rows), "entries": rows})
        except Exception as exc:
            return _error(f"list_journal_entries failed: {exc}")

    @tool(
        "create_journal_entry",
        "Write an entry to the case journal. Use this to record session "
        "starts/ends, milestones reached, decisions made, phase changes, "
        "findings discovered, or general notes. The journal provides "
        "cross-session continuity — the next session reads it to understand "
        "where things stand.\n\n"
        "Entry types:\n"
        "  session_start  — Beginning of a work session\n"
        "  session_end    — End of a work session (summary + next steps)\n"
        "  milestone      — Key accomplishment reached\n"
        "  decision       — Strategic choice and the reasoning behind it\n"
        "  phase_change   — Moving from one stage of work to another\n"
        "  finding        — Discovery worth remembering\n"
        "  note           — General observation",
        {
            "type": "object",
            "properties": {
                "entry_type": {
                    "type": "string",
                    "enum": [
                        "session_start", "session_end",
                        "milestone", "decision", "phase_change",
                        "finding", "note",
                    ],
                    "description": "Type of journal entry.",
                },
                "content": {
                    "type": "string",
                    "description": "Markdown body of the entry. Use headers, "
                    "bullet lists, and structured formatting as appropriate. "
                    "For session_start: note what case/matter, what was "
                    "previously accomplished, what's planned this session. "
                    "For session_end: summarize accomplishments and list "
                    "concrete next steps.",
                },
                "title": {
                    "type": "string",
                    "description": "Optional one-line title summarizing the entry.",
                },
            },
            "required": ["entry_type", "content"],
        },
    )
    async def create_journal_entry(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                from core.db import insert_journal_entry as _insert_journal
                entry_id = _insert_journal(
                    conn,
                    case_id=case_id,
                    entry_type=args["entry_type"],
                    content=args["content"],
                    title=args.get("title"),
                )
            finally:
                conn.close()
            return _result({
                "entry_id": entry_id,
                "entry_type": args["entry_type"],
                "title": args.get("title"),
            })
        except Exception as exc:
            return _error(f"create_journal_entry failed: {exc}")

    @tool(
        "search_vendors",
        "Search the unified vendor registry for small businesses matching "
        "specific criteria. Use this to find vendors for an opportunity, "
        "identify potential teaming partners, or verify a vendor's "
        "certifications and capabilities.\n\n"
        "The registry combines GSA Schedule holders and SBA-certified "
        "businesses. Each vendor has NAICS codes, socioeconomic flags "
        "(small business, woman-owned, SDVOSB, HUBZone, 8a), contact "
        "info, capabilities narratives, and GSA contract vehicle data "
        "when available.\n\n"
        "Common patterns:\n"
        "  - Match vendors to an opportunity: search by the opportunity's "
        "NAICS code + the required set-aside type\n"
        "  - Find SDVOSB IT vendors in Virginia: "
        "naics='541511', state='VA', set_aside='sdvosb'\n"
        "  - Keyword search for capability: q='cybersecurity' + "
        "set_aside='small_business'",
        {
            "type": "object",
            "properties": {
                "q": {
                    "type": "string",
                    "description": "Free-text search across business name "
                    "and capabilities narrative. Use keywords describing "
                    "the product, service, or expertise needed.",
                },
                "naics": {
                    "type": "string",
                    "description": "Filter by NAICS code. Partial match — "
                    "'5415' matches 541511, 541512, 541513, etc. "
                    "Use the opportunity's NAICS code here.",
                },
                "state": {
                    "type": "string",
                    "description": "Two-letter state abbreviation (e.g. 'VA', "
                    "'MD', 'CA') to find vendors in a specific location.",
                },
                "set_aside": {
                    "type": "string",
                    "enum": [
                        "small_business", "woman_owned", "veteran_owned",
                        "sdvosb", "hubzone", "8a",
                    ],
                    "description": "Filter by socioeconomic designation. "
                    "Match this to the opportunity's set-aside requirement.",
                },
                "limit": {
                    "type": "integer",
                    "description": "Max results (default: 10, max: 50).",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def search_vendors(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                limit = min(args.get("limit", 10), 50)
                filters = []
                filter_params: list = []

                if args.get("naics"):
                    filters.append("naics_codes_all LIKE %s")
                    filter_params.append(f"%{args['naics']}%")

                if args.get("state"):
                    filters.append("state = UPPER(%s)")
                    filter_params.append(args["state"])

                set_aside_map = {
                    "small_business": "is_small_business",
                    "woman_owned": "is_woman_owned",
                    "veteran_owned": "is_veteran_owned",
                    "sdvosb": "is_sdvosb",
                    "hubzone": "is_hubzone",
                    "8a": "is_8a",
                }
                sa = args.get("set_aside")
                if sa and sa in set_aside_map:
                    filters.append(f"{set_aside_map[sa]} = TRUE")

                filter_clause = (" AND " + " AND ".join(filters)) if filters else ""

                cols = (
                    "vendor_name, trade_name, source, uei, cage_code, "
                    "contact_name, contact_email, contact_phone, website, "
                    "city, state, county, "
                    "sba_certifications, business_types, "
                    "naics_codes_all, naics_code_primary, capabilities, "
                    "gsa_contract_number, gsa_large_category, gsa_sub_category, "
                    "gsa_option_end_date, gsa_ultimate_end_date, "
                    "is_small_business, is_woman_owned, is_veteran_owned, "
                    "is_sdvosb, is_hubzone, is_8a"
                )

                with conn.cursor() as cur:
                    if args.get("q"):
                        q = args["q"]
                        like_q = f"%{q}%"
                        kw_params = [q, like_q, like_q]

                        if filters:
                            filter_where = " AND ".join(filters)
                            cur.execute(
                                f"SELECT {cols} FROM vendors "
                                f"WHERE {filter_where} AND ("
                                f"to_tsvector('english', coalesce(vendor_name, '') || ' ' || coalesce(capabilities, '')) "
                                f"@@ plainto_tsquery('english', %s) "
                                f"OR vendor_name ILIKE %s "
                                f"OR capabilities ILIKE %s) "
                                f"ORDER BY vendor_name LIMIT %s",
                                filter_params + kw_params + [limit],
                            )
                        else:
                            inner = (
                                f"SELECT id FROM vendors "
                                f"WHERE to_tsvector('english', coalesce(vendor_name, '') || ' ' || coalesce(capabilities, '')) "
                                f"@@ plainto_tsquery('english', %s) "
                                f"UNION ALL "
                                f"SELECT id FROM vendors "
                                f"WHERE vendor_name ILIKE %s "
                                f"UNION ALL "
                                f"SELECT id FROM vendors "
                                f"WHERE capabilities ILIKE %s"
                            )
                            cur.execute(
                                f"SELECT {cols} FROM vendors v "
                                f"INNER JOIN ({inner}) AS m ON v.id = m.id "
                                f"ORDER BY v.vendor_name LIMIT %s",
                                kw_params + [limit],
                            )
                    else:
                        where_clause = f" WHERE {filters[0]}" if filters else ""
                        cur.execute(
                            f"SELECT {cols} FROM vendors{where_clause} "
                            f"ORDER BY vendor_name LIMIT %s",
                            filter_params + [limit],
                        )
                    rows = cur.fetchall()
                    cols = [desc[0] for desc in cur.description]
            finally:
                conn.close()

            vendors = [dict(zip(cols, row)) for row in rows]
            # Convert boolean flags for cleaner output
            for v in vendors:
                for flag in ("is_small_business", "is_woman_owned",
                             "is_veteran_owned", "is_sdvosb",
                             "is_hubzone", "is_8a"):
                    v[flag] = bool(v[flag])

            return _result({"count": len(vendors), "vendors": vendors})
        except Exception as exc:
            return _error(f"search_vendors failed: {exc}")

    # -- Layer 13: PDF Form Filling (T11) ------------------------------------

    @tool(
        "download_document",
        "Download the original binary file for a document from MinIO storage "
        "to a local temp path. Use this when you need to work with the "
        "original PDF, DOCX, or other file — for example, before filling a "
        "PDF form or converting a document. Returns the local file path.\n\n"
        "The file is downloaded to /tmp/vision-downloads/ with the original "
        "filename. Caller is responsible for cleanup after processing.\n\n"
        "Only works for documents that have a storage_path (original files "
        "uploaded through the API or fetched from SAM.gov). Documents "
        "created from inbound email replies do not have storage_paths.",
        {
            "type": "object",
            "properties": {
                "document_id": {
                    "type": "integer",
                    "description": "The document ID to download. Use "
                    "list_documents to find the document_id.",
                },
            },
            "required": ["document_id"],
        },
    )
    async def download_document(args: dict[str, Any]) -> dict[str, Any]:
        """Download a document's original binary from MinIO to a local path."""
        try:
            doc_id = args["document_id"]
            conn = _conn()
            try:
                with conn.cursor() as cur:
                    cur.execute(
                        "SELECT id, name, storage_path FROM documents WHERE id = %s",
                        (doc_id,),
                    )
                    row = cur.fetchone()
            finally:
                conn.close()

            if not row:
                return _error(f"Document {doc_id} not found.")

            _db_id, name, storage_path = row
            if not storage_path:
                return _error(
                    f"Document '{name}' has no storage_path — it was created "
                    f"inline (e.g. email reply) and has no binary file to download."
                )

            from pathlib import Path as _Path
            from ingestion.storage import download_file as _download_file

            dest_dir = _Path("/tmp/vision-downloads")
            dest_dir.mkdir(parents=True, exist_ok=True)

            parts = storage_path.split("/", 1)
            if len(parts) != 2:
                return _error(f"Invalid storage_path: {storage_path}")

            bucket, object_key = parts
            dest_path = dest_dir / name

            _download_file(bucket, object_key, dest_path)

            return _result({
                "document_id": doc_id,
                "name": name,
                "local_path": str(dest_path),
                "size_bytes": dest_path.stat().st_size,
            })
        except Exception as exc:
            return _error(f"download_document failed: {exc}")

    @tool(
        "fill_pdf_form",
        "Fill form fields in a PDF document. Use pymupdf to place text at "
        "specified positions on the page — overwriting blanks, underscores, "
        "or empty cells with the provided values.\n\n"
        "PREREQUISITE: Call download_document first to get the PDF locally, "
        "then pass the local_path to this tool.\n\n"
        "HOW IT WORKS:\n"
        "  1. Opens the PDF at local_path\n"
        "  2. For each field in field_data, searches the document text for "
        "     the field label (e.g. 'UNIT PRICE', '30a. SIGNATURE')\n"
        "  3. Finds the blank/empty area adjacent to the label\n"
        "  4. Draws a white rectangle over the blank and places the value "
        "     text right-aligned in the cell\n"
        "  5. Saves to output_path (defaults to same dir, suffixed with "
        "     '_filled')\n\n"
        "FIELD DATA FORMAT:\n"
        "  {\n"
        "    \"field_label\": \"value\",\n"
        "    \"UNIT PRICE\": \"$48,500.00\",\n"
        "    \"30a. SIGNATURE OF OFFEROR\": \"Jane Doe, CEO\"\n"
        "  }\n\n"
        "TIPS:\n"
        "  - Use EXACT label text from the PDF (copy from get_document_structure "
        "    or search_blocks output)\n"
        "  - For tables: use column header text that appears in the table "
        "    (e.g. 'UNIT PRICE', 'AMOUNT')\n"
        "  - For signature blocks: use the full block label (e.g. "
        "    '30a. SIGNATURE OF OFFEROR/CONTRACTOR')\n"
        "  - For 'if equal, name here' blanks: use the surrounding context "
        "    as the label\n"
        "  - Dollar amounts should be formatted with $ and commas\n"
        "  - Dates should be MM/DD/YYYY format\n\n"
        "LIMITATIONS:\n"
        "  - Works on text-based PDFs (not scanned images) — the PDF must "
        "    have searchable text for label matching\n"
        "  - Text placement is approximate — verify the output visually\n"
        "  - Complex multi-line forms may need manual adjustment",
        {
            "type": "object",
            "properties": {
                "local_path": {
                    "type": "string",
                    "description": "Absolute path to the PDF file. Get this "
                    "from download_document first.",
                },
                "field_data": {
                    "type": "object",
                    "description": "Dict mapping field labels to values. "
                    "Keys are the EXACT text labels as they appear in the "
                    "PDF. Values are the text to place in each field.",
                },
                "output_path": {
                    "type": "string",
                    "description": "Where to save the filled PDF. Defaults "
                    "to the same directory with '_filled' suffix.",
                },
                "page": {
                    "type": "integer",
                    "description": "Specific page to fill (0-indexed). If "
                    "omitted, searches all pages for each label.",
                },
            },
            "required": ["local_path", "field_data"],
        },
    )
    async def fill_pdf_form(args: dict[str, Any]) -> dict[str, Any]:
        """Fill a PDF form using pymupdf."""
        try:
            import fitz  # pymupdf
            from pathlib import Path as _Path

            local_path = _Path(args["local_path"])
            if not local_path.exists():
                return _error(f"File not found: {local_path}")

            field_data = args["field_data"]
            output_path = args.get("output_path")
            target_page = args.get("page")  # None = search all pages

            if not output_path:
                stem = local_path.stem
                output_path = str(local_path.parent / f"{stem}_filled.pdf")

            doc = fitz.open(str(local_path))
            filled = []
            not_found = []

            for label, value in field_data.items():
                found = False
                pages_to_search = (
                    [doc[target_page]] if target_page is not None
                    else [doc[i] for i in range(doc.page_count)]
                )

                for page in pages_to_search:
                    instances = page.search_for(label)
                    if not instances:
                        continue

                    # Use the first match as the anchor
                    label_rect = instances[0]

                    # Search for blank/underscore areas to the right of
                    # the label on the same line
                    blanks = page.search_for("__")
                    if not blanks:
                        blanks = page.search_for("_")

                    # Find the closest blank to the right of the label
                    best_blank = None
                    best_dist = float("inf")
                    for blank_rect in blanks:
                        # Must be on roughly the same line
                        if abs(blank_rect.y0 - label_rect.y0) > 15:
                            continue
                        # Must be to the right (or at) the label
                        if blank_rect.x0 < label_rect.x0 - 5:
                            continue
                        dist = blank_rect.x0 - label_rect.x1
                        if 0 <= dist < best_dist:
                            best_dist = dist
                            best_blank = blank_rect

                    if best_blank:
                        # Expand the blank rect slightly for clean coverage
                        expanded = fitz.Rect(
                            best_blank.x0 - 2, best_blank.y0 - 1,
                            best_blank.x1 + 2, best_blank.y1 + 1,
                        )
                        page.draw_rect(expanded, color=None, fill=(1, 1, 1), width=0)

                        # Right-align the value within the blank area
                        fontsize = 10
                        tw = fitz.get_text_length(str(value), fontname="hebo", fontsize=fontsize)
                        text_x = expanded.x1 - tw - 3
                        text_y = expanded.y1 - 2

                        page.insert_text(
                            fitz.Point(text_x, text_y),
                            str(value),
                            fontname="hebo",
                            fontsize=fontsize,
                            color=(0, 0, 0),
                        )
                        found = True
                        break

                if found:
                    filled.append(label)
                else:
                    not_found.append(label)

            doc.save(output_path, deflate=True)
            doc.close()

            return _result({
                "output_path": output_path,
                "fields_filled": filled,
                "fields_not_found": not_found,
                "total_fields": len(field_data),
            })
        except Exception as exc:
            return _error(f"fill_pdf_form failed: {exc}")

    @tool(
        "upload_filled_document",
        "Upload a filled or modified document to MinIO storage and register "
        "it as a new document in the case. Use this after fill_pdf_form or "
        "convert_docx_to_pdf to persist the result.\n\n"
        "The uploaded file becomes a permanent document in the case, "
        "viewable in the Documents tab. It is tagged with source='filled_form' "
        "to distinguish it from original solicitation documents.\n\n"
        "After upload, the document goes through the standard ingestion "
        "pipeline (OCR, embedding) so its content is searchable.",
        {
            "type": "object",
            "properties": {
                "file_path": {
                    "type": "string",
                    "description": "Absolute path to the filled/modified "
                    "file to upload.",
                },
                "name": {
                    "type": "string",
                    "description": "Display name for the document in the "
                    "case. Use a descriptive name like 'SF 1449 — Filled' "
                    "or 'Proposal Package — Signed'.",
                },
                "document_type": {
                    "type": "string",
                    "enum": ["filled_form", "proposal", "supporting", "other"],
                    "description": "Document type classification. Use "
                    "'filled_form' for filled government forms (SF 1449, "
                    "SF 30, etc.), 'proposal' for proposal documents, "
                    "'supporting' for past performance or certifications, "
                    "'other' for anything else. Default: filled_form.",
                },
            },
            "required": ["file_path", "name"],
        },
    )
    async def upload_filled_document(args: dict[str, Any]) -> dict[str, Any]:
        """Upload a filled/modified file to MinIO and register as a document."""
        try:
            from pathlib import Path as _Path
            from core.db import tx, insert_document
            from ingestion.storage import upload_file as _upload_file
            from ingestion.jobs import enqueue as _enqueue_job

            file_path = _Path(args["file_path"])
            if not file_path.exists():
                return _error(f"File not found: {file_path}")

            name = args["name"]
            doc_type = args.get("document_type", "filled_form")

            # Upload to MinIO
            storage_ref = _upload_file(str(file_path), original_name=name)
            bucket = storage_ref["bucket"]
            object_key = storage_ref["object_key"]
            storage_path = f"{bucket}/{object_key}"
            size_bytes = storage_ref["size_bytes"]

            # Create documents row
            with tx() as conn:
                doc_id = insert_document(
                    conn,
                    case_id=case_id,
                    name=name,
                    page_count=None,
                    source="filled_form" if doc_type == "filled_form" else "user_upload",
                )
                with conn.cursor() as cur:
                    cur.execute(
                        """UPDATE documents
                           SET storage_path = %s,
                               ocr_status = 'pending'
                           WHERE id = %s""",
                        (storage_path, doc_id),
                    )

            # Enqueue ingestion
            job = _enqueue_job(
                case_id=case_id,
                job_type="ingest",
                storage_ref={"bucket": bucket, "object_key": object_key,
                             "original_name": name},
            )

            # Auto-create workspace PDF item so the filled form appears
            # in the Workspace tab alongside the proposal narrative
            from core.db import insert_draft as _insert_draft
            ws_item_id = None
            try:
                with tx() as conn:
                    ws_item_id = _insert_draft(
                        conn,
                        case_id=case_id,
                        name=name,
                        document_type=doc_type,
                        content=[{"document_id": doc_id, "name": name}],
                        created_by="agent",
                        file_type="pdf",
                        folder="artifacts",
                    )
            except Exception:
                pass  # non-fatal — the document still exists

            return _result({
                "document_id": doc_id,
                "name": name,
                "storage_path": storage_path,
                "size_bytes": size_bytes,
                "job_id": job["id"],
                "job_status": job["status"],
                "workspace_item_id": ws_item_id,
            })
        except Exception as exc:
            return _error(f"upload_filled_document failed: {exc}")

    @tool(
        "convert_docx_to_pdf",
        "Convert a DOCX file to PDF format. Uses python-docx to read the "
        "DOCX and pymupdf (fitz) to create a new PDF with the same content.\n\n"
        "PREREQUISITE: Call download_document first to get the DOCX locally, "
        "then pass the local_path to this tool.\n\n"
        "The conversion preserves:\n"
        "  - Paragraphs and headings (h1-h6)\n"
        "  - Tables with cell content\n"
        "  - Bold formatting on paragraphs\n"
        "  - Font sizes from style hierarchy\n\n"
        "Limitations:\n"
        "  - Complex formatting (nested tables, images, headers/footers) "
        "    may not convert perfectly\n"
        "  - For forms with precise layout requirements, prefer filling "
        "    the original PDF directly with fill_pdf_form\n"
        "  - Text-only conversion — images and embedded objects are skipped",
        {
            "type": "object",
            "properties": {
                "local_path": {
                    "type": "string",
                    "description": "Absolute path to the DOCX file. Get "
                    "this from download_document first.",
                },
                "output_path": {
                    "type": "string",
                    "description": "Where to save the PDF. Defaults to the "
                    "same directory with .pdf extension.",
                },
            },
            "required": ["local_path"],
        },
    )
    async def convert_docx_to_pdf(args: dict[str, Any]) -> dict[str, Any]:
        """Convert a DOCX file to PDF."""
        try:
            from pathlib import Path as _Path
            import fitz  # pymupdf

            local_path = _Path(args["local_path"])
            if not local_path.exists():
                return _error(f"File not found: {local_path}")

            output_path = args.get("output_path")
            if not output_path:
                output_path = str(local_path.with_suffix(".pdf"))

            try:
                from docx import Document
            except ImportError:
                return _error("python-docx not installed.")

            docx = Document(str(local_path))

            # Create a new PDF
            pdf_doc = fitz.open()
            page = pdf_doc.new_page(width=612, height=792)  # letter size
            y = 72  # start 1 inch from top
            margin_left = 72
            margin_right = 540  # 612 - 72
            line_height = 14

            for para in docx.paragraphs:
                text = para.text.strip()
                if not text:
                    y += line_height
                    continue

                # Determine font size from style
                fontsize = 11
                style_name = (para.style.name if para.style else "").lower()
                if "heading 1" in style_name:
                    fontsize = 18
                elif "heading 2" in style_name:
                    fontsize = 15
                elif "heading 3" in style_name:
                    fontsize = 13

                # Check bold
                is_bold = any(r.bold for r in para.runs if r.bold)

                # Word-wrap text
                words = text.split()
                lines = []
                current_line = ""
                for word in words:
                    test_line = f"{current_line} {word}".strip()
                    tw = fitz.get_text_length(test_line, fontname="hebo" if is_bold else "helv", fontsize=fontsize)
                    if tw < (margin_right - margin_left):
                        current_line = test_line
                    else:
                        if current_line:
                            lines.append(current_line)
                        current_line = word
                if current_line:
                    lines.append(current_line)

                for line in lines:
                    if y > 720:  # new page
                        page = pdf_doc.new_page(width=612, height=792)
                        y = 72

                    page.insert_text(
                        fitz.Point(margin_left, y + fontsize),
                        line,
                        fontname="hebo" if is_bold else "helv",
                        fontsize=fontsize,
                        color=(0, 0, 0),
                    )
                    y += fontsize + 4

                y += 4  # paragraph spacing

            # Handle tables
            for table in docx.tables:
                y += 12
                if y > 700:
                    page = pdf_doc.new_page(width=612, height=792)
                    y = 72

                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    row_text = " | ".join(cells)
                    tw = fitz.get_text_length(row_text, fontname="helv", fontsize=9)
                    if tw < (margin_right - margin_left):
                        page.insert_text(
                            fitz.Point(margin_left, y + 9),
                            row_text,
                            fontname="helv",
                            fontsize=9,
                            color=(0, 0, 0),
                        )
                    else:
                        # Truncate or wrap
                        page.insert_text(
                            fitz.Point(margin_left, y + 9),
                            row_text[:100],
                            fontname="helv",
                            fontsize=9,
                            color=(0, 0, 0),
                        )
                    y += 16
                y += 8

            pdf_doc.save(output_path, deflate=True)
            pdf_doc.close()

            output_size = _Path(output_path).stat().st_size
            return _result({
                "output_path": output_path,
                "input": str(local_path),
                "size_bytes": output_size,
            })
        except Exception as exc:
            return _error(f"convert_docx_to_pdf failed: {exc}")

    # -- Layer 14: SAM.gov Databank Notices -----------------------------------

    @tool(
        "query_sam_notices",
        "Query the SAM.gov databank of federal contract opportunities. "
        "This contains thousands of notices imported from SAM.gov CSV "
        "exports, searchable by full-text query, NAICS code, set-aside "
        "type, agency, place of performance, date ranges, and more.\n\n"
        "Use this for:\n"
        "  - Finding opportunities matching specific NAICS codes\n"
        "  - Searching for small business set-asides in a state\n"
        "  - Identifying trends (e.g. 'how many IT contracts in VA?') \n"
        "  - Researching what agencies buy in your NAICS\n"
        "  - Finding active solicitations with approaching deadlines\n\n"
        "FILTERS (all optional — combine freely):\n"
        "  q: full-text search across title, description, NAICS, agency\n"
        "  naics_code: exact NAICS code match (e.g. '541511')\n"
        "  naics_description: partial NAICS description match\n"
        "  psc_code: Product/Service Code\n"
        "  contract_opportunity_type: 'Combined Synopsis/Solicitation', "
        "'Sources Sought', 'Award Notice', 'Presolicitation', etc.\n"
        "  current_set_aside: partial match (e.g. 'Small Business', 'SDVOSB')\n"
        "  sub_tier_name: agency name (e.g. 'DEPT OF THE ARMY')\n"
        "  pop_state: 2-letter state code (e.g. 'VA', 'MD')\n"
        "  pop_city: city name\n"
        "  status: 'active' or 'inactive'\n"
        "  response_date_from / response_date_to: date range for due dates\n"
        "  has_attachments: true to only show opportunities with attachments\n"
        "  limit: max results (default 100, max 1000)\n"
        "  offset: for pagination\n\n"
        "EXAMPLES:\n"
        "  - IT services in Virginia: {naics_code: '541511', pop_state: 'VA'}\n"
        "  - Active SDVOSB set-asides: {current_set_aside: 'SDVOSB', status: 'active'}\n"
        "  - Search for roofing: {q: 'roofing repair', pop_state: 'CA'}\n"
        "  - Army solicitations due soon: {sub_tier_name: 'DEPT OF THE ARMY', "
        "response_date_from: '2026-07-22', response_date_to: '2026-08-22'}",
        {
            "type": "object",
            "properties": {
                "q": {
                    "type": "string",
                    "description": "Full-text search query. Searches title, "
                    "description, NAICS, agency, and POC name.",
                },
                "naics_code": {"type": "string"},
                "naics_description": {"type": "string"},
                "psc_code": {"type": "string"},
                "contract_opportunity_type": {"type": "string"},
                "current_set_aside": {"type": "string"},
                "current_set_aside_code": {
                    "type": "string",
                    "description": "Set-aside code: SBA, SDVOSBC, WOSB, "
                    "HZC, 8A, etc.",
                },
                "sub_tier_name": {"type": "string"},
                "pop_state": {"type": "string"},
                "pop_city": {"type": "string"},
                "status": {"type": "string"},
                "awardee_name": {"type": "string"},
                "awardee_uei": {"type": "string"},
                "notice_id": {"type": "string"},
                "response_date_from": {"type": "string"},
                "response_date_to": {"type": "string"},
                "published_date_from": {"type": "string"},
                "published_date_to": {"type": "string"},
                "has_attachments": {"type": "boolean"},
                "ivl_enabled": {"type": "boolean"},
                "limit": {"type": "integer", "description": "Max results "
                    "(default: 100, max: 1000)."},
                "offset": {"type": "integer"},
                "order_by": {
                    "type": "string",
                    "description": "Column to sort by. Default: last_published_date.",
                },
                "order_dir": {
                    "type": "string",
                    "enum": ["ASC", "DESC"],
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def query_sam_notices(args: dict[str, Any]) -> dict[str, Any]:
        """Query the SAM.gov databank notices table."""
        try:
            from api.routes.sam_notices import (
                _SORTABLE_COLUMNS as _cols,
                _FILTERABLE_COLUMNS,
            )

            conn = _conn()
            try:
                where_parts = []
                params: list[Any] = []

                # Full-text search (OR logic for multi-word queries)
                q = args.get("q")
                if q and str(q).strip():
                    words = str(q).strip().split()
                    if len(words) == 1:
                        where_parts.append("search_vector @@ plainto_tsquery('english', %s)")
                        params.append(words[0])
                    else:
                        or_expr = " || ".join(["plainto_tsquery('english', %s)"] * len(words))
                        where_parts.append(f"search_vector @@ ({or_expr})")
                        params.extend(words)

                # String filters
                for col, arg_name, match_type in [
                    ("naics_code", "naics_code", "exact"),
                    ("naics_description", "naics_description", "like"),
                    ("psc_code", "psc_code", "exact"),
                    ("contract_opportunity_type", "contract_opportunity_type", "exact"),
                    ("current_set_aside", "current_set_aside", "like"),
                    ("current_set_aside_code", "current_set_aside_code", "exact"),
                    ("sub_tier_name", "sub_tier_name", "like"),
                    ("pop_state", "pop_state", "exact"),
                    ("pop_city", "pop_city", "like"),
                    ("status", "status", "exact"),
                    ("awardee_name", "awardee_name", "like"),
                    ("awardee_uei", "awardee_uei", "exact"),
                    ("notice_id", "notice_id", "exact"),
                    ("contracting_office", "contracting_office", "like"),
                ]:
                    val = args.get(arg_name)
                    if val and str(val).strip():
                        if match_type == "exact":
                            where_parts.append(f"{col} = %s")
                            params.append(str(val).strip())
                        else:
                            where_parts.append(f"{col} ILIKE %s")
                            params.append(f"%{str(val).strip()}%")

                # Booleans
                if args.get("has_attachments") is True:
                    where_parts.append("attachment_count > 0")
                if args.get("ivl_enabled") is True:
                    where_parts.append("ivl_enabled = true")

                # Date ranges
                for db_col, arg_name in [
                    ("current_response_date", "response_date_from"),
                    ("current_response_date", "response_date_to"),
                    ("last_published_date", "published_date_from"),
                    ("last_published_date", "published_date_to"),
                ]:
                    val = args.get(arg_name)
                    if val:
                        op = ">=" if "from" in arg_name else "<="
                        where_parts.append(f"{db_col} {op} %s")
                        params.append(str(val))

                where_clause = ""
                if where_parts:
                    where_clause = "WHERE " + " AND ".join(where_parts)

                order_col = args.get("order_by", "last_published_date")
                if order_col not in _cols:
                    order_col = "last_published_date"
                order_dir = "DESC" if args.get("order_dir", "DESC").upper() == "DESC" else "ASC"

                limit = min(args.get("limit", 100), 1000)
                offset = max(args.get("offset", 0), 0)

                with conn.cursor() as cur:
                    cur.execute(
                        f"SELECT COUNT(*) FROM sam_notices {where_clause}",
                        tuple(params),
                    )
                    total = cur.fetchone()[0]

                    cur.execute(
                        f"""SELECT id, notice_id, opportunity_title,
                                   contract_opportunity_type,
                                   naics_code, psc_code,
                                   current_set_aside, current_set_aside_code,
                                   sub_tier_name, pop_city, pop_state,
                                   current_response_date, last_published_date,
                                   status, poc_name, poc_email,
                                   awardee_name, attachment_count,
                                   ivl_enabled, description
                            FROM sam_notices
                            {where_clause}
                            ORDER BY {order_col} {order_dir}
                            LIMIT %s OFFSET %s""",
                        tuple(params + [limit, offset]),
                    )
                    rows = cur.fetchall()
                    columns = [desc[0] for desc in cur.description]

                results = [dict(zip(columns, row)) for row in rows]

                # Truncate descriptions for agent context
                for r in results:
                    if r.get("description") and len(r["description"]) > 1000:
                        r["description"] = r["description"][:1000] + "..."

                return _result({
                    "total": total,
                    "limit": limit,
                    "offset": offset,
                    "count": len(results),
                    "results": results,
                })
            finally:
                conn.close()
        except Exception as exc:
            return _error(f"query_sam_notices failed: {exc}")

    # -- Layer 15: Acquisition Gateway Forecasts ------------------------------

    @tool(
        "query_forecast_opportunities",
        "Query the Acquisition Gateway forecast opportunities database. "
        "These are future procurement projections from federal agencies — "
        "what they plan to buy, estimated values, timelines, and set-aside "
        "strategies. Use this to identify upcoming opportunities before "
        "they hit SAM.gov.\n\n"
        "FILTERS (all optional):\n"
        "  q: full-text search across title, description, agency\n"
        "  agency: agency name (e.g. 'Department of Labor')\n"
        "  naics_code: exact NAICS code\n"
        "  set_aside: partial match (e.g. 'Small Business', 'To Be Determined')\n"
        "  fiscal_year: fiscal year (e.g. '2026')\n"
        "  estimated_value_text: value range (e.g. 'Below $150K')\n"
        "  value_under: max estimated value in dollars (e.g. 350000 for SAT)\n"
        "  value_over: min estimated value in dollars (e.g. 1000000 for $1M+)\n"
        "  limit: max results (default 100, max 1000)\n\n"
        "EXAMPLES:\n"
        "  - IT forecasts for 2026: {q: 'IT services', fiscal_year: '2026'}\n"
        "  - Under SAT + SB: {value_under: 350000, set_aside: 'Small Business'}\n"
        "  - VA forecasts by agency: {agency: 'Veterans', fiscal_year: '2026'}",
        {
            "type": "object",
            "properties": {
                "q": {"type": "string", "description": "Full-text search."},
                "agency": {"type": "string"},
                "naics_code": {"type": "string"},
                "set_aside": {"type": "string"},
                "fiscal_year": {"type": "string"},
                "estimated_value_text": {"type": "string"},
                "value_under": {"type": "number", "description": "Max value in dollars (e.g. 350000)."},
                "value_over": {"type": "number", "description": "Min value in dollars (e.g. 1000000)."},
                "office": {"type": "string"},
                "place_of_performance": {"type": "string"},
                "limit": {"type": "integer"},
                "offset": {"type": "integer"},
                "order_by": {"type": "string"},
                "order_dir": {"type": "string", "enum": ["ASC", "DESC"]},
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def query_forecast_opportunities(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                where_parts = []
                params: list[Any] = []

                q = args.get("q")
                if q and str(q).strip():
                    words = str(q).strip().split()
                    if len(words) == 1:
                        where_parts.append("search_vector @@ plainto_tsquery('english', %s)")
                        params.append(words[0])
                    else:
                        or_expr = " || ".join(["plainto_tsquery('english', %s)"] * len(words))
                        where_parts.append(f"search_vector @@ ({or_expr})")
                        params.extend(words)

                # Numeric value filters
                val_under = args.get("value_under")
                if val_under is not None:
                    where_parts.append("(estimated_value_high <= %s OR (estimated_value_high IS NULL AND estimated_value_low <= %s))")
                    params.extend([float(val_under), float(val_under)])
                val_over = args.get("value_over")
                if val_over is not None:
                    where_parts.append("estimated_value_low >= %s")
                    params.append(float(val_over))

                for col, arg_name, match_type in [
                    ("agency", "agency", "like"), ("naics_code", "naics_code", "exact"),
                    ("set_aside", "set_aside", "like"), ("fiscal_year", "fiscal_year", "exact"),
                    ("estimated_value_text", "estimated_value_text", "like"),
                    ("office", "office", "like"), ("place_of_performance", "place_of_performance", "like"),
                ]:
                    val = args.get(arg_name)
                    if val and str(val).strip():
                        if match_type == "exact":
                            where_parts.append(f"{col} = %s")
                            params.append(str(val).strip())
                        else:
                            where_parts.append(f"{col} ILIKE %s")
                            params.append(f"%{str(val).strip()}%")

                where_clause = "WHERE " + " AND ".join(where_parts) if where_parts else ""
                limit = min(args.get("limit", 100), 1000)
                offset = max(args.get("offset", 0), 0)

                with conn.cursor() as cur:
                    cur.execute(f"SELECT COUNT(*) FROM forecast_opportunities {where_clause}", tuple(params))
                    total = cur.fetchone()[0]

                    cur.execute(
                        f"""SELECT id, title, description, source_url,
                                   agency, office, naics_code, naics_description,
                                   set_aside, place_of_performance, period_of_performance,
                                   fiscal_year, estimated_value_text,
                                   estimated_value_low, estimated_value_high,
                                   created_date, last_updated_date
                            FROM forecast_opportunities {where_clause}
                            ORDER BY created_date DESC LIMIT %s OFFSET %s""",
                        tuple(params + [limit, offset]),
                    )
                    rows = cur.fetchall()
                    columns = [desc[0] for desc in cur.description]

                results = [dict(zip(columns, row)) for row in rows]
                for r in results:
                    if r.get("description") and len(r["description"]) > 1000:
                        r["description"] = r["description"][:1000] + "..."

                return _result({"total": total, "limit": limit, "offset": offset, "count": len(results), "results": results})
            finally:
                conn.close()
        except Exception as exc:
            return _error(f"query_forecast_opportunities failed: {exc}")

    # -- Layer 0: Job Queue Inspection & Repair ------------------------------

    @tool(
        "list_jobs",
        "List background jobs for the current case. Use this to check the "
        "status of ingestion, solicitation triage, vendor matching, and "
        "other async operations. Returns job id, type, status, progress, "
        "error messages, and timestamps.\n\n"
        "The worker processes jobs in order. Failed jobs have error_message "
        "set. Stuck jobs (status='processing' for >10 min) may need a retry.\n\n"
        "Filter by status to find problems: 'failed' shows what broke, "
        "'processing' shows what's running, 'queued' shows what's waiting.",
        {
            "type": "object",
            "properties": {
                "status": {
                    "type": "string",
                    "enum": ["queued", "processing", "complete", "failed"],
                    "description": "Filter by status.",
                },
                "job_type": {
                    "type": "string",
                    "description": "Filter by job type: ingest, enrich, "
                    "synthesize, sam_fetch, solicitation_triage, "
                    "vendor_matching, inbound_email, etc.",
                },
                "limit": {
                    "type": "integer",
                    "description": "Max jobs to return (default: 20, max: 100).",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_jobs(args: dict[str, Any]) -> dict[str, Any]:
        try:
            from ingestion.jobs import list_jobs as _list_jobs
            status = args.get("status")
            job_type = args.get("job_type")
            limit = min(args.get("limit", 20), 100)

            jobs = _list_jobs(case_id=case_id, status=status, limit=limit)
            if job_type:
                jobs = [j for j in jobs if j.get("job_type") == job_type]

            # Strip large fields for readability
            for j in jobs:
                j.pop("storage_ref", None)
                if j.get("metadata") and isinstance(j["metadata"], dict):
                    # Keep only key metadata fields
                    j["metadata"] = {
                        k: v for k, v in j["metadata"].items()
                        if k in ("vendor_match_id", "sender", "subject",
                                 "original_name", "batch_id")
                    } or None

            return _result({
                "case_id": case_id,
                "count": len(jobs),
                "jobs": jobs,
            })
        except Exception as exc:
            return _error(f"list_jobs failed: {exc}")

    @tool(
        "get_job",
        "Get full details for a specific job, including its error message "
        "and metadata. Use this after list_jobs to inspect a failed or "
        "stuck job before deciding whether to retry it.",
        {
            "type": "object",
            "properties": {
                "job_id": {
                    "type": "integer",
                    "description": "Job ID from list_jobs.",
                },
            },
            "required": ["job_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def get_job(args: dict[str, Any]) -> dict[str, Any]:
        try:
            from ingestion.jobs import get_job as _get_job
            job_id = args["job_id"]
            job = _get_job(job_id)
            if not job:
                return _error(f"Job {job_id} not found.")
            if job.get("case_id") != case_id:
                return _error(f"Job {job_id} does not belong to this case.")
            return _result({"job": job})
        except Exception as exc:
            return _error(f"get_job failed: {exc}")

    @tool(
        "retry_job",
        "Reset a failed or stuck job back to 'queued' so the worker picks "
        "it up again. Use this when:\n"
        "  - A job failed with a transient error (network timeout, API rate limit)\n"
        "  - A job is stuck in 'processing' state for >10 min (worker crashed)\n"
        "  - Vendor matching or solicitation triage didn't complete\n\n"
        "The job's error_message is cleared, attempts counter is preserved, "
        "and status is reset to 'queued'. The worker will claim it on the "
        "next poll cycle.\n\n"
        "Only works on jobs in 'failed' or 'processing' status. Completed "
        "jobs cannot be retried (they succeeded). Queued jobs don't need "
        "retry (they're already waiting).",
        {
            "type": "object",
            "properties": {
                "job_id": {
                    "type": "integer",
                    "description": "Job ID from list_jobs.",
                },
            },
            "required": ["job_id"],
        },
    )
    async def retry_job(args: dict[str, Any]) -> dict[str, Any]:
        try:
            from core.db import tx
            job_id = args["job_id"]

            with tx() as conn:
                with conn.cursor() as cur:
                    # Verify job exists and belongs to this case
                    cur.execute(
                        "SELECT id, status, case_id FROM jobs WHERE id = %s",
                        (job_id,),
                    )
                    row = cur.fetchone()
                    if not row:
                        return _error(f"Job {job_id} not found.")
                    _id, status, job_case_id = row
                    if job_case_id != case_id:
                        return _error(f"Job {job_id} does not belong to this case.")
                    if status not in ("failed", "processing"):
                        return _error(
                            f"Job {job_id} is '{status}' — only 'failed' or "
                            f"'processing' jobs can be retried."
                        )

                    cur.execute(
                        """UPDATE jobs
                           SET status = 'queued',
                               error_message = NULL,
                               updated_at = now()
                           WHERE id = %s
                           RETURNING id, status, job_type, attempts, error_message""",
                        (job_id,),
                    )
                    updated = cur.fetchone()
                    columns = [d[0] for d in cur.description]

            return _result({
                "retried": True,
                "job": dict(zip(columns, updated)),
            })
        except Exception as exc:
            return _error(f"retry_job failed: {exc}")

    # -- Layer 15b: Summarization --------------------------------------------

    _DIMENSIONS_SAM = ["agency", "naics_code", "set_aside_code", "notice_type", "state", "office"]
    _DIMENSIONS_FC = ["agency", "naics_code", "set_aside", "fiscal_year", "value_range", "office", "state"]

    @tool(
        "summarize_sam_notices",
        "Get aggregate breakdowns of SAM.gov notices by a dimension. "
        "Returns counts, SB set-aside counts, unrestricted counts, and "
        "solicitation counts grouped by the specified dimension. "
        "Optionally filter before aggregating.\n\n"
        "Dimensions: agency, naics_code, set_aside_code, notice_type, state, office\n\n"
        "Use this to answer questions like:\n"
        "  - 'Which agencies have the most active solicitations?'\n"
        "  - 'What NAICS codes are most common in SB set-asides?'\n"
        "  - 'Break down active notices by state'\n"
        "  - 'Show me the set-aside distribution for the Army'",
        {
            "type": "object",
            "properties": {
                "group_by": {
                    "type": "string",
                    "enum": _DIMENSIONS_SAM,
                    "description": "Dimension to group by.",
                },
                "agency": {"type": "string", "description": "Optional: filter by agency (ILIKE)."},
                "naics_code": {"type": "string", "description": "Optional: filter by NAICS."},
                "set_aside_code": {"type": "string", "description": "Optional: filter by set-aside code (SBA, SDVOSBC, etc.)."},
                "notice_type": {"type": "string", "description": "Optional: filter by notice type."},
                "status": {"type": "string", "description": "Optional: filter by status (default: active)."},
                "limit": {"type": "integer", "description": "Max groups to return (default: 20)."},
            },
            "required": ["group_by"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def summarize_sam_notices(args: dict[str, Any]) -> dict[str, Any]:
        try:
            group_by = args["group_by"]
            col_map = {
                "agency": "sub_tier_name", "naics_code": "naics_code",
                "set_aside_code": "current_set_aside_code", "notice_type": "contract_opportunity_type",
                "state": "pop_state", "office": "contracting_office",
            }
            col = col_map[group_by]
            limit = min(args.get("limit", 20), 100)

            where_parts = ["status = 'active'"]
            params: list[Any] = []
            for arg_name, db_col in [("agency", "sub_tier_name"), ("naics_code", "naics_code"),
                                       ("set_aside_code", "current_set_aside_code"),
                                       ("notice_type", "contract_opportunity_type")]:
                val = args.get(arg_name)
                if val and str(val).strip():
                    where_parts.append(f"{db_col} ILIKE %s")
                    params.append(f"%{str(val).strip()}%")
            if args.get("status"):
                where_parts[0] = f"status = %s"
                params.insert(0, args["status"])

            where_clause = "WHERE " + " AND ".join(where_parts)

            conn = _conn()
            try:
                with conn.cursor() as cur:
                    cur.execute(f"""
                        SELECT {col} as dimension, COUNT(*) as total,
                               COUNT(*) FILTER (WHERE current_set_aside_code = 'SBA') as sba,
                               COUNT(*) FILTER (WHERE current_set_aside_code = 'SDVOSBC') as sdvosb,
                               COUNT(*) FILTER (WHERE current_set_aside_code IS NULL OR current_set_aside_code = '' OR current_set_aside_code = 'NONE') as unrestricted,
                               COUNT(*) FILTER (WHERE contract_opportunity_type = 'Combined Synopsis/Solicitation') as sols,
                               COUNT(*) FILTER (WHERE contract_opportunity_type = 'Sources Sought') as ss,
                               COUNT(*) FILTER (WHERE current_response_date IS NOT NULL AND current_response_date >= NOW()) as upcoming,
                               COUNT(*) FILTER (WHERE current_response_date IS NOT NULL AND current_response_date >= NOW() AND current_response_date <= NOW() + INTERVAL '30 days') as due_30d
                        FROM sam_notices {where_clause}
                        GROUP BY {col} ORDER BY total DESC LIMIT %s
                    """, tuple(params + [limit]))
                    rows = cur.fetchall()
                    columns = [d[0] for d in cur.description]
                results = [dict(zip(columns, r)) for r in rows]
                return _result({"group_by": group_by, "groups": len(results), "results": results})
            finally:
                conn.close()
        except Exception as exc:
            return _error(f"summarize_sam_notices failed: {exc}")

    @tool(
        "summarize_forecasts",
        "Get aggregate breakdowns of forecast opportunities by a dimension. "
        "Returns counts, SB set-aside counts, under-SAT counts, $1M+ counts, "
        "and TBD counts grouped by the specified dimension. "
        "Optionally filter before aggregating.\n\n"
        "Dimensions: agency, naics_code, set_aside, fiscal_year, value_range, office, state\n"
        "'value_range' groups by estimated_value_text buckets. Other dimensions "
        "return standard breakdowns.\n\n"
        "Use this to answer questions like:\n"
        "  - 'Which agencies have the most forecast opportunities?'\n"
        "  - 'Break down VA forecasts by NAICS'\n"
        "  - 'What's the value distribution for Interior?'\n"
        "  - 'Show SB vs unrestricted by fiscal year'",
        {
            "type": "object",
            "properties": {
                "group_by": {
                    "type": "string",
                    "enum": _DIMENSIONS_FC,
                    "description": "Dimension to group by.",
                },
                "agency": {"type": "string", "description": "Optional: filter by agency (ILIKE)."},
                "naics_code": {"type": "string", "description": "Optional: filter by NAICS."},
                "set_aside": {"type": "string", "description": "Optional: filter by set-aside."},
                "fiscal_year": {"type": "string", "description": "Optional: filter by fiscal year."},
                "value_under": {"type": "number", "description": "Optional: max estimated value."},
                "value_over": {"type": "number", "description": "Optional: min estimated value."},
                "limit": {"type": "integer", "description": "Max groups to return (default: 20)."},
            },
            "required": ["group_by"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def summarize_forecasts(args: dict[str, Any]) -> dict[str, Any]:
        try:
            group_by = args["group_by"]
            if group_by == "value_range":
                col = "estimated_value_text"
            else:
                col_map = {
                    "agency": "agency", "naics_code": "naics_code",
                    "set_aside": "set_aside", "fiscal_year": "fiscal_year",
                    "office": "office", "state": "place_of_performance",
                }
                col = col_map[group_by]
            limit = min(args.get("limit", 20), 100)

            where_parts = []
            params: list[Any] = []
            for arg_name, db_col, match in [
                ("agency", "agency", "like"), ("naics_code", "naics_code", "exact"),
                ("set_aside", "set_aside", "like"), ("fiscal_year", "fiscal_year", "exact"),
            ]:
                val = args.get(arg_name)
                if val and str(val).strip():
                    where_parts.append(f"{db_col} {'ILIKE' if match == 'like' else '='} %s")
                    params.append(f"%{str(val).strip()}%" if match == "like" else str(val).strip())
            val_under = args.get("value_under")
            if val_under is not None:
                where_parts.append("(estimated_value_high <= %s OR (estimated_value_high IS NULL AND estimated_value_low <= %s))")
                params.extend([float(val_under), float(val_under)])
            val_over = args.get("value_over")
            if val_over is not None:
                where_parts.append("estimated_value_low >= %s")
                params.append(float(val_over))

            where_clause = "WHERE " + " AND ".join(where_parts) if where_parts else ""

            conn = _conn()
            try:
                with conn.cursor() as cur:
                    cur.execute(f"""
                        SELECT {col} as dimension, COUNT(*) as total,
                               COUNT(*) FILTER (WHERE estimated_value_high <= 350000) as under_sat,
                               COUNT(*) FILTER (WHERE estimated_value_low >= 1000000) as over_1m,
                               COUNT(*) FILTER (WHERE set_aside ILIKE '%small business%') as sb,
                               COUNT(*) FILTER (WHERE set_aside ILIKE '%veteran%' OR set_aside ILIKE '%sdvosb%') as sdvosb,
                               COUNT(*) FILTER (WHERE set_aside = 'To Be Determined') as tbd,
                               ROUND(AVG(estimated_value_high) FILTER (WHERE estimated_value_high IS NOT NULL)) as avg_high_value
                        FROM forecast_opportunities {where_clause}
                        GROUP BY {col} ORDER BY total DESC LIMIT %s
                    """, tuple(params + [limit]))
                    rows = cur.fetchall()
                    columns = [d[0] for d in cur.description]
                results = [dict(zip(columns, r)) for r in rows]
                return _result({"group_by": group_by, "groups": len(results), "results": results})
            finally:
                conn.close()
        except Exception as exc:
            return _error(f"summarize_forecasts failed: {exc}")

    # -- Layer 16: Saved Reports ----------------------------------------------

    @tool(
        "list_reports",
        "List saved filter-preset reports for the current case. Reports are "
        "named, reusable queries that persist filter criteria for the "
        "Forecasts and Sam Notices tabs. Use this to see what reports "
        "already exist before creating or executing one.\n\n"
        "Optionally filter by data_source ('forecasts' or 'sam_notices').",
        {
            "type": "object",
            "properties": {
                "data_source": {
                    "type": "string",
                    "enum": ["forecasts", "sam_notices"],
                    "description": "Optional: filter by data source.",
                },
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def list_reports(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                ds = args.get("data_source")
                with conn.cursor() as cur:
                    if ds:
                        cur.execute(
                            """SELECT id, name, data_source, query_filters,
                                      sort_by, sort_dir, created_by, created_at
                               FROM saved_reports
                               WHERE case_id = %s AND data_source = %s
                               ORDER BY updated_at DESC""",
                            (case_id, ds),
                        )
                    else:
                        cur.execute(
                            """SELECT id, name, data_source, query_filters,
                                      sort_by, sort_dir, created_by, created_at
                               FROM saved_reports
                               WHERE case_id = %s
                               ORDER BY updated_at DESC""",
                            (case_id,),
                        )
                    rows = cur.fetchall()
                    columns = [d[0] for d in cur.description]
                return _result({"count": len(rows), "reports": [dict(zip(columns, r)) for r in rows]})
            finally:
                conn.close()
        except Exception as exc:
            return _error(f"list_reports failed: {exc}")

    @tool(
        "create_report",
        "Save a new filter-preset report for later use. The query_filters "
        "object should be the exact filter object you would pass to "
        "query_forecast_opportunities or query_sam_notices.\n\n"
        "data_source must be 'forecasts' or 'sam_notices'.\n\n"
        "EXAMPLES:\n"
        "  {name: 'IT Under 150K', data_source: 'forecasts', "
        "   query_filters: {q: 'IT software', value_under: 150000}, "
        "   sort_by: 'estimated_value_low'}\n"
        "  {name: 'Active SB Solicitations', data_source: 'sam_notices', "
        "   query_filters: {current_set_aside_code: 'SBA', status: 'active'}}",
        {
            "type": "object",
            "properties": {
                "name": {"type": "string", "description": "Report display name."},
                "data_source": {
                    "type": "string",
                    "enum": ["forecasts", "sam_notices"],
                    "description": "Which data table this report queries.",
                },
                "query_filters": {
                    "type": "object",
                    "description": "Filter object matching the query tool's parameters.",
                },
                "sort_by": {"type": "string", "description": "Optional: column to sort by."},
                "sort_dir": {"type": "string", "enum": ["ASC", "DESC"]},
            },
            "required": ["name", "data_source", "query_filters"],
        },
    )
    async def create_report(args: dict[str, Any]) -> dict[str, Any]:
        try:
            import json as _json
            with tx() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        """INSERT INTO saved_reports (case_id, name, data_source, query_filters, sort_by, sort_dir)
                           VALUES (%s, %s, %s, %s::jsonb, %s, %s)
                           RETURNING id, name, data_source, sort_by, sort_dir, created_at""",
                        (
                            case_id, args["name"].strip(), args["data_source"],
                            _json.dumps(args["query_filters"]),
                            args.get("sort_by"), args.get("sort_dir", "ASC"),
                        ),
                    )
                    row = cur.fetchone()
                    columns = [d[0] for d in cur.description]
            return _result({"created": True, "report": dict(zip(columns, row))})
        except Exception as exc:
            return _error(f"create_report failed: {exc}")

    @tool(
        "execute_report",
        "Run a saved report by ID. Loads the stored query_filters and "
        "executes them against the appropriate data source. Returns "
        "paginated results.\n\n"
        "Use list_reports first to find the report ID, then call this "
        "to get the actual data.",
        {
            "type": "object",
            "properties": {
                "report_id": {"type": "integer", "description": "Report ID from list_reports."},
                "limit": {"type": "integer", "description": "Override page size."},
                "offset": {"type": "integer", "description": "Override page offset."},
            },
            "required": ["report_id"],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def execute_report(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                with conn.cursor() as cur:
                    cur.execute(
                        "SELECT * FROM saved_reports WHERE id = %s AND case_id = %s",
                        (args["report_id"], case_id),
                    )
                    row = cur.fetchone()
                    if not row:
                        return _error(f"Report {args['report_id']} not found in this case.")
                    columns = [d[0] for d in cur.description]
                    report = dict(zip(columns, row))
            finally:
                conn.close()

            filters = report.get("query_filters") or {}
            if args.get("limit"):
                filters["limit"] = args["limit"]
            if args.get("offset"):
                filters["offset"] = args["offset"]
            if report.get("sort_by"):
                filters["order_by"] = report["sort_by"]
            if report.get("sort_dir"):
                filters["order_dir"] = report["sort_dir"]

            if report["data_source"] == "forecasts":
                return await query_forecast_opportunities(filters)
            else:
                return await query_sam_notices(filters)
        except Exception as exc:
            return _error(f"execute_report failed: {exc}")

    @tool(
        "delete_report",
        "Delete a saved report by ID. Irreversible.",
        {
            "type": "object",
            "properties": {
                "report_id": {"type": "integer", "description": "Report ID from list_reports."},
            },
            "required": ["report_id"],
        },
    )
    async def delete_report(args: dict[str, Any]) -> dict[str, Any]:
        try:
            with tx() as conn:
                with conn.cursor() as cur:
                    cur.execute(
                        "DELETE FROM saved_reports WHERE id = %s AND case_id = %s",
                        (args["report_id"], case_id),
                    )
                    if cur.rowcount == 0:
                        return _error(f"Report {args['report_id']} not found in this case.")
            return _result({"deleted": args["report_id"]})
        except Exception as exc:
            return _error(f"delete_report failed: {exc}")

    # -- Layer 17: Solicitation Creation --------------------------------------

    @tool(
        "create_solicitation",
        "Create a new federal solicitation from a SAM.gov URL. This enqueues "
        "a sam_fetch job that downloads all documents and metadata, then "
        "auto-triggers triage and vendor matching when complete.",
        {
            "type": "object",
            "properties": {
                "url": {
                    "type": "string",
                    "description": "Full SAM.gov opportunity URL (e.g. https://sam.gov/opp/abc123/view)",
                },
            },
            "required": ["url"],
        },
    )
    async def create_solicitation_from_url(args: dict) -> dict:
        """Create a solicitation from a SAM.gov URL, enqueue sam_fetch."""
        try:
            from core.solicitation import SolicitationManager, DuplicateNoticeError
            from ingestion.jobs import enqueue
            from ingestion.sam_client import extract_notice_id

            url = args["url"]
            notice_id = extract_notice_id(url)
            if not notice_id:
                return _error("Could not extract notice ID from URL. Make sure it's a valid SAM.gov opportunity URL.")

            mgr = SolicitationManager()
            try:
                sol = mgr.create(source_type="federal", url=url, notice_id=notice_id)
            except DuplicateNoticeError as e:
                return _error(str(e))

            job = enqueue(
                case_id=sol["case_id"],
                job_type="sam_fetch",
                metadata={"solicitation_id": sol["id"], "notice_id": notice_id},
            )
            return _result({
                "solicitation_id": sol["id"],
                "case_id": sol["case_id"],
                "notice_id": notice_id,
                "job_id": job["id"],
                "title": sol.get("title", ""),
            })
        except Exception as exc:
            return _error(f"create_solicitation failed: {exc}")

    # -- Layer 17: GA DOAS Opportunities ------------------------------------

    @tool(
        "query_ga_doas_opportunities",
        "Query Georgia DOAS state and local procurement opportunities. "
        "These are bids, RFPs, and RFQs from Georgia cities, counties, "
        "school boards, and state agencies.\n\n"
        "FILTERS: q (full-text), government_entity (ILIKE), event_id, status",
        {
            "type": "object",
            "properties": {
                "q": {"type": "string"},
                "government_entity": {"type": "string"},
                "event_id": {"type": "string"},
                "status": {"type": "string"},
                "limit": {"type": "integer"},
                "offset": {"type": "integer"},
            },
            "required": [],
        },
        annotations=ToolAnnotations(readOnlyHint=True),
    )
    async def query_ga_doas_opportunities(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn = _conn()
            try:
                where_parts = []; params: list[Any] = []
                q = args.get("q")
                if q and str(q).strip():
                    words = str(q).strip().split()
                    if len(words) == 1:
                        where_parts.append("search_vector @@ plainto_tsquery('english', %s)"); params.append(words[0])
                    else:
                        or_expr = " || ".join(["plainto_tsquery('english', %s)"] * len(words))
                        where_parts.append(f"search_vector @@ ({or_expr})"); params.extend(words)
                for col, arg, match in [("government_entity","government_entity","like"),("event_id","event_id","exact"),("status","status","exact")]:
                    v = args.get(arg)
                    if v and str(v).strip():
                        where_parts.append(f"{col} {'ILIKE' if match=='like' else '='} %s")
                        params.append(f"%{str(v).strip()}%" if match=="like" else str(v).strip())
                where_clause = "WHERE " + " AND ".join(where_parts) if where_parts else ""
                limit = min(args.get("limit",100),1000); offset = max(args.get("offset",0),0)
                with conn.cursor() as cur:
                    cur.execute(f"SELECT COUNT(*) FROM ga_doas_opportunities {where_clause}", tuple(params))
                    total = cur.fetchone()[0]
                    cur.execute(f"SELECT * FROM ga_doas_opportunities {where_clause} ORDER BY end_date ASC LIMIT %s OFFSET %s", tuple(params+[limit,offset]))
                    rows = cur.fetchall(); cols = [d[0] for d in cur.description]
                return _result({"total":total,"limit":limit,"offset":offset,"count":len(rows),"results":[dict(zip(cols,r)) for r in rows]})
            finally: conn.close()
        except Exception as exc: return _error(f"query_ga_doas_opportunities failed: {exc}")

    # -- Layer 18: DIBBS RFQs ------------------------------------------------

    @tool("query_dibbs_rfqs", "Query DIBBS RFQ opportunities. These are DLA rapid-turnaround RFQs for NSN items. Filter by NSN, FSC code, nomenclature, status.",
        {"type":"object","properties":{"q":{"type":"string"},"nsn":{"type":"string"},"fsc_code":{"type":"string"},"solicitation":{"type":"string"},"status":{"type":"string"},"limit":{"type":"integer"},"offset":{"type":"integer"}},"required":[]},
        annotations=ToolAnnotations(readOnlyHint=True))
    async def query_dibbs_rfqs(args: dict[str, Any]) -> dict[str, Any]:
        try:
            conn=_conn()
            try:
                where=[]; params=[]
                q=args.get("q")
                if q and str(q).strip():
                    w=str(q).strip().split()
                    if len(w)==1: where.append("search_vector @@ plainto_tsquery('english',%s)"); params.append(w[0])
                    else: where.append(f"search_vector @@ ({' || '.join(['plainto_tsquery(%s)']*len(w))})"); params.extend(w)
                for col,arg,m in [("nsn","nsn","exact"),("fsc_code","fsc_code","exact"),("solicitation","solicitation","exact"),("status","status","exact")]:
                    v=args.get(arg)
                    if v and str(v).strip(): where.append(f"{col}={'=' if m=='exact' else 'ILIKE'} %s"); params.append(str(v).strip())
                wc="WHERE "+" AND ".join(where) if where else ""
                lim=min(args.get("limit",100),1000); off=max(args.get("offset",0),0)
                with conn.cursor() as cur:
                    cur.execute(f"SELECT COUNT(*) FROM dibbs_rfqs {wc}",tuple(params)); total=cur.fetchone()[0]
                    cur.execute(f"SELECT * FROM dibbs_rfqs {wc} ORDER BY return_by ASC LIMIT %s OFFSET %s",tuple(params+[lim,off]))
                    rows=cur.fetchall(); cols=[d[0] for d in cur.description]
                return _result({"total":total,"limit":lim,"offset":off,"count":len(rows),"results":[dict(zip(cols,r)) for r in rows]})
            finally: conn.close()
        except Exception as exc: return _error(f"query_dibbs_rfqs failed: {exc}")

    # -- Build server --------------------------------------------------------

    return create_sdk_mcp_server(
        name="vision",
        version="1.0.0",
        tools=[
            list_jobs,
            get_job,
            create_solicitation_from_url,
            retry_job,
            get_case,
            list_workspaces,
            create_workspace,
            list_documents,
            search_blocks,
            semantic_search,
            search_hybrid,
            get_document_structure,
            search_sections,
            get_block_context,
            get_blocks_in_section,
            get_strategies,
            get_strategy_tree,
            list_drafts,
            get_draft,
            create_draft,
            update_draft,
            list_workspace_items,
            get_workspace_item,
            create_workspace_item,
            update_workspace_item,
            list_folders,
            create_folder,
            list_tasks,
            create_task,
            update_task,
            delete_task,
            create_calendar_event,
            list_calendar_events,
            get_calendar_event,
            create_reminder,
            list_reminders,
            get_reminder,
            list_correspondence_threads,
            create_correspondence_thread,
            update_correspondence_thread,
            list_correspondence_items,
            create_correspondence_item,
            update_correspondence_item,
            delete_correspondence_item,
            list_company_profiles,
            get_company_profile,
            get_case_profile,
            create_knowledge_entry,
            search_knowledge,
            list_knowledge_tags,
            far_lookup,
            far_status,
            statute_lookup,
            list_vault_items,
            get_vault_item,
            create_vault_item,
            update_vault_item,
            attach_vault_documents,
            list_journal_entries,
            create_journal_entry,
            search_vendors,
            download_document,
            fill_pdf_form,
            upload_filled_document,
            convert_docx_to_pdf,
            query_sam_notices,
            query_forecast_opportunities,
            summarize_sam_notices,
            summarize_forecasts,
            list_reports,
            create_report,
            execute_report,
            delete_report,
            query_ga_doas_opportunities,
            query_dibbs_rfqs,
        ],
    )

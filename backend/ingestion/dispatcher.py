"""
Vision — Document Ingestion Pipeline.

Takes a PDF + case_id, submits to DataLab OCR, polls for completion,
and normalizes the structured JSON into the evidence store:
documents → sections → blocks → block_headings.

Port of:
  - datalab_eval/run_convert.py      (DataLab client)
  - section_mapping_20260505/pipeline/stage1_datalab.py  (pipeline flow)
  - section_mapping_20260505/pipeline/stage2_index.py    (normalization)
"""

from __future__ import annotations

import json
import os
import re
import time
from html import unescape as html_unescape
from pathlib import Path
from typing import Any

from core.db import (
    connect,
    ensure_schema,
    insert_block,
    insert_block_heading,
    insert_document,
    insert_section,
    tx,
)

_TAG_RE = re.compile(r"<[^>]+>")


# ---------------------------------------------------------------------------
# HTML/text helpers
# ---------------------------------------------------------------------------

def strip_html(html: str | None) -> str:
    """Strip HTML tags and decode entities, returning plain text."""
    if not html:
        return ""
    return html_unescape(_TAG_RE.sub(" ", html)).strip()


# ---------------------------------------------------------------------------
# DataLab client
# ---------------------------------------------------------------------------

def _load_api_key() -> str:
    """Read DATALAB_API_KEY from env or .env files."""
    for key in ("DATALAB_API_KEY", "DATALABS_API_KEY"):
        v = os.environ.get(key)
        if v:
            return v.strip()

    # Fall back to .env files
    candidates = [
        Path(__file__).resolve().parents[3] / ".env",  # scripts/.env
        Path(__file__).resolve().parent / ".env",       # vision/.env
    ]
    for env_path in candidates:
        if not env_path.exists():
            continue
        for line in env_path.read_text().splitlines():
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            k, _, v = line.partition("=")
            if k.strip() in ("DATALAB_API_KEY", "DATALABS_API_KEY"):
                v = v.strip().strip('"').strip("'")
                if v:
                    return v

    raise SystemExit(
        "ERROR: DATALAB_API_KEY not found in environment or .env file."
    )


def _datalab_convert(
    pdf_path: Path,
    output_format: str = "json",
    mode: str = "accurate",
    page_range: str | None = None,
    max_polls: int = 600,
    poll_interval: int = 2,
) -> dict:
    """Submit a PDF to DataLab and return the parsed result dict.

    Uses the datalab_sdk if available, otherwise the requests-based fallback.
    """
    api_key = _load_api_key()

    try:
        from datalab_sdk import ConvertOptions, DatalabClient
    except ImportError:
        raise SystemExit(
            "ERROR: datalab-python-sdk not installed.\n"
            "  pip install -r scripts/datalab_eval/requirements.txt"
        )

    client = DatalabClient(api_key=api_key)
    opts = ConvertOptions(
        output_format=output_format,
        mode=mode,
        paginate=True,
        page_range=page_range,
    )

    result = client.convert(
        str(pdf_path),
        options=opts,
        max_polls=max_polls,
        poll_interval=poll_interval,
    )

    # Extract the raw dict from the SDK result object
    raw = getattr(result, "json", None) or getattr(result, "json_data", None)
    if raw is None:
        # Try model_dump / dict / to_dict
        for method in ("model_dump", "dict", "to_dict"):
            fn = getattr(result, method, None)
            if callable(fn):
                try:
                    raw = fn()
                    break
                except Exception:
                    continue

    if raw is None:
        raise RuntimeError("Could not extract JSON from DataLab result")

    return raw if isinstance(raw, dict) else json.loads(raw)


# ---------------------------------------------------------------------------
# Normalization — DataLab JSON → PostgreSQL
# ---------------------------------------------------------------------------

def _walk_blocks(pages: list[dict]):
    """Yield every block dict in document order (pages → children)."""
    for page in sorted(pages, key=lambda p: p.get("page", 0)):
        for child in page.get("children") or []:
            yield child


def _resolve_section_title(datalab_id: str, id_to_block: dict) -> str:
    """Resolve a SectionHeader ID to its title text."""
    blk = id_to_block.get(datalab_id)
    if blk and blk.get("block_type") == "SectionHeader":
        return strip_html(blk.get("html"))
    return ""


def _extract_heading_chain(
    heading_registers: dict[int, str | None],
    id_to_block: dict,
) -> list[str]:
    """Build the heading chain array from current heading registers."""
    chain = []
    for level in range(1, 7):
        hid = heading_registers.get(level)
        if hid is None:
            break
        title = _resolve_section_title(hid, id_to_block)
        if title:
            chain.append(title)
        else:
            break
    return chain


def _normalize_datalab_json(
    conn,
    data: dict,
    document_name: str,
    document_id: int,
) -> None:
    """Parse a DataLab JSON dict into the evidence store tables.

    Two passes:
      1. Build the section tree (structural + local SectionHeaders).
      2. Insert blocks, assign to sections, write heading ancestry.
    """
    pages = data.get("children") or []

    # Build block index for title resolution
    id_to_block: dict[str, dict] = {}
    for blk in _walk_blocks(pages):
        bid = blk.get("id")
        if bid:
            id_to_block[bid] = blk

    # ------------------------------------------------------------------ pass 1: sections
    heading_registers: dict[int, str | None] = {i: None for i in range(1, 7)}
    db_section_by_datalab_id: dict[str, int] = {}
    current_section_db_id: int | None = None
    section_page_end: dict[int, int] = {}
    sections_to_update: list[dict] = []

    for blk in _walk_blocks(pages):
        hierarchy = blk.get("section_hierarchy") or {}
        page = blk.get("page", 0)

        # Detect heading register changes
        changed_levels = []
        for level in range(1, 7):
            new_val = hierarchy.get(str(level))
            if new_val != heading_registers[level]:
                changed_levels.append(level)
                heading_registers[level] = new_val

        if changed_levels:
            deepest = max(changed_levels)
            section_datalab_id = heading_registers.get(deepest)
            if section_datalab_id:
                title = _resolve_section_title(section_datalab_id, id_to_block)
                parent_datalab_id = (
                    heading_registers.get(deepest - 1)
                    if deepest > 1 else None
                )
                parent_db_id = (
                    db_section_by_datalab_id.get(parent_datalab_id)
                    if parent_datalab_id else None
                )
                heading_chain = _extract_heading_chain(
                    heading_registers, id_to_block
                )

                section_db_id = insert_section(
                    conn,
                    document_id=document_id,
                    datalab_id=section_datalab_id,
                    parent_id=parent_db_id,
                    heading_level=deepest,
                    title=title,
                    page_start=page,
                    page_end=None,
                    heading_chain=heading_chain,
                )
                db_section_by_datalab_id[section_datalab_id] = section_db_id
                section_page_end[section_db_id] = page
                current_section_db_id = section_db_id
                sections_to_update.append({
                    "db_id": section_db_id,
                    "datalab_id": section_datalab_id,
                })

        # Check for local SectionHeader (not in any hierarchy)
        if blk.get("block_type") == "SectionHeader":
            bid = blk.get("id")
            is_structural = any(
                bid == heading_registers[i] for i in range(1, 7)
            )
            if not is_structural:
                title = strip_html(blk.get("html"))
                if title and title not in ("",):
                    section_db_id = insert_section(
                        conn,
                        document_id=document_id,
                        datalab_id=bid,
                        parent_id=current_section_db_id,
                        heading_level=None,
                        title=title,
                        page_start=page,
                        page_end=None,
                        heading_chain=_extract_heading_chain(
                            heading_registers, id_to_block
                        ),
                    )
                    db_section_by_datalab_id[bid] = section_db_id
                    section_page_end[section_db_id] = page
                    current_section_db_id = section_db_id
                    sections_to_update.append({
                        "db_id": section_db_id,
                        "datalab_id": bid,
                    })

    # --------------------------------------------------------------- pass 2: blocks
    heading_registers = {i: None for i in range(1, 7)}
    current_section_db_id = None
    section_block_counts: dict[int, int] = {}
    section_text_parts: dict[int, list[str]] = {}

    for blk in _walk_blocks(pages):
        hierarchy = blk.get("section_hierarchy") or {}
        page = blk.get("page", 0)
        block_type = blk.get("block_type", "Block")
        datalab_id = blk.get("id", "")
        html = blk.get("html") or ""
        text = strip_html(html)
        bbox = blk.get("bbox")

        # Track heading changes
        for level in range(1, 7):
            new_val = hierarchy.get(str(level))
            if new_val != heading_registers[level]:
                heading_registers[level] = new_val

        # Detect current section from local SectionHeaders
        if block_type == "SectionHeader" and datalab_id in db_section_by_datalab_id:
            is_structural = any(
                datalab_id == heading_registers[i] for i in range(1, 7)
            )
            if not is_structural:
                current_section_db_id = db_section_by_datalab_id[datalab_id]

        # Fall back to deepest structural ancestor
        deepest_structural = None
        for level in range(6, 0, -1):
            hid = heading_registers[level]
            if hid and hid in db_section_by_datalab_id:
                if deepest_structural is None:
                    deepest_structural = db_section_by_datalab_id[hid]

        section_id = current_section_db_id or deepest_structural
        if section_id is None:
            continue  # skip blocks with no section context

        # Insert block
        block_db_id = insert_block(
            conn,
            document_id=document_id,
            datalab_id=datalab_id,
            section_id=section_id,
            block_type=block_type,
            page=page,
            html_content=html,
            text_content=text,
            bbox=tuple(bbox) if bbox else None,
        )

        # Track section metadata
        section_page_end[section_id] = max(
            section_page_end.get(section_id, page), page
        )
        section_block_counts[section_id] = (
            section_block_counts.get(section_id, 0) + 1
        )
        if text:
            section_text_parts.setdefault(section_id, []).append(text)

        # Write heading ancestry
        for level in range(1, 7):
            hid = heading_registers[level]
            if hid and hid in db_section_by_datalab_id:
                insert_block_heading(
                    conn, block_db_id, db_section_by_datalab_id[hid], level,
                    depth=1,  # closest ancestor — depth counting is handled by level
                )

    # ------------------------------------------------------- post-pass: update sections
    with conn.cursor() as cur:
        for sec in sections_to_update:
            sid = sec["db_id"]
            blk_count = section_block_counts.get(sid, 0)
            end_page = section_page_end.get(sid)
            search_text = " ".join(
                section_text_parts.get(sid, [])
            )[:100000]
            cur.execute(
                """UPDATE sections
                   SET block_count = %s, page_end = %s, search_text = %s
                   WHERE id = %s""",
                (blk_count, end_page, search_text, sid),
            )

    total_blocks = sum(section_block_counts.values())
    print(
        f"  sections: {len(sections_to_update)} | "
        f"blocks: {total_blocks}"
    )


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def ingest_pdf(
    case_id: int,
    pdf_path: str | Path,
    document_name: str | None = None,
    mode: str = "accurate",
    page_range: str | None = None,
) -> dict:
    """Ingest a PDF into the evidence store for a given case.

    Args:
        case_id: The case this document belongs to.
        pdf_path: Path to the PDF file.
        document_name: Display name. Defaults to the PDF filename.
        mode: DataLab conversion mode — 'fast', 'balanced', or 'accurate'.
        page_range: Optional page range string (e.g. '0,44,45,46').

    Returns:
        dict with document_id, page_count, section_count, block_count.
    """
    pdf_path = Path(pdf_path).resolve()
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF not found: {pdf_path}")

    if document_name is None:
        document_name = pdf_path.name

    ensure_schema()

    print(f"Ingest: {document_name}")
    print(f"  PDF: {pdf_path}")
    print(f"  Case ID: {case_id}")

    # -- Step 1: Insert placeholder document row (visible immediately) ---------
    with tx() as conn:
        doc_id = insert_document(
            conn,
            case_id=case_id,
            name=document_name,
            page_count=None,     # unknown until OCR completes
            source="user_upload",
        )
        with conn.cursor() as cur:
            cur.execute(
                """UPDATE documents
                   SET ocr_status = 'processing',
                       ocr_provider = 'datalab'
                   WHERE id = %s""",
                (doc_id,),
            )

    # -- Step 2: DataLab OCR ------------------------------------------------
    print("  → DataLab OCR (this may take a minute)...")
    t0 = time.time()
    data = _datalab_convert(
        pdf_path,
        output_format="json",
        mode=mode,
        page_range=page_range,
    )
    elapsed = time.time() - t0
    pages = data.get("children") or []
    page_count = len(pages)
    print(f"  ← DataLab returned {page_count} pages in {elapsed:.1f}s")

    # -- Step 3: Update document with page_count and ocr_status ---------------
    with tx() as conn:
        with conn.cursor() as cur:
            cur.execute(
                """UPDATE documents
                   SET page_count = %s,
                       ocr_status = 'complete',
                       ocr_result_path = NULL
                   WHERE id = %s""",
                (page_count, doc_id),
            )

    # -- Step 4: Normalize into evidence store ------------------------------
    print("  → Normalizing blocks...")
    t0 = time.time()
    with tx() as conn:
        _normalize_datalab_json(
            conn, data, document_name, doc_id,
        )
    elapsed = time.time() - t0
    print(f"  ← Normalized in {elapsed:.1f}s")

    # -- Step 5: Gather stats -----------------------------------------------
    conn = connect()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT count(*) FROM sections WHERE document_id = %s",
                (doc_id,),
            )
            section_count = cur.fetchone()[0]
            cur.execute(
                "SELECT count(*) FROM blocks WHERE document_id = %s",
                (doc_id,),
            )
            block_count = cur.fetchone()[0]
    finally:
        conn.close()

    # -- Step 6: Image with no extractable text → Mistral visual description --
    _IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"}
    if block_count == 0 and pdf_path.suffix.lower() in _IMAGE_EXTENSIONS:
        try:
            description = _describe_image_with_mistral(pdf_path)
            if description:
                with tx() as conn:
                    _normalize_image_description(conn, doc_id, document_name, description)
                # Recount
                conn = connect()
                try:
                    with conn.cursor() as cur:
                        cur.execute("SELECT count(*) FROM sections WHERE document_id = %s", (doc_id,))
                        section_count = cur.fetchone()[0]
                        cur.execute("SELECT count(*) FROM blocks WHERE document_id = %s", (doc_id,))
                        block_count = cur.fetchone()[0]
                finally:
                    conn.close()
                print(f"  → Image described: {len(description)} chars")
            else:
                print("  → Image description returned empty — skipping")
        except Exception as e:
            print(f"  → Image description failed (non-fatal): {e}")

    result = {
        "document_id": doc_id,
        "document_name": document_name,
        "page_count": page_count,
        "section_count": section_count,
        "block_count": block_count,
    }
    print(
        f"  Done: doc_id={doc_id}, "
        f"{page_count} pages, "
        f"{section_count} sections, "
        f"{block_count} blocks"
    )
    return result


def ingest_datalab_json(
    case_id: int,
    json_path: str | Path,
    document_name: str,
    page_count: int | None = None,
) -> dict:
    """Ingest an existing DataLab JSON file directly (skip OCR step).

    Useful for re-processing a previous DataLab run without re-spending
    on OCR. The JSON must be the output of a prior DataLab convert call.
    """
    json_path = Path(json_path).resolve()
    if not json_path.exists():
        raise FileNotFoundError(f"JSON not found: {json_path}")

    data = json.loads(json_path.read_text())
    pages = data.get("children") or []
    if page_count is None:
        page_count = len(pages)

    ensure_schema()

    print(f"Ingest (from JSON): {document_name}")
    print(f"  JSON: {json_path}")
    print(f"  Case ID: {case_id}")
    print(f"  Pages: {page_count}")

    # Insert document
    with tx() as conn:
        doc_id = insert_document(
            conn,
            case_id=case_id,
            name=document_name,
            page_count=page_count,
            source="data_lab",
        )
        with conn.cursor() as cur:
            cur.execute(
                """UPDATE documents
                   SET ocr_status = 'complete',
                       ocr_provider = 'datalab',
                       ocr_result_path = %s
                   WHERE id = %s""",
                (str(json_path), doc_id),
            )

    # Normalize
    print("  → Normalizing blocks...")
    t0 = time.time()
    with tx() as conn:
        _normalize_datalab_json(conn, data, document_name, doc_id)
    elapsed = time.time() - t0
    print(f"  ← Normalized in {elapsed:.1f}s")

    # Stats
    conn = connect()
    try:
        with conn.cursor() as cur:
            cur.execute(
                "SELECT count(*) FROM sections WHERE document_id = %s",
                (doc_id,),
            )
            section_count = cur.fetchone()[0]
            cur.execute(
                "SELECT count(*) FROM blocks WHERE document_id = %s",
                (doc_id,),
            )
            block_count = cur.fetchone()[0]
    finally:
        conn.close()

    result = {
        "document_id": doc_id,
        "document_name": document_name,
        "page_count": page_count,
        "section_count": section_count,
        "block_count": block_count,
    }
    print(
        f"  Done: doc_id={doc_id}, "
        f"{page_count} pages, "
        f"{section_count} sections, "
        f"{block_count} blocks"
    )
    return result


# ---------------------------------------------------------------------------
# DOCX Normalization
# ---------------------------------------------------------------------------

def _normalize_docx(conn, docx_path: Path, document_id: int) -> None:
    """Extract paragraphs and tables from a DOCX into blocks and sections."""
    try:
        from docx import Document
    except ImportError:
        raise SystemExit("ERROR: python-docx not installed.")

    doc = Document(str(docx_path))
    heading_registers: dict[int, int | None] = {i: None for i in range(1, 7)}
    section_block_counts: dict[int, int] = {}
    section_text_parts: dict[int, list[str]] = {}
    sections_inserted: list[int] = []
    current_section_id: int | None = None
    block_seq = 0
    _section_titles: dict[int, str] = {}

    def _heading_level(paragraph) -> int | None:
        style_name = (paragraph.style.name if paragraph.style else "").lower()
        for level in range(1, 7):
            if f"heading {level}" in style_name or f"heading{level}" in style_name:
                return level
        text = paragraph.text.strip()
        if text and len(text) < 120 and text.isupper():
            runs = paragraph.runs
            if runs and any(r.bold for r in runs):
                return 3
        return None

    def _ensure_section(level: int, title: str) -> int:
        nonlocal current_section_id
        parent_id = heading_registers.get(level - 1) if level > 1 else None
        chain = [
            _section_titles.get(heading_registers.get(i, 0), "")
            for i in range(1, level)
            if heading_registers.get(i) is not None
        ] + [title]

        section_id = insert_section(
            conn, document_id=document_id, heading_level=level,
            title=title, page_start=1, page_end=None, heading_chain=chain,
        )
        heading_registers[level] = section_id
        for deeper in range(level + 1, 7):
            heading_registers[deeper] = None
        sections_inserted.append(section_id)
        current_section_id = section_id
        _section_titles[section_id] = title
        return section_id

    root_id = insert_section(
        conn, document_id=document_id, heading_level=0,
        title=docx_path.name, page_start=1, page_end=None,
        heading_chain=[docx_path.name],
    )
    heading_registers[0] = root_id  # type: ignore
    sections_inserted.append(root_id)
    current_section_id = root_id
    _section_titles[root_id] = docx_path.name

    for para in doc.paragraphs:
        text = para.text.strip()
        level = _heading_level(para)
        if level and text:
            sec_id = _ensure_section(level, text)
            _section_titles[sec_id] = text
            insert_block(
                conn, document_id=document_id, section_id=sec_id,
                block_type="SectionHeader", page=1,
                html_content=f"<h{level}>{text}</h{level}>", text_content=text,
            )
            block_seq += 1
        elif text:
            if current_section_id is None:
                current_section_id = root_id
            insert_block(
                conn, document_id=document_id, section_id=current_section_id,
                block_type="Text", page=1,
                html_content=f"<p>{text}</p>", text_content=text,
            )
            block_seq += 1
            section_block_counts[current_section_id] = (
                section_block_counts.get(current_section_id, 0) + 1
            )
            section_text_parts.setdefault(current_section_id, []).append(text)

    for table in doc.tables:
        if current_section_id is None:
            current_section_id = root_id
        rows = [" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        table_html = "<table>" + "".join(
            f"<tr>{''.join(f'<td>{cell.text}</td>' for cell in row.cells)}</tr>"
            for row in table.rows
        ) + "</table>"
        table_text = "\n".join(rows)
        insert_block(
            conn, document_id=document_id, section_id=current_section_id,
            block_type="Table", page=1,
            html_content=table_html, text_content=table_text,
        )
        block_seq += 1
        section_block_counts[current_section_id] = (
            section_block_counts.get(current_section_id, 0) + 1
        )
        section_text_parts.setdefault(current_section_id, []).append(table_text)

    with conn.cursor() as cur:
        for sid in sections_inserted:
            blk_count = section_block_counts.get(sid, 0)
            search_text = " ".join(section_text_parts.get(sid, []))[:100000]
            cur.execute(
                "UPDATE sections SET block_count = %s, page_end = 1, "
                "search_text = %s WHERE id = %s",
                (blk_count, search_text, sid),
            )

    print(f"  sections: {len(sections_inserted)} | blocks: {block_seq}")


# ---------------------------------------------------------------------------
# Markdown Normalization
# ---------------------------------------------------------------------------

_MD_EXTENSIONS = {".md", ".markdown"}

# Regex patterns for markdown parsing
_RE_HEADING = re.compile(r"^(#{1,6})\s+(.+)$")
_RE_LIST_ITEM = re.compile(r"^(\s*)([-*+]|\d+[.)])\s+(.+)$")
_RE_HR = re.compile(r"^(\s*[-*_]\s*){3,}$")
_RE_FRONTMATTER_DELIM = re.compile(r"^---\s*$")

# Block types used in markdown ingestion
_MD_BLOCK_TYPES = {
    "heading": "SectionHeader",
    "paragraph": "Text",
    "list_item": "List",
    "code_block": "Code",
    "blockquote": "Text",
    "divider": "Divider",
}


def _try_parse_frontmatter(lines: list[str]) -> tuple[dict | None, int]:
    """If ``lines`` starts with YAML frontmatter (--- delimited), parse it
    and return the dict + index of the first line after the closing ``---``.
    Otherwise return (None, 0)."""
    if not lines or not _RE_FRONTMATTER_DELIM.match(lines[0]):
        return None, 0
    try:
        import yaml
    except ImportError:
        # No PyYAML — collect raw key-value pairs between --- delimiters
        meta: dict[str, str] = {}
        for i in range(1, len(lines)):
            if _RE_FRONTMATTER_DELIM.match(lines[i]):
                return meta, i + 1
            line = lines[i].strip()
            if ":" in line:
                k, _, v = line.partition(":")
                meta[k.strip()] = v.strip()
        return meta, 0  # never closed — treat whole file as content
    else:
        buf: list[str] = []
        for i in range(1, len(lines)):
            if _RE_FRONTMATTER_DELIM.match(lines[i]):
                try:
                    meta = yaml.safe_load("\n".join(buf))
                    return (meta if isinstance(meta, dict) else None), i + 1
                except Exception:
                    return None, i + 1
            buf.append(lines[i])
        return None, 0


def _heading_level_and_text(line: str) -> tuple[int, str] | None:
    """Return (level, title) if *line* is an ATX heading, else None."""
    m = _RE_HEADING.match(line)
    if not m:
        return None
    return len(m.group(1)), m.group(2).strip()


def _list_item_info(line: str) -> tuple[int, str, str] | None:
    """Return (indent_level, marker, content) if *line* is a list item, else None."""
    m = _RE_LIST_ITEM.match(line)
    if not m:
        return None
    indent = len(m.group(1))
    marker = m.group(2).rstrip(".")
    # Normalize numbered markers to "1.", "2.", etc.
    if marker.isdigit():
        marker = f"{marker}."
    return indent, marker, m.group(3).strip()


def _is_hr(line: str) -> bool:
    """Return True if *line* is a horizontal rule."""
    return bool(_RE_HR.match(line))


def _normalize_markdown(conn, md_path: Path, document_id: int) -> None:
    """Ingest a Markdown file into the evidence store.

    Parses headings into **sections** (with parent-child hierarchy) and
    body content into **blocks**.  Each block is linked to the nearest
    ancestor section via ``block_headings``.
    """
    with open(md_path, encoding="utf-8") as fh:
        raw_lines = fh.readlines()

    lines = [l.rstrip("\n") for l in raw_lines]

    # -- frontmatter ----------------------------------------------------------
    content_start = 0
    frontmatter, content_start = _try_parse_frontmatter(lines)
    if frontmatter:
        # Convert any non-JSON-serializable types (dates, datetimes from YAML)
        def _serialize(v: Any) -> Any:
            if hasattr(v, "isoformat"):
                return v.isoformat()
            if isinstance(v, (set, frozenset)):
                return list(v)
            return v
        safe = {k: _serialize(v) for k, v in frontmatter.items()}
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE documents SET metadata = metadata || %s::jsonb WHERE id = %s",
                (json.dumps({"frontmatter": safe}), document_id),
            )

    # -- state ----------------------------------------------------------------
    sections: list[dict] = []       # {id, heading_level, title}
    heading_registers: dict[int, int] = {}  # level → section_id
    blocks_out: list[dict] = []     # collected for batch insert

    def _current_section_id() -> int | None:
        """Return the most recently added section ID, or None."""
        return sections[-1]["id"] if sections else None

    # helpers
    def _add_section(level: int, title: str) -> int:
        # Parent = most recent section with strictly lower heading level
        parent_id = None
        for s in reversed(sections):
            if s["heading_level"] is not None and s["heading_level"] < level:
                parent_id = s["id"]
                break

        heading_chain = [s["title"] for s in sections if s["heading_level"] and s["heading_level"] <= level]
        heading_chain.append(title)

        sid = insert_section(
            conn, document_id=document_id,
            heading_level=level, title=title,
            page_start=1, page_end=1, block_count=0,
            search_text=title, heading_chain=heading_chain,
        )
        # Update parent link
        if parent_id and parent_id != sid:
            with conn.cursor() as cur:
                cur.execute(
                    "UPDATE sections SET parent_id = %s WHERE id = %s",
                    (parent_id, sid),
                )

        sections.append({"id": sid, "heading_level": level, "title": title})
        heading_registers[level] = sid
        # Clear any deeper heading registers
        for lvl in list(heading_registers.keys()):
            if lvl > level:
                del heading_registers[lvl]

        return sid

    def _add_block(section_id: int, block_type: str, text: str, html: str | None = None, meta: dict | None = None) -> int:
        if html is None:
            html = f"<p>{text}</p>"
        bid = insert_block(
            conn, document_id=document_id, section_id=section_id,
            block_type=block_type, page=1,
            html_content=html, text_content=text,
            metadata=meta,
        )
        # Link to all ancestor sections
        for s in sections:
            if s["heading_level"] is not None:
                depth = heading_registers.get(s["heading_level"])
                if depth:
                    insert_block_heading(
                        conn, block_id=bid, section_id=s["id"],
                        heading_level=s["heading_level"], depth=1,
                    )
        blocks_out.append({"id": bid, "section_id": section_id})
        return bid

    # -- parsing state machine -------------------------------------------------
    i = content_start
    pending_text: list[str] = []     # accumulated paragraph lines
    in_code_block = False
    code_lang: str | None = None
    code_lines: list[str] = []

    def _flush_text():
        nonlocal pending_text
        if not pending_text:
            return
        text = "\n".join(pending_text).strip()
        pending_text.clear()
        if not text:
            return
        # Determine current section
        sid = _current_section_id()
        if sid is None:
            sid = _add_section(1, "Untitled")
        html = f"<p>{text}</p>"
        _add_block(sid, _MD_BLOCK_TYPES["paragraph"], text, html)

    def _flush_code():
        nonlocal in_code_block, code_lang, code_lines
        if not code_lines:
            in_code_block = False
            code_lang = None
            return
        text = "\n".join(code_lines)
        lang_tag = f' class="language-{code_lang}"' if code_lang else ""
        html = f"<pre><code{lang_tag}>{text}</code></pre>"
        meta = {"code_language": code_lang} if code_lang else None
        sid = _current_section_id()
        if sid is None:
            sid = _add_section(1, "Untitled")
        _add_block(sid, _MD_BLOCK_TYPES["code_block"], text, html, meta)
        code_lines.clear()
        code_lang = None
        in_code_block = False

    # -- main loop ------------------------------------------------------------
    while i < len(lines):
        line = lines[i]

        # Code fences
        if line.strip().startswith("```"):
            if not in_code_block:
                _flush_text()
                in_code_block = True
                code_lang = line.strip()[3:].strip() or None
            else:
                _flush_code()
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Blank line → paragraph boundary
        if not line.strip():
            _flush_text()
            i += 1
            continue

        # Horizontal rule
        if _is_hr(line):
            _flush_text()
            sid = _current_section_id() or _add_section(1, "Untitled")
            _add_block(sid, _MD_BLOCK_TYPES["divider"], "---", "<hr/>")
            i += 1
            continue

        # Heading
        heading_info = _heading_level_and_text(line)
        if heading_info:
            _flush_text()
            level, title = heading_info
            sid = _add_section(level, title)
            _add_block(sid, _MD_BLOCK_TYPES["heading"], title, f"<h{level}>{title}</h{level}>")
            i += 1
            continue

        # List item
        list_info = _list_item_info(line)
        if list_info:
            _flush_text()
            indent, marker, content = list_info
            sid = _current_section_id() or _add_section(1, "Untitled")
            html = f"<li>{content}</li>"
            _add_block(
                sid, _MD_BLOCK_TYPES["list_item"], content, html,
                {"list_marker": marker, "list_level": indent // 2 + 1},
            )
            i += 1
            # Check for continuation lines (indented text without list marker)
            while i < len(lines) and lines[i].startswith("  ") and not _RE_LIST_ITEM.match(lines[i]) and not lines[i].strip().startswith("```"):
                # Continuation text — append to previous list item
                extra = lines[i].strip()
                if extra:
                    # Update the last inserted block's text
                    prev = blocks_out[-1]
                    new_text = prev.get("_text", content) + " " + extra
                    prev["_text"] = new_text
                    with conn.cursor() as cur:
                        cur.execute(
                            "UPDATE blocks SET text_content = %s, html_content = %s WHERE id = %s",
                            (new_text, f"<li>{new_text}</li>", prev["id"]),
                        )
                i += 1
            continue

        # Blockquote
        if line.lstrip().startswith("> "):
            _flush_text()
            quote_text = line.lstrip()[2:]
            sid = _current_section_id() or _add_section(1, "Untitled")
            html = f"<blockquote><p>{quote_text}</p></blockquote>"
            _add_block(sid, _MD_BLOCK_TYPES["blockquote"], quote_text, html, {"block_type": "blockquote"})
            i += 1
            continue

        # Regular paragraph text
        pending_text.append(line)
        i += 1

    # -- flush remaining buffers ----------------------------------------------
    if in_code_block:
        _flush_code()
    _flush_text()

    # -- update section block counts ------------------------------------------
    for s in sections:
        bc = sum(1 for b in blocks_out if b["section_id"] == s["id"])
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE sections SET block_count = %s WHERE id = %s",
                (bc, s["id"]),
            )

    print(f"  sections: {len(sections)} | blocks: {len(blocks_out)}")


# ---------------------------------------------------------------------------
# CSV Normalization
# ---------------------------------------------------------------------------

def _normalize_csv(conn, csv_path: Path, document_id: int) -> None:
    """Ingest a CSV as a single-section document with one block per row."""
    import csv
    with open(csv_path, newline="") as f:
        rows = list(csv.reader(f))
    if not rows:
        return
    section_id = insert_section(
        conn, document_id=document_id, heading_level=0,
        title=csv_path.name, page_start=1, page_end=1,
        block_count=len(rows),
        search_text=" ".join(" ".join(row) for row in rows)[:100000],
        heading_chain=[csv_path.name],
    )
    for i, row in enumerate(rows):
        text = " | ".join(row)
        insert_block(
            conn, document_id=document_id, section_id=section_id,
            block_type="SectionHeader" if i == 0 else "Text", page=1,
            text_content=text,
        )
    print(f"  sections: 1 | blocks: {len(rows)}")


# ---------------------------------------------------------------------------
# File-Type Dispatcher
# ---------------------------------------------------------------------------

_DATALAB_EXTENSIONS = {".pdf", ".jpg", ".jpeg", ".png", ".tiff", ".tif", ".bmp"}
_DOCX_EXTENSIONS = {".docx"}
_CSV_EXTENSIONS = {".csv"}
_XLSX_EXTENSIONS = {".xlsx", ".xls"}


def ingest_file(
    case_id: int,
    file_path: str | Path,
    document_name: str | None = None,
    mode: str = "accurate",
) -> dict:
    """Ingest any supported file type. Dispatches to the right extractor."""
    file_path = Path(file_path).resolve()
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    if document_name is None:
        document_name = file_path.name
    suffix = file_path.suffix.lower()
    ensure_schema()

    # DataLab path (PDF + images)
    if suffix in _DATALAB_EXTENSIONS:
        return ingest_pdf(
            case_id=case_id, pdf_path=file_path,
            document_name=document_name, mode=mode,
        )

    # DOCX path
    if suffix in _DOCX_EXTENSIONS:
        print(f"Ingest (DOCX): {document_name}")
        with tx() as conn:
            doc_id = insert_document(conn, case_id=case_id, name=document_name,
                                     page_count=1, source="user_upload")
            with conn.cursor() as cur:
                cur.execute("UPDATE documents SET ocr_status='complete', ocr_provider='python-docx' WHERE id=%s", (doc_id,))
        t0 = time.time()
        with tx() as conn:
            _normalize_docx(conn, file_path, doc_id)
        print(f"  ← Extracted in {time.time() - t0:.1f}s")
        conn = connect()
        try:
            with conn.cursor() as cur:
                cur.execute("SELECT count(*) FROM sections WHERE document_id=%s", (doc_id,)); sc = cur.fetchone()[0]
                cur.execute("SELECT count(*) FROM blocks WHERE document_id=%s", (doc_id,)); bc = cur.fetchone()[0]
        finally:
            conn.close()
        result = {"document_id": doc_id, "document_name": document_name, "page_count": 1, "section_count": sc, "block_count": bc}
        print(f"  Done: doc_id={doc_id}, {sc} sections, {bc} blocks")
        return result

    # CSV path
    if suffix in _CSV_EXTENSIONS:
        print(f"Ingest (CSV): {document_name}")
        with tx() as conn:
            doc_id = insert_document(conn, case_id=case_id, name=document_name, page_count=1, source="user_upload")
            with conn.cursor() as cur:
                cur.execute("UPDATE documents SET ocr_status='complete', ocr_provider='csv-parser' WHERE id=%s", (doc_id,))
        with tx() as conn:
            _normalize_csv(conn, file_path, doc_id)
        conn = connect()
        try:
            with conn.cursor() as cur:
                cur.execute("SELECT count(*) FROM blocks WHERE document_id=%s", (doc_id,)); bc = cur.fetchone()[0]
        finally:
            conn.close()
        result = {"document_id": doc_id, "document_name": document_name, "page_count": 1, "section_count": 1, "block_count": bc}
        print(f"  Done: doc_id={doc_id}, {bc} blocks")
        return result

    # Markdown path
    if suffix in _MD_EXTENSIONS:
        print(f"Ingest (Markdown): {document_name}")
        with tx() as conn:
            doc_id = insert_document(conn, case_id=case_id, name=document_name,
                                     page_count=1, source="user_upload")
            with conn.cursor() as cur:
                cur.execute("UPDATE documents SET ocr_status='complete', ocr_provider='markdown-splitter' WHERE id=%s", (doc_id,))
        t0 = time.time()
        with tx() as conn:
            _normalize_markdown(conn, file_path, doc_id)
        print(f"  ← Extracted in {time.time() - t0:.1f}s")
        conn = connect()
        try:
            with conn.cursor() as cur:
                cur.execute("SELECT count(*) FROM sections WHERE document_id=%s", (doc_id,)); sc = cur.fetchone()[0]
                cur.execute("SELECT count(*) FROM blocks WHERE document_id=%s", (doc_id,)); bc = cur.fetchone()[0]
        finally:
            conn.close()
        result = {"document_id": doc_id, "document_name": document_name, "page_count": 1, "section_count": sc, "block_count": bc}
        print(f"  Done: doc_id={doc_id}, {sc} sections, {bc} blocks")
        return result

    # XLSX path
    if suffix in _XLSX_EXTENSIONS:
        print(f"Ingest (XLSX): {document_name}")
        try:
            import openpyxl
        except ImportError:
            raise SystemExit("ERROR: openpyxl not installed.")
        wb = openpyxl.load_workbook(file_path, read_only=True)
        total_blocks, doc_id = 0, None
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = [[str(cell.value or "") for cell in row] for row in ws.iter_rows()]
            if not rows:
                continue
            with tx() as conn:
                if doc_id is None:
                    doc_id = insert_document(conn, case_id=case_id, name=document_name, page_count=len(wb.sheetnames), source="user_upload")
                    with conn.cursor() as cur:
                        cur.execute("UPDATE documents SET ocr_status='complete', ocr_provider='openpyxl' WHERE id=%s", (doc_id,))
                section_id = insert_section(conn, document_id=doc_id, heading_level=0, title=sheet_name, page_start=1, page_end=1, block_count=len(rows), search_text=" ".join(" ".join(r) for r in rows)[:100000], heading_chain=[sheet_name])
                for i, row in enumerate(rows):
                    insert_block(conn, document_id=doc_id, section_id=section_id, block_type="SectionHeader" if i == 0 else "Text", page=1, text_content=" | ".join(row))
                total_blocks += len(rows)
            print(f"  Sheet '{sheet_name}': {len(rows)} rows")
        wb.close()
        conn = connect()
        try:
            with conn.cursor() as cur:
                cur.execute("SELECT count(*) FROM sections WHERE document_id=%s", (doc_id,)); sc = cur.fetchone()[0]
        finally:
            conn.close()
        result = {"document_id": doc_id, "document_name": document_name, "page_count": len(wb.sheetnames), "section_count": sc, "block_count": total_blocks}
        print(f"  Done: doc_id={doc_id}, {sc} sections, {total_blocks} blocks")
        return result

    # Audio path
    if suffix in _AUDIO_EXTENSIONS:
        return ingest_audio(case_id=case_id, audio_path=file_path,
                            document_name=document_name)

    raise ValueError(
        f"Unsupported file type: {suffix}. "
        f"Supported: {sorted(_DATALAB_EXTENSIONS | _DOCX_EXTENSIONS | _CSV_EXTENSIONS | _XLSX_EXTENSIONS | _MD_EXTENSIONS | _AUDIO_EXTENSIONS)}"
    )


# ---------------------------------------------------------------------------
# ---------------------------------------------------------------------------
# Image Description (Mistral vision)
# ---------------------------------------------------------------------------

def _load_mistral_key() -> str:
    """Load MISTRAL_API_KEY from env or .env files."""
    key = os.environ.get("MISTRAL_API_KEY")
    if key:
        return key.strip()
    for env_path in [
        Path(__file__).resolve().parents[3] / ".env",
        Path(__file__).resolve().parents[3] / "mcp-server" / ".env",
    ]:
        if env_path.exists():
            for line in env_path.read_text().splitlines():
                if line.startswith("MISTRAL_API_KEY="):
                    key = line.split("=", 1)[1].strip().strip('"').strip("'")
                    if key:
                        return key
    raise RuntimeError("MISTRAL_API_KEY not found in env or .env files")


def _describe_image_with_mistral(image_path: Path) -> str | None:
    """Return a textual description of an image using Mistral's vision model.

    Encodes the image as base64, sends to pixtral-large, returns the
    assistant's content text. Returns None on failure.
    """
    import base64
    from mistralai.client import Mistral

    mistral_key = _load_mistral_key()
    client = Mistral(api_key=mistral_key)

    suffix = image_path.suffix.lower()
    mime_map = {
        ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
        ".png": "image/png", ".tiff": "image/tiff",
        ".tif": "image/tiff", ".bmp": "image/bmp",
    }
    mime_type = mime_map.get(suffix, "image/jpeg")

    with open(image_path, "rb") as f:
        img_b64 = base64.b64encode(f.read()).decode("utf-8")

    data_url = f"data:{mime_type};base64,{img_b64}"

    print("    → Sending to Mistral vision (pixtral-large)...")
    resp = client.chat.complete(
        model="pixtral-large-latest",
        messages=[{
            "role": "user",
            "content": [
                {
                    "type": "text",
                    "text": (
                        "Describe this image in detail. Include any text visible in "
                        "the image, the setting, people, objects, actions, and any "
                        "other relevant details that would help someone understand "
                        "what this image contains without seeing it."
                    ),
                },
                {"type": "image_url", "image_url": data_url},
            ],
        }],
        max_tokens=1000,
    )

    content = resp.choices[0].message.content
    if isinstance(content, list):
        content = " ".join(
            c.get("text", "") if isinstance(c, dict) else str(c)
            for c in content
        )
    return str(content).strip() if content else None


def _normalize_image_description(
    conn, document_id: int, document_name: str, description: str,
) -> None:
    """Insert a text description of an image into the evidence store.

    Creates one section spanning page 1 with a single Text block containing
    the Mistral-generated description. This makes image-only documents
    searchable and readable by the agent.
    """
    section_id = insert_section(
        conn,
        document_id=document_id,
        heading_level=0,
        title=f"Visual Description — {document_name}",
        page_start=1,
        page_end=1,
        search_text=description[:100000],
        heading_chain=[f"Visual Description — {document_name}"],
    )

    insert_block(
        conn,
        document_id=document_id,
        section_id=section_id,
        block_type="Text",
        page=1,
        html_content=f"<p>{description}</p>",
        text_content=description,
        metadata={"source": "mistral-pixtral-vision"},
    )


# ---------------------------------------------------------------------------
# Audio Transcription (Mistral voxtral)
# ---------------------------------------------------------------------------

_AUDIO_EXTENSIONS = {".m4a", ".mp3", ".wav", ".ogg", ".flac", ".webm", ".mp4"}


def _normalize_audio(conn, audio_path: Path, document_id: int, transcript_text: str,
                     segments: list[dict] | None = None) -> None:
    """Ingest a transcript into the evidence store.

    Creates one section for the full transcript. If the transcription API
    returns timestamped segments, each segment becomes a block. Otherwise,
    splits on paragraph breaks, then on sentence boundaries for reasonable
    block granularity.
    """
    section_id = insert_section(
        conn, document_id=document_id, heading_level=0,
        title=audio_path.name, page_start=1, page_end=1,
        search_text=transcript_text[:100000],
        heading_chain=[audio_path.name],
    )

    block_count = 0

    if segments:
        for i, seg in enumerate(segments):
            start = seg.get("start", 0)
            end = seg.get("end", 0)
            text = seg.get("text", "").strip()
            if not text:
                continue
            timestamp = f"[{int(start//60):02d}:{int(start%60):02d}]"
            insert_block(
                conn, document_id=document_id, section_id=section_id,
                block_type="Text", page=1,
                html_content=f'<p data-start="{start}" data-end="{end}">{timestamp} {text}</p>',
                text_content=f"{timestamp} {text}",
                metadata={"start": start, "end": end, "segment_index": i},
            )
            block_count += 1
    else:
        # No segments — split by paragraphs, then by sentences
        paragraphs = [p.strip() for p in transcript_text.split("\n\n") if p.strip()]
        if not paragraphs:
            paragraphs = [transcript_text.strip()]

        # If a paragraph is long (>500 chars), split into sentences
        import re
        sentence_re = re.compile(r'(?<=[.!?])\s+(?=[A-Z])')
        blocks_out = []
        for para in paragraphs:
            if len(para) > 500:
                sentences = sentence_re.split(para)
                blocks_out.extend(s.strip() for s in sentences if s.strip())
            else:
                blocks_out.append(para)

        for i, text in enumerate(blocks_out):
            insert_block(
                conn, document_id=document_id, section_id=section_id,
                block_type="Text", page=1,
                html_content=f"<p>{text}</p>",
                text_content=text,
                metadata={"chunk_index": i},
            )
            block_count = len(blocks_out)

    with conn.cursor() as cur:
        cur.execute(
            "UPDATE sections SET block_count = %s WHERE id = %s",
            (block_count, section_id),
        )
    print(f"  sections: 1 | blocks: {block_count}")


def ingest_audio(
    case_id: int,
    audio_path: str | Path,
    document_name: str | None = None,
) -> dict:
    """Transcribe an audio file via Mistral voxtral and ingest into evidence store.

    Supports: m4a, mp3, wav, ogg, flac, webm, mp4.
    Requires: MISTRAL_API_KEY in environment or .env.
    Pattern ported from agent_workspace_prototype/scripts/process_intake.py
    """
    audio_path = Path(audio_path).resolve()
    if not audio_path.exists():
        raise FileNotFoundError(f"Audio file not found: {audio_path}")
    if document_name is None:
        document_name = audio_path.name

    # Load Mistral API key (same pattern as OCR)
    api_key = _load_api_key()
    # But for audio we need MISTRAL_API_KEY specifically (not DATALAB)
    mistral_key = os.environ.get("MISTRAL_API_KEY")
    if not mistral_key:
        for env_path in [
            Path(__file__).resolve().parents[3] / ".env",
            Path(__file__).resolve().parents[3] / "mcp-server" / ".env",
        ]:
            if env_path.exists():
                for line in env_path.read_text().splitlines():
                    if line.startswith("MISTRAL_API_KEY="):
                        mistral_key = line.split("=", 1)[1].strip().strip('"').strip("'")
                        break
        if not mistral_key:
            raise RuntimeError("MISTRAL_API_KEY not found in env or .env files")

    print(f"Ingest (Audio): {document_name}")
    print(f"  File: {audio_path} ({audio_path.stat().st_size / 1024 / 1024:.1f}MB)")

    print("  → Transcribing via Mistral voxtral...")
    t0 = time.time()

    try:
        from mistralai.client import Mistral
    except ImportError:
        raise SystemExit(
            "ERROR: mistralai package not installed.\n"
            "  pip install mistralai"
        )

    client = Mistral(api_key=mistral_key)
    file_id = None

    try:
        # 1. Upload audio file
        print("    Uploading...")
        with open(audio_path, "rb") as f:
            uploaded = client.files.upload(
                file={"file_name": audio_path.name, "content": f},
                purpose="audio",
            )
        file_id = uploaded.id

        # 2. Get signed URL
        signed = client.files.get_signed_url(file_id=file_id)

        # 3. Transcribe
        print("    Transcribing...")
        transcription = client.audio.transcriptions.complete(
            model="voxtral-mini-latest",
            file_url=signed.url,
        )

        full_text = transcription.text
        # voxtral returns segments if available
        segments = []
        if hasattr(transcription, "segments") and transcription.segments:
            segments = [
                {"start": s.start, "end": s.end, "text": s.text}
                for s in transcription.segments
            ]

        elapsed = time.time() - t0
        lang = getattr(transcription, "language", "unknown")
        print(
            f"  ← Transcribed in {elapsed:.1f}s — "
            f"{len(full_text)} chars, {len(segments)} segments, lang={lang}"
        )

    finally:
        # Clean up uploaded file
        if file_id:
            try:
                client.files.delete(file_id=file_id)
            except Exception:
                pass

    # Insert document
    with tx() as conn:
        doc_id = insert_document(
            conn, case_id=case_id, name=document_name,
            page_count=1, source="user_upload",
        )
        with conn.cursor() as cur:
            cur.execute(
                "UPDATE documents SET ocr_status = 'complete', "
                "ocr_provider = 'voxtral-mini' WHERE id = %s",
                (doc_id,),
            )

    # Normalize
    print("  → Indexing transcript...")
    with tx() as conn:
        _normalize_audio(conn, audio_path, doc_id, full_text, segments)

    # Stats
    conn = connect()
    try:
        with conn.cursor() as cur:
            cur.execute("SELECT count(*) FROM blocks WHERE document_id = %s", (doc_id,))
            bc = cur.fetchone()[0]
    finally:
        conn.close()

    result = {
        "document_id": doc_id, "document_name": document_name,
        "page_count": 1, "section_count": 1, "block_count": bc,
        "duration_seconds": segments[-1]["end"] if segments else None,
        "transcript_length": len(full_text),
    }
    print(f"  Done: doc_id={doc_id}, {bc} blocks, {len(full_text)} chars")
    return result
